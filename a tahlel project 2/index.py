from datetime import date, datetime
from os import path, system, remove
from sys import argv
import re
import random
from MySQLdb import connect
# from PyQt5.QtCore import Qt,QStringListModel
# from PyQt5.QtGui import *
# from PyQt5.QtWidgets import (QComboBox,QDateEdit,QLineEdit,QListWidget,QPushButton,QSpinBox,QTableWidgetItem,QWidget,QDoubleSpinBox,QListWidgetItem,QCompleter,QMainWindow,QMessageBox,QApplication)
# from PyQt5.uic import loadUiType
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
from docx.shared import RGBColor
from smtplib import SMTP
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from win32com import client
import searchDialog
import searchDialog2
import multyclass
import time
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import add_buys_completepy
import add_delete_category_dialogpy
import add_delete_analyst_choices
import add_delete_analyst_choicesrupy2
import mandobuipy
import NormalDialogClass
import NormalTextDialogClass
import EditClientClass
from mymain import Ui_MainWindow as main_wind
from ast import literal_eval
import HandelPriceClass
show_clients_check = False
# main_wind, _ = loadUiType("design.ui")
user_id = ''
analyst_name_for_update_before = ''
analyst_name_for_update = ''
client_id_glob = 0
chick_if_add_new = False
if_print = False
analysts_name_glo = []
clients_name_glo = []
clients_name_glo2 = []
clients_name_glo_clients_page = []
from_start = False
addTrue = True
select_by_date = False
echo_mode_num = 0
echo_mode_num2 = 0
global_from_save = None
edit_employee_check = True
Edit_employee = False
CJ = None
search_info_by_date = False
all_doctors = []
Edit_Doctor = True
Delete_Doctor = True
current_group_box = ''
word_files = None
save_word_files = None
MAX_ID = 0

class mainapp(QMainWindow, main_wind):
	def __init__(self, parent=None):
		global from_start
		global word_files
		global save_word_files
		global word_data
		super(mainapp, self).__init__(parent)
		QMainWindow.__init__(self)
		self.setupUi(self)
		self.tabWidget.tabBar().setVisible(False)
		self.DB()
		self.Show_default_statics()
		self.Delete_Files()
		self.add_clients_to_combo()
		from_start = True
		self.Show_All_The_Analysts()
		self.handel_buttons()
		# self.Show_All_The_Sales()
		self.Show_all_analysts_in_combo()
		self.Show_all_buys()
		# self.History()
		self.groupBox.setEnabled(False)
		self.Show_paths()
		# self.add_Analyst_to_list()
		self.Auto_complete_combo()
		self.add_client_to_list()
		self.Auto_complete_combo2()
		self.add_client_to_list4()
		self.Auto_complete_combo4()
		# self.add_today_client_to_list()
		# self.Auto_complete_combo7()
		self.Show_Word_Doc_Data()
		self.Add_Doctor_Data()
		self.dateEdit_6.setDate(date.today())
		self.dateEdit_5.setDate(date.today())
		# self.dateEdit.setDate(datetime.date.today())
		self.tableWidget_5.setColumnHidden(4, True)
		self.tableWidget_7.setColumnHidden(0, True)
		self.tableWidget_7.setColumnHidden(3, True)
		self.tableWidget_3.setColumnHidden(0, True)
		self.cur.execute(''' SELECT * FROM paths WHERE id=1 ''')
		mydata = self.cur.fetchone()
		word_files = mydata[1]
		save_word_files = mydata[2]
		self.cur.execute(''' SELECT * FROM  word WHERE id=1''')
		word_data = self.cur.fetchone()
		self.clear_data_in_sales()
		self.tableWidget_7.resizeColumnsToContents()
		self.tableWidget_2.resizeColumnsToContents()
		self.Show_all_human_type_in_combos()
		# data = [(5,'Random  blood sugar',2,'حقل كتابة','Biochemistry','2021-05-11 18:39:48','mg / dl','( 80 - 140 )','',None,1),(6,'Blood Urea',3,'حقل كتابة','Biochemistry','2021-02-22 23:11:50',' mg / dl','( 20 - 45 )',None,None,2),(7,'S. Creatinin',3,'حقل كتابة','Biochemistry','2021-02-22 23:13:53','mg / dl','( 0.7 - 1.4 )',None,None,3),(8,'S. Uric acid',3,'حقل كتابة','Biochemistry','2021-06-06 00:07:15','mg/dl','( 3-7 )','',None,4),(9,'S. Cholesterol',3,'حقل كتابة','Biochemistry','2021-02-22 23:15:47','mg / dl','( 150 - 250 )',None,None,5),(10,'S. Triglycerid',3,'حقل كتابة','Biochemistry','2021-02-22 23:16:20','mg / dl','( 65 - 180 )',None,None,6),(11,'Total serum Bilirubin',3,'حقل كتابة','Biochemistry','2021-02-22 23:17:26','mg / dl','( 0.3 - 1.0 )',None,6,7),(12,'S.Calcium',3,'حقل كتابة','Biochemistry','2021-02-22 23:18:14','mg / dl','( 8.8 - 10.2 )',None,None,8),(13,'Vitamin D',15,'حقل كتابة','Biochemistry','2021-02-22 23:18:58','ng / dl','( 30 - 70 )',None,None,9),(14,'Color ',3,'خيارات','General Stool Examination','2021-06-12 10:59:48',' ',' ','\'milky\', \'yallow\', \'brown\', \'green\', \'\'',None,1),(15,'Consistency',0,'خيارات','General Stool Examination','2021-06-08 22:19:45','','','\'Solid\', \'Liquid\', \'Semi solid\', \'Semi liquid\', \'Mucoid\', \'  \'',None,2),(16,'R.B.Cs',0,'خيارات','General Stool Examination','2021-06-10 17:25:44','/ H.P.F','','\'Nil\', \'0-1\', \'0-2\', \'1-2\', \'1-3\', \'2-4\', \'3-5\', \'4-6\', \'6-8\', \'Few\', \'+\', \'++\', \'+++\', \'++++\', \'\'',None,3),(18,'E.Histolytica',0,'خيارات','General Stool Examination','2021-06-10 22:38:57','','','\'Nil\', \'Cyst\', \'Trophozoite\', \'\'',None,5),(19,'G.Lamblia',0,'خيارات','General Stool Examination','2021-06-10 22:40:16','','','\'Nil\', \'Cyst\', \'Trophozoite\', \'\'',None,6),(22,'Appearance',3,'خيارات','General Urine Examination','2021-02-22 23:26:16','','','\'Turbid\', \'Clear\', \' \'',None,1),(23,'Reaction.',0,'خيارات','General Urine Examination','2021-06-12 17:05:55','','','\'Acidic\', \' Alkaline\', \'\'',None,2),(24,'Albumin',0,'خيارات','General Urine Examination','2021-02-22 23:26:57','','','\'Nil\', \'+\', \'++\', \'+++\', \'Trace\', \' \'',None,3),(25,'Pus cells..',0,'خيارات','General Urine Examination','2021-06-10 17:52:18','/ H.P.F','','\'Nil\', \'0-1\', \'0-2\', \'1-2\', \'2-3\', \'3-5\', \'4-6\', \'6-8\', \'+\', \'++\', \'+++\', \'++++\', \'\', \'\'',None,7),(26,'RBCs',0,'خيارات','General Urine Examination','2021-06-10 17:27:28','/ H.P.F','','\'Nil\', \'0-1\', \'0-2\', \'1-2\', \'1-3\', \'2-3\', \'2-4\', \'3-5\', \'4-6\', \'6-8\', \'Few\', \'+\', \'++\', \'+++\', \'++++\', \'\'',None,6),(27,'Pus cells,',0,'خيارات','General Stool Examination','2021-06-10 17:23:42','/ H.P.F','','\'Nil\', \'0-1\', \'0-2\', \'1-2\', \'2-3\', \'3-5\', \'4-6\', \'5-7\', \'6-8\', \'+\', \'++\', \'+++\', \'++++\', \'\'',None,4),(28,'Epith .cells',0,'خيارات','General Urine Examination','2021-06-10 17:17:20','/ H.P.F','','\'Nil\', \'Few\', \'+\', \'++\', \'+++\', \'++++\', \'\'',None,8),(29,'Crystals',0,'خيارات','General Urine Examination','2021-06-12 16:47:15','','','\'Nil\', \'Am.Urate few\', \'Am.Urate +\', \'Am.Urate ++\', \'Am.Urate +++\', \'Am.Urate ++++\', \' Ca.Oxalate few\', \' Ca.Oxalate +\', \' Ca.Oxalate ++\', \' Ca.Oxalate +++\', \' Ca.Oxalate ++++\', \' Uric Acid few\', \' Uric Acid +\', \' Uric Acid ++\', \' Uric Acid +++\', \' Uric Acid ++++\', \' Am.Phosphatase +\', \' Am.Phosphatase ++\', \' Am.Phosphatase +++\', \' Am.Phosphatase ++++\', \' Am.Phosphatase few\', \'\'',None,9),(30,'Casts',0,'خيارات','General Urine Examination','2021-06-10 17:15:26','','','\'Nil\', \'Granular cast +\', \'Granular cast ++\', \'Granular cast +++\', \'\'',None,10),(32,'Hb',3,'حقل كتابة',' Hematology and Serology','2021-02-22 23:31:12','gm/dl','','4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, \' \'',None,7),(33,'PCV',0,'حقل كتابة',' Hematology and Serology','2021-07-04 18:43:20','%','','',None,8),(34,'WBCs',3,'حقل كتابة',' Hematology and Serology','2021-07-04 18:51:14','cells/cumm','','',None,8),(35,'E.S.R',3,'حقل كتابة',' Hematology and Serology','2021-02-22 23:32:33','mm/1 hr','','',None,10),(36,'Blood Group',3,'خيارات',' Hematology and Serology','2021-06-12 10:13:41',' ',' ','\'\', \'A (+ve)\', \'B (+ve)\', \'AB (+ve)\', \'O (+ve)\', \'O (-ve)\', \'A (-ve)\', \'B (-ve)\', \'AB (-ve)\', \'\'',None,1),(37,'Rh',0,'خيارات',' Hematology and Serology','2021-06-08 22:59:07','','','\' \', \'Positive (+ve)\', \'Negative (-ve)\', \'\'',None,2),(38,'Pregnancy test  in urine',3,'خيارات',' Hematology and Serology','2021-06-11 21:02:36','','','\'\', \'Positive (+ve)\', \'Negative (-ve)\', \'Weak Positive\', \'\'',None,11),(39,'Pregnancy test  in serum',3,'خيارات',' Hematology and Serology','2021-07-05 17:32:51','','','\'\', \'Positive (+ve)\', \'Negative (-ve)\', \'Weak Positive\', \'\'',None,1),(40,'R.B.Sugar',2,'حقل كتابة',' Hematology and Serology','2021-02-22 23:35:07','mg/dl','',None,None,3),(41,'Bl. Urea',3,'حقل كتابة',' Hematology and Serology','2021-02-22 23:35:28','mg/dl','',None,None,22),(42,'Salmonella typhi  IgG',4,'خيارات',' Hematology and Serology','2021-06-12 10:15:04','','','\'\', \'Positive (+ve)\', \'Negative (-ve)\', \'\'',None,13),(43,'Salmonella typhi  IgM',0,'خيارات',' Hematology and Serology','2021-06-12 10:15:38','','','\'\', \'Positive (+ve)\', \'Negative (-ve)\', \'\'',None,14),(44,'Rose-Bengal test',3,'خيارات',' Hematology and Serology','2021-06-12 10:16:21','','','\'\', \'Positive (+ve)\', \'Negative (-ve)\', \'\'',None,15),(45,'T3',7,'حقل كتابة','Hormones and Viruses','2021-02-23 23:10:14','ng / ml','( 0.6 - 1.85 )',None,None,13),(46,'T4',7,'حقل كتابة','Hormones and Viruses','2021-02-23 23:11:06','g / dl','Females (4.8 - 12.0 )  males (4.4 - 10.8 )',None,None,14),(47,'TSH',7,'حقل كتابة','Hormones and Viruses','2021-02-23 23:11:36','IU /ml','(0.4 - 7.0 )',None,None,15),(48,'LH',10,'حقل كتابة','Hormones and Viruses','2021-02-23 23:12:03','m Iu/ml','',None,None,16),(49,'FSH',10,'حقل كتابة','Hormones and Viruses','2021-02-23 23:12:41','m IU/ml','',None,None,17),(50,'Prolactin',15,'حقل كتابة','Hormones and Viruses','2021-02-23 23:13:23','ng / ml','women-non gestation 4.5-25 men ( 3.7 - 17.5 )',None,None,18),(51,'Testosterone',10,'حقل كتابة','Hormones and Viruses','2021-06-08 23:31:03','ng/ml','','',None,19),(52,'Toxoplasma IgG',25,'خيارات','Hormones and Viruses','2021-06-12 09:38:40','','','\'\', \'Negative (-ve)\', \'Positive (+ve)\', \'0.5 Negative\', \'0.6 Negative\', \'0.7 Negative\', \'0.8 Negative\', \'1.1 Positive\', \'1.1 Positive\', \'1.2 Positive\', \'1.3 Positive\', \'1.4 Positive\', \'1.5 Positive\', \'\'',None,1),(53,'Toxoplasma IgM',0,'خيارات','Hormones and Viruses','2021-06-12 09:44:10','','','\'\', \'Negative (-ve)\', \'Positive (+ve)\', \'0.5 Negative\', \'0.6 Negative\', \'0.7 Negative\', \'0.8 Negative\', \'1.1 Positive\', \'1.1 Positive\', \'1.2 Positive\', \'1.3 Positive\', \'1.4 Positive\', \'1.5 Positive\', \'\'',None,2),(54,'Cytomegalo Virus IgG',0,'خيارات','Hormones and Viruses','2021-06-12 09:47:23','','','\'\', \'Negative (-ve)\', \'Positive (+ve)\', \'0.5 Negative\', \'0.6 Negative\', \'0.7 Negative\', \'0.8 Negative\', \'1.1 Positive\', \'1.1 Positive\', \'1.2 Positive\', \'1.3 Positive\', \'1.4 Positive\', \'1.5 Positive\', \'\'',None,3),(55,'Cytomegalo Virus IgM',0,'خيارات','Hormones and Viruses','2021-06-12 09:50:06','','','\'\', \'Negative (-ve)\', \'Positive (+ve)\', \'0.5 Negative\', \'0.6 Negative\', \'0.7 Negative\', \'0.8 Negative\', \'1.1 Positive\', \'1.1 Positive\', \'1.2 Positive\', \'1.3 Positive\', \'1.4 Positive\', \'1.5 Positive\', \'\'',None,4),(56,'Rubella IgG',0,'خيارات','Hormones and Viruses','2021-06-12 09:54:39','','','\'\', \'Negative (-ve)\', \'Positive (+ve)\', \'0.5 Negative\', \'0.6 Negative\', \'0.7 Negative\', \'0.8 Negative\', \'1.1 Positive\', \'1.1 Positive\', \'1.2 Positive\', \'1.3 Positive\', \'1.4 Positive\', \'1.5 Positive\', \'\'',None,5),(57,'Rubella IgM',0,'خيارات','Hormones and Viruses','2021-06-12 09:57:15','','','\'\', \'Negative (-ve)\', \'Positive (+ve)\', \'0.5 Negative\', \'0.6 Negative\', \'0.7 Negative\', \'0.8 Negative\', \'1.1 Positive\', \'1.1 Positive\', \'1.2 Positive\', \'1.3 Positive\', \'1.4 Positive\', \'1.5 Positive\', \'\'',None,6),(58,'Anti - Phspholipin IgG',35,'خيارات','Hormones and Viruses','2021-02-23 23:41:17','','','\'0.5 Negative\', \'0.6 Negative\', \'0.7 Negative\', \'0.8 Negative\', \'1.1 Positive\', \'1.1 Positive\',\'1.2 Positive\', \'1.3 Positive\', \'1.4 Positive\', \'1.5 Positive\', \' \'',None,9),(59,'Anti - Phspholipin  IgM',0,'خيارات','Hormones and Viruses','2021-02-23 23:41:44',' ',' ','\'0.5 Negative\', \'0.6 Negative\', \'0.7 Negative\', \'0.8 Negative\', \'1.1 Positive\', \'1.1 Positive\',\'1.2 Positive\', \'1.3 Positive\', \'1.4 Positive\', \'1.5 Positive\', \' \'',None,10),(60,'Anti - Cardiolipin  IgG',0,'خيارات','Hormones and Viruses','2021-02-23 23:41:58','','','\'0.5 Negative\', \'0.6 Negative\', \'0.7 Negative\', \'0.8 Negative\', \'1.1 Positive\', \'1.1 Positive\',\'1.2 Positive\', \'1.3 Positive\', \'1.4 Positive\', \'1.5 Positive\', \' \'',None,11),(61,'Anti - Cardiolipin  IgM',0,'خيارات','Hormones and Viruses','2021-02-23 23:42:14',' ',' ','\'0.5 Negative\', \'0.6 Negative\', \'0.7 Negative\', \'0.8 Negative\', \'1.1 Positive\', \'1.1 Positive\',\'1.2 Positive\', \'1.3 Positive\', \'1.4 Positive\', \'1.5 Positive\', \' \'',None,12),(62,'Herps   IgG',0,'خيارات','Hormones and Viruses','2021-06-12 09:59:58','','','\'\', \'Negative (-ve)\', \'Positive (+ve)\', \'0.5 Negative\', \'0.6 Negative\', \'0.7 Negative\', \'0.8 Negative\', \'1.1 Positive\', \'1.1 Positive\', \'1.2 Positive\', \'1.3 Positive\', \'1.4 Positive\', \'1.5 Positive\', \'\'',None,7),(63,'Herpes  IgM',0,'خيارات','Hormones and Viruses','2021-06-12 10:02:40','','','\'\', \'Negative (-ve)\', \'Positive (+ve)\', \'0.5 Negative\', \'0.6 Negative\', \'0.7 Negative\', \'0.8 Negative\', \'1.1 Positive\', \'1.1 Positive\', \'1.2 Positive\', \'1.3 Positive\', \'1.4 Positive\', \'1.5 Positive\', \'\'',None,8),(65,'Volume',3,'خيارات','Seminal Fluid Analysis','2021-06-12 10:48:06','ml','','\' \', \'1\', \'2\', \'3\', \'4\', \'5\', \'0.2\', \'0.3\', \'0.3\', \'0.4\', \'0.5\', \'0.6\', \'0.7\', \'0.8\', \'0.9\', \'\'',None,1),(66,'Reaction',0,'خيارات','Seminal Fluid Analysis','2021-06-08 23:22:40','','','\'Acidic\', \' Alkaline\', \' \'',None,2),(68,'Liquefaction',0,'خيارات','Seminal Fluid Analysis','2021-06-12 10:51:41','min.','','\'30\', \'5\', \'10\', \'15\', \'20\', \'25\', \'35\', \'40\', \'45\', \' \'',None,1),(69,'Count',0,'حقل كتابة','Seminal Fluid Analysis',None,'million/ml','',None,None,5),(70,'Motility:Active',0,'خيارات','Seminal Fluid Analysis','2021-05-21 18:41:52','%','','\'  10\', \'  15\', \'  20\', \'  25\', \'  30\', \'  35\', \'  40\', \'  45\', \'  50\', \'  55\', \'  60\', \'  65\', \'  70\', \'  75\', \'  80\', \' \', \' 5\', \' \'',None,6),(71,'Motility:Sluggish',0,'خيارات','Seminal Fluid Analysis','2021-05-21 18:41:12','%','','\' 5\', \'10\', \'15\', \'20\', \'25\', \'30\', \'35\', \'40\', \'45\', \'50\', \'55\', \'60\', \'65\', \'70\', \'75\', \'80\', \' \'',None,7),(72,'Motility:Dead',0,'خيارات','Seminal Fluid Analysis','2021-05-21 18:41:27','%','','\' 5\', \'10\', \'15\', \'20\', \'25\', \'30\', \'35\', \'40\', \'45\', \'50\', \'55\', \'60\', \'65\', \'70\', \'75\', \'80\', \' \'',None,8),(73,'Morphology:Normal',0,'خيارات','Seminal Fluid Analysis',None,'%','','\' 5\', \'10\', \'15\', \'20\', \'25\', \'30\', \'35\', \'40\', \'45\', \'50\', \'55\', \'60\', \'65\', \'70\', \'75\', \'80\', \' \'',None,9),(74,'Morphology:Abnormal',0,'خيارات','Seminal Fluid Analysis',None,'%','','\' 5\', \'10\', \'15\', \'20\', \'25\', \'30\', \'35\', \'40\', \'45\', \'50\', \'55\', \'60\', \'65\', \'70\', \'75\', \'80\', \' \'',None,10),(75,' Pus cells:',0,'خيارات','Seminal Fluid Analysis','2021-06-08 22:44:48','/ H.P.F','','\'1-2\', \'1-3\', \'2-3\', \'2-4\', \'4-6\', \'3-5\', \'5-7\', \'6-8\', \'4-5\', \'+\', \'++\', \'+++\', \'++++\', \' \', \' \', \' \'',None,11),(77,'HBS Ag',5,'خيارات',' Hematology and Serology','2021-06-08 23:06:00','','','\' \', \'Positive (+ve)\', \'Negative (-ve)\', \'\'',None,4),(78,'HCV Ab',5,'خيارات',' Hematology and Serology','2021-06-08 23:08:03','','','\' \', \'Positive (+ve)\', \'Negative (-ve)\', \'\'',None,5),(79,'HIV',5,'خيارات',' Hematology and Serology','2021-06-08 23:08:58','','','\' \', \'Positive (+ve)\', \'Negative (-ve)\', \'\'',None,6),(80,'Bacteria',0,'خيارات','General Stool Examination','2021-06-10 17:07:57','','','\'Nil\', \'Few\', \'+\', \'++\', \'+++\', \'++++\', \'\'',None,8),(81,'Monillia.',0,'خيارات','General Stool Examination','2021-06-10 22:29:16','','','\'Nil\', \'Few\', \'+\', \'++\', \'+++\', \'++++\', \'\'',None,9),(82,'Fatty drop',0,'خيارات','General Stool Examination','2021-06-10 22:36:03','','','\'Nil\', \'Few\', \'+\', \'++\', \'+++\', \'++++\', \'\'',None,10),(83,'Bacteria. ',0,'خيارات','General Urine Examination','2021-06-10 17:08:57','','','\'Nil\', \'Few\', \'+\', \'++\', \'+++\', \'++++\', \'\'',None,11),(84,'Monillia',0,'خيارات','General Urine Examination','2021-06-10 22:30:08','','','\'Nil\', \'Few\', \'+\', \'++\', \'+++\', \'++++\', \'\'',None,12),(85,'sugar',0,'خيارات','General Urine Examination','2021-06-10 17:10:09','','','\'Nil\', \'Few\', \'+\', \'++\', \'+++\', \'++++\', \'\'',None,4),(98,'HbA1C',15,'حقل كتابة','Biochemistry','2021-06-06 00:06:40','%','( 4.2 - 6.2 )','',None,10),(99,'B - HCG titer',15,'حقل كتابة',' Hematology and Serology','2021-06-06 01:18:37','mIU/ml','( Less than 10 )','',None,21),(100,'lha',0,'حقل كتابة','Hormones and Viruses','2021-06-06 01:25:46','','','',None,20),(101,'H.Pylori in Serum',10,'خيارات',' Hematology and Serology','2021-06-10 09:30:50','','','\'\', \'Positive (+ve)\', \'Negative (-ve)\'',None,19),(102,'H.Pylori in Stool',10,'خيارات',' Hematology and Serology','2021-06-10 09:31:46','','','\'\', \'Positive (+ve)\', \'Negative (-ve)\'',None,20),(103,'Reumatoid Facter',3,'خيارات',' Hematology and Serology','2021-06-10 09:50:51','','','\'\', \'Positive (+ve)\', \'Negative (-ve)\'',None,16),(104,'A.S.O titer',3,'خيارات',' Hematology and Serology','2021-06-10 09:55:23','','','\'\', \'Negative (-ve)\', \'Positive (+ve) 1/200\', \'Positive (+ve) 1/400\', \'Positive (+ve) 1/600\', \'Positive (+ve) 1/800\', \'Positive (+ve) 1/1000\', \'Positive (+ve) 1/1200\', \'Positive (+ve) 1/1400\', \'Positive (+ve) 1/1600\', \'Positive (+ve) 1/1800\', \'\'',None,17),(105,'C.Reactive Protein',3,'خيارات',' Hematology and Serology','2021-06-10 09:58:29','','','\'\', \'Positive (+ve)\', \'Negative (-ve)\', \'\'',None,18),(106,'HAV',5,'خيارات',' Hematology and Serology','2021-06-10 10:01:04','','','\'\', \'Positive (+ve)\', \'Negative (-ve)\', \'\'',None,23),(108,'Mucus',0,'خيارات','General Urine Examination','2021-06-10 17:56:26','','','\'Nil\', \'Few\', \'+\', \'++\', \'+++\', \'++++\', \'\'',None,13),(114,'F.Blood Sugar:',0,'حقل كتابة','Oral Glucose Tolerance Test (O.G.T.T)','2021-06-19 19:32:03','mg/dl','','',None,1),(115,'Blood sugar:',0,'حقل كتابة','Oral Glucose Tolerance Test (O.G.T.T)','2021-06-19 19:33:18','mg/dl','after 1/2 hr.','',None,2),(116,'Blood Sugar:-',0,'حقل كتابة','Oral Glucose Tolerance Test (O.G.T.T)','2021-06-19 19:34:18','mg/dl','after 1  hr.','',None,3),(117,'Blood Sugar,,',0,'حقل كتابة','Oral Glucose Tolerance Test (O.G.T.T)','2021-06-19 19:35:32','mg/dl','after 1 hr. and 30 min.','',None,4),(118,'Blood Sugar;',12,'حقل كتابة','Oral Glucose Tolerance Test (O.G.T.T)','2021-06-19 19:36:38','mg/dl','after  2 hr.','',None,5)]
		# print(data[0][1])
		# for zindex,row in enumerate(data):
		#     self.cur.execute(""" insert into addanalyst(id,name,price,category,sub_category,date,unit,defult,results,quantity,analyst_index) VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)  """,(data[zindex][0],data[zindex][1],data[zindex][2],data[zindex][3],data[zindex][4],data[zindex][5],data[zindex][6],data[zindex][7],data[zindex][8],data[zindex][9],data[zindex][10]))
		# self.db.commit()
		self.Show_statics()
		self.ADD_MANDOBS_TO_COMBO()

	def DB(self):
		self.db = connect(host='localhost', user='root', password='12345', db='tahlel2', charset="utf8",
						  use_unicode=True, port=3306)
		self.cur = self.db.cursor()

	def Show_Word_Doc_Data(self):
		self.cur.execute(''' SELECT * FROM word ''')
		data = self.cur.fetchone()
		self.lineEdit_18.setText(data[10])
		self.lineEdit_20.setText(data[11])
		self.lineEdit_19.setText(data[12])
		self.lineEdit_23.setText(data[2])
		self.lineEdit_22.setText(data[3])
		self.lineEdit_37.setText(data[4])
		self.lineEdit_41.setText(data[5])
		self.lineEdit_35.setText(data[6])
		self.lineEdit_42.setText(data[7])
		self.lineEdit_43.setText(data[1])
		self.lineEdit_34.setText(data[8])
		# self.lineEdit_45.setText(data[13])
		# self.lineEdit_44.setText(data[14])
		self.lineEdit_46.setText(data[15])
		self.lineEdit_47.setText(data[16])

	def Update_Word_Doc_Data(self):
		var1 = self.lineEdit_16.text()
		var2 = self.lineEdit_18.text()
		var3 = self.lineEdit_20.text()
		var4 = self.lineEdit_19.text()
		var5 = self.lineEdit_23.text()
		var6 = self.lineEdit_22.text()
		var7 = self.lineEdit_37.text()
		var8 = self.lineEdit_41.text()
		var9 = self.lineEdit_35.text()
		var10 = self.lineEdit_42.text()
		var11 = self.lineEdit_43.text()
		var12 = self.lineEdit_34.text()
		var13 = ''
		var14 = ''
		var15 = self.lineEdit_46.text()
		var16 = self.lineEdit_47.text()
		self.cur.execute(
			''' UPDATE word SET  shop_name=%s,phone1=%s,phone2=%s,employee_name1=%s,employee_name2=%s,employee1_shahada=%s,employee2_shahada=%s,gps=%s,client_name=%s,client_lqb=%s,doctor_name=%s,doctor_lqb=%s,client_name2=%s,client_lqb2=%s,doctor_name2=%s,doctor_lqb2=%s WHERE id=1''',
			(var11, var5, var6, var7, var8, var9, var10, var12, var1, var2, var3, var4, var13, var14, var15, var16,))
		self.db.commit()
		QMessageBox.information(self, 'info', 'تم تحديث المعلومات بنجاح')
		self.Show_Word_Doc_Data()
		self.Add_Data_To_history(4, 7)
		# self.History()
		self.Update_Word_Info()

	def Update_Word_Info(self):
		global word_files
		global save_word_files
		global word_data
		self.cur.execute(''' SELECT * FROM  word WHERE id=1''')
		word_data = self.cur.fetchone()
		self.cur.execute(''' SELECT * FROM paths WHERE id=1 ''')
		mydata = self.cur.fetchone()
		word_files = mydata[1]
		save_word_files = mydata[2]

	def all_per(self):
		if self.checkBox.isChecked():
			self.checkBox_58.setCheckState(True)
			self.checkBox_57.setCheckState(True)
			self.checkBox_56.setCheckState(True)
			self.checkBox_9.setCheckState(True)
			self.checkBox_55.setCheckState(True)
			self.checkBox_44.setCheckState(True)
			self.checkBox_40.setCheckState(True)
			self.checkBox_13.setCheckState(True)
			self.checkBox_8.setCheckState(True)
			self.checkBox_15.setCheckState(True)
			self.checkBox_9.setCheckState(True)
			self.checkBox_55.setCheckState(True)
			self.checkBox_44.setCheckState(True)
			self.checkBox_40.setCheckState(True)
			self.checkBox_13.setCheckState(True)
			self.checkBox_16.setCheckState(True)
			self.checkBox_14.setCheckState(True)
			self.checkBox_7.setCheckState(True)
			self.checkBox_11.setCheckState(True)
			self.checkBox_12.setCheckState(True)
			self.checkBox_46.setCheckState(True)
			self.checkBox_49.setCheckState(True)
			self.checkBox_50.setCheckState(True)
			self.checkBox_47.setCheckState(True)
			self.checkBox_48.setCheckState(True)
			self.checkBox_51.setCheckState(True)
			self.checkBox_52.setCheckState(True)
			self.checkBox_53.setCheckState(True)
			self.checkBox_54.setCheckState(True)
			self.checkBox_44.setCheckState(True)
			self.checkBox_45.setCheckState(True)
			self.checkBox_42.setCheckState(True)
			self.checkBox_43.setCheckState(True)
			self.checkBox_41.setCheckState(True)
		else:
			self.False_checkState()

	def handel_buttons(self):
		global addTrue
		global edit_employee_check
		if not edit_employee_check:
			self.comboBox_3.hide()
		self.pushButton_15.clicked.connect(self.Light_Blue_Theme)
		self.pushButton_9.clicked.connect(self.Dark_Orange_Theme)
		self.pushButton_13.clicked.connect(self.Dark_Blue_Theme)
		self.pushButton_14.clicked.connect(self.Dark_Theme)
		self.pushButton_11.clicked.connect(self.Dark_Gray_Theme)
		self.pushButton.clicked.connect(self.Open_Sales_Page)
		self.pushButton_21.clicked.connect(self.Open_Login_Page)
		self.pushButton_4.clicked.connect(self.Open_Settings_Page)
		self.pushButton_12.clicked.connect(self.Open_Login_Page)
		self.pushButton_3.clicked.connect(self.Open_History_Page)
		self.pushButton_5.clicked.connect(self.Open_clients_Page)
		self.pushButton_2.clicked.connect(self.Open_Analyse_Page)
		self.pushButton_8.clicked.connect(self.Open_ResetPassword_Page)
		self.pushButton_17.clicked.connect(lambda arg='no': self.Sales_Page(for_loop2=arg))
		self.pushButton_30.clicked.connect(self.get_client_id)
		self.comboBox_16.currentIndexChanged.connect(self.Show_Type_of_result_category)
		self.comboBox_2.currentIndexChanged.connect(self.Show_permissions)
		self.comboBox_23.currentIndexChanged.connect(lambda: self.tslsol_wout_b('not in edit'))
		self.comboBox_26.currentIndexChanged.connect(lambda: self.tslsol_wout_b('in edit'))
		self.pushButton_41.clicked.connect(self.Show_all_taslsol)
		self.pushButton_50.clicked.connect(self.Show_all_taslsol)
		self.comboBox_21.currentTextChanged.connect(self.Show_analyst_in_Edit_Or_Delete)
		self.comboBox_28.currentTextChanged.connect(self.Show_Doctor_Data)
		self.comboBox_3.currentTextChanged.connect(self.Show_employee_data)
		self.pushButton_29.clicked.connect(self.Clients_Page)
		self.pushButton_16.clicked.connect(self.Show_Add_Buys_Complete_Dialog)
		self.pushButton_20.clicked.connect(self.Show_All_The_Analysts)
		self.pushButton_27.clicked.connect(self.Edit_Analyst)
		self.pushButton_7.clicked.connect(self.Log_In_Chieck)
		self.pushButton_28.clicked.connect(self.Delete_Analyst)
		self.pushButton_22.clicked.connect(self.Delete_All_History_Data)
		self.pushButton_18.clicked.connect(self.Print_Sale_Data)
		self.pushButton_36.clicked.connect(self.Search_In_All_Sales)
		self.comboBox_20.currentIndexChanged.connect(self.Search_In_History)
		self.comboBox_24.currentIndexChanged.connect(self.Search_In_History)
		self.pushButton_35.clicked.connect(self.clear_data_in_sales)
		self.pushButton_10.clicked.connect(self.Reset_password)
		self.pushButton_26.clicked.connect(self.Add_Analyst)
		self.pushButton_34.clicked.connect(self.Preview)
		self.pushButton_24.clicked.connect(self.Add_Path)
		self.pushButton_54.clicked.connect(self.Show_All_Clients)
		self.pushButton_43.clicked.connect(self.Show_search_Widget)
		self.pushButton_44.clicked.connect(self.Show_multy_Dialog)
		self.comboBox_16.view().pressed.connect(self.Set_Chick_State2)
		self.pushButton_53.clicked.connect(self.Add_permissions)
		self.pushButton_25.clicked.connect(self.Add_permissions)
		self.pushButton_42.clicked.connect(self.Show_search_Widget2)
		self.pushButton_53.clicked.connect(self.Add_employee)
		self.pushButton_25.clicked.connect(self.Add_employee)
		self.pushButton_23.clicked.connect(self.Show_statics)
		self.pushButton_31.clicked.connect(self.Update_Word_Doc_Data)
		self.pushButton_33.pressed.connect(self.echo_mode)
		self.pushButton_33.released.connect(self.echo_mode)
		self.pushButton_38.pressed.connect(self.echo_mode2)
		self.pushButton_38.released.connect(self.echo_mode2)
		self.checkBox.clicked.connect(self.all_per)
		self.Add_all_employee_to_comboBox()
		self.comboBox_4.currentTextChanged.connect(self.Show_client_data_by_current_index)
		# self.my_def2()
		self.add_all_subCategory_toList()
		self.pushButton_47.clicked.connect(self.Add_Doctor)
		self.pushButton_48.clicked.connect(self.Update_doctor)
		self.pushButton_49.clicked.connect(self.Delete_doctor)
		self.add_doctors_to_list()
		self.tabWidget_4.currentChanged.connect(self.Show_default_statics)
		self.pushButton_4.clicked.connect(self.Show_default_statics)
		self.pushButton_46.clicked.connect(self.Show_Category_Dialog)
		self.pushButton_45.clicked.connect(self.Show_Category_Dialog)
		self.tabWidget_3.currentChanged.connect(self.ShowIsInCategory)
		self.pushButton_39.clicked.connect(self.Show_analyst_chioces_dialog)
		self.pushButton_40.clicked.connect(self.Show_analyst_chioces_dialog)
		self.pushButton_51.clicked.connect(self.show_mandob_dialog)
		self.pushButton_52.clicked.connect(self.Show_all_buys)
		self.tabWidget_2.currentChanged.connect(self.CC)
		self.tabWidget_5.currentChanged.connect(self.CC2)
		self.pushButton_5.clicked.connect(self.CC2)
		self.pushButton_55.clicked.connect(self.Add_Human_Type)
		self.pushButton_56.clicked.connect(self.Edit_Human_Type)
		self.pushButton_57.clicked.connect(self.Delete_human_type)
		self.comboBox_34.currentTextChanged.connect(self.show_Data_in_Human_Edit)
		self.pushButton_62.clicked.connect(self.Add_name_lqb_to_human_type)
		self.comboBox_32.currentTextChanged.connect(self.VB)
		self.pushButton_58.clicked.connect(lambda: self.Show_NormalTextDialog('edit'))
		self.pushButton_59.clicked.connect(self.Show_NormalTextDialog)
		self.pushButton_60.clicked.connect(self.ShowEditClientDialog)
		self.pushButton_61.clicked.connect(self.GoBack)
		self.pushButton_63.clicked.connect(self.GoUp)
		self.spinBox.valueChanged.connect(self.HideOrShowGoButtons)
		self.pushButton_32.clicked.connect(self.ShowHandelPriceDialog)
	def ShowHandelPriceDialog(self):
		if not self.comboBox_4.isEnabled():
			self.get_total_price()
			self.PriceDialog = HandelPriceClass.Dialog()
			self.PriceDialog.lineEdit.setText(self.lineEdit_24.text())
			self.cur.execute(''' select price_added,full_price,pushed_price,discount,minus_price,latest_price from price_task where client_name=%s and date(date)=%s ''',(self.comboBox_4.currentText(),date.today()))
			data = self.cur.fetchone()
			self.PriceDialog.spinBox_10.setValue(int(self.lineEdit_24.text()))
			if data:
				if float(self.lineEdit_24.text()) == float(data[1]):
					self.PriceDialog.spinBox_7.setValue(data[0])
					self.PriceDialog.spinBox_8.setValue(data[3])
					self.PriceDialog.spinBox_9.setValue(data[4])
					self.PriceDialog.spinBox_10.setValue(data[2])		
					self.ShowRealTiemePrice()
				else:
					self.cur.execute(''' delete from price_task where client_name=%s and date(date)=%s ''',(self.comboBox_4.currentText(),date.today(),))
					self.db.commit()
			self.PriceDialog.spinBox_7.setMaximum(2147483647)
			self.PriceDialog.spinBox_7.setMinimum(0)
			# self.PriceDialog.spinBox_9.setMaximum(float(self.PriceDialog.lineEdit.text()))
			# self.PriceDialog.spinBox_9.setMinimum(0)
			self.PriceDialog.spinBox_10.setMaximum(int(float(self.PriceDialog.lineEdit.text())))
			self.PriceDialog.spinBox_10.setMinimum(0)
			
			self.PriceDialog.pushButton_16.clicked.connect(self.FF6)
			self.PriceDialog.spinBox_7.valueChanged.connect(self.ShowRealTiemePrice)
			self.PriceDialog.spinBox_8.valueChanged.connect(self.ShowRealTiemePrice)
			self.PriceDialog.spinBox_9.valueChanged.connect(self.ShowRealTiemePrice)
			self.PriceDialog.spinBox_10.valueChanged.connect(self.ShowRealTiemePrice)
			self.PriceDialog.lineEdit.textChanged.connect(self.GGH)
			self.PriceDialog.show()
		else:
			QMessageBox.information(self,'','لا يوجد اي نموذج ليتم تطبيق السعر عليه')
	def GGH(self):
		self.PriceDialog.spinBox_10.setMaximum(int(float(self.PriceDialog.lineEdit.text())))
		self.PriceDialog.spinBox_10.setMinimum(0)
	def ShowRealTiemePrice(self):
		# self.PriceDialog.spinBox_9.setMaximum(float(self.PriceDialog.lineEdit.text()))
		# self.PriceDialog.spinBox_9.setMinimum(0)
		latest = 0
		now_price = float(self.lineEdit_24.text())
		price_added = self.PriceDialog.spinBox_7.value()
		khsm = self.PriceDialog.spinBox_8.value()
		minus_price = self.PriceDialog.spinBox_9.value()
		latest += now_price
		latest += price_added
		latest -= minus_price
		latest = latest - (latest *(khsm/100))
		self.PriceDialog.lineEdit.setText(str(latest))
		self.PriceDialog.lineEdit_2.setText(str(float(self.PriceDialog.lineEdit.text()) - self.PriceDialog.spinBox_10.value()))
	def FF6(self):
		self.cur.execute(''' select price_added,full_price,pushed_price,discount,minus_price,latest_price from price_task where client_name=%s and date(date)=%s ''',(self.comboBox_4.currentText(),date.today()))
		data = self.cur.fetchone()
		latest = 0
		now_price = int(self.lineEdit_24.text())
		price_added = self.PriceDialog.spinBox_7.value()
		khsm = self.PriceDialog.spinBox_8.value()
		minus_price = self.PriceDialog.spinBox_9.value()
		latest += now_price
		latest += price_added
		latest -= minus_price
		latest = latest - (latest *(khsm/100))
		if data:
			self.cur.execute(''' update price_task set client_name=%s,full_price=%s,pushed_price=%s,discount=%s,minus_price=%s,latest_price=%s,date=%s,price_added=%s where client_name=%s and date(date)=%s''',(self.comboBox_4.currentText(),int(self.lineEdit_24.text()),self.PriceDialog.spinBox_10.value(),khsm,minus_price,latest,datetime.now(),self.PriceDialog.spinBox_7.value(),self.comboBox_4.currentText(),date.today(),))
			self.db.commit()
		else:
			self.cur.execute(''' insert into price_task (client_name,full_price,pushed_price,discount,minus_price,latest_price,date,price_added) VALUES(%s,%s,%s,%s,%s,%s,%s,%s)''',(self.comboBox_4.currentText(),int(self.lineEdit_24.text()),self.PriceDialog.spinBox_10.value(),khsm,minus_price,latest,datetime.now(),self.PriceDialog.spinBox_7.value(),))
			self.db.commit()
		self.lineEdit_24.setText(str(latest - float(self.PriceDialog.lineEdit_2.text())))
		self.PriceDialog.close()
		# print(latest,'WOW')
	def VBG(self):
		global MAX_ID
		self.cur.execute(''' select MAX(id) from addclient ''')
		data = self.cur.fetchone()
		if data:
			MAX_ID = int(data[0])

	def HideOrShowGoButtons(self):
		self.VBG()
		global MAX_ID
		if self.spinBox.value() >= MAX_ID:
			self.pushButton_61.show()
			self.pushButton_63.hide()
		elif self.spinBox.value() < MAX_ID and self.spinBox.value() > 1:
			self.pushButton_61.show()
			self.pushButton_63.show()
		if self.spinBox.value() == 1 or self.spinBox.value() == 0:
			self.pushButton_61.hide()
			self.pushButton_63.show()
	def GoBack(self):
		self.VBG()
		global select_by_date
		# try:
		self.Update_addNewItem_Data()
		self.spinBox.setValue(self.spinBox.value() - 1)
		select_by_date = 'gg'
		self.Show_All_one_client_analyst()
		self.HideOrShowGoButtons()
		# except:
		# 	QMessageBox.information(self,'','هنالك خطأ يرجى مراجعة العملية')

	def GoUp(self):
		global select_by_date
		# try:
		select_by_date = 'gg'
		self.Update_addNewItem_Data()
		self.spinBox.setValue(self.spinBox.value() + 1)
		self.Show_All_one_client_analyst()
		self.HideOrShowGoButtons()
		# except:
		# 	QMessageBox.information(self,'','هنالك خطأ يرجى مراجعة العملية')
	def ShowEditClientDialog(self):
		self.EditClientDialog = EditClientClass.Dialog()
		global all_doctors
		if self.spinBox.value() :
			self.EditClientDialog.lineEdit_17.setText(self.comboBox_4.currentText())
			self.EditClientDialog.lineEdit_16.hide()
			self.EditClientDialog.spinBox_7.setValue(self.spinBox_7.value())
			self.EditClientDialog.comboBox_14.addItems(self.RETURN_ALL_human_type())
			self.EditClientDialog.comboBox_14.setCurrentIndex(self.comboBox_14.currentIndex())
			self.EditClientDialog.comboBox_15.addItems(all_doctors)
			self.EditClientDialog.comboBox_15.setCurrentIndex(self.comboBox_15.currentIndex())
			self.EditClientDialog.textEdit.setPlainText(self.textEdit.toPlainText())
			self.EditClientDialog.pushButton_17.clicked.connect(self.HandelEditClientDialog)
			self.EditClientDialog.pushButton_16.clicked.connect(self.HandelEditClientDialog)
			self.EditClientDialog.show()
		else:
			QMessageBox.information(self,'','يرجى اختيار اسم مراجع صحيح')
	def HandelEditClientDialog(self):
		if self.sender().text() == 'حفظ':
			try:
				self.cur.execute(''' UPDATE addclient SET client_name=%s,client_age=%s,client_genus=%s,client_doctor=%s where client_name=%s ''',(self.EditClientDialog.lineEdit_17.text(),self.EditClientDialog.spinBox_7.value(),self.EditClientDialog.comboBox_14.currentText(),self.EditClientDialog.comboBox_15.currentText(),self.comboBox_4.currentText(),))
				self.cur.execute(''' UPDATE addnewitem SET client_name=%s,client_age=%s,genus=%s,doctor_name=%s,notes=%s where client_name=%s ''',(self.EditClientDialog.lineEdit_17.text(),self.EditClientDialog.spinBox_7.value(),self.EditClientDialog.comboBox_14.currentText(),self.EditClientDialog.comboBox_15.currentText(),self.EditClientDialog.textEdit.toPlainText(),self.comboBox_4.currentText(),))
				self.cur.execute(''' UPDATE price_task set client_name=%s where client_name=%s ''',(self.EditClientDialog.lineEdit_17.text(),self.comboBox_4.currentText(),))
				self.db.commit()
				QMessageBox.information(self.EditClientDialog,'','تم تعديل معلومات المراجع بنجاح')
				self.EditClientDialog.close()
				self.add_clients_to_combo()
				self.comboBox_4.setCurrentText('s%^# ')
				self.comboBox_4.setCurrentText(self.EditClientDialog.lineEdit_17.text())
				# self.comboBox_14.setCurrentIndex(self.EditClientDialog.comboBox_14.currentIndex())
				# self.spinBox_7.setValue(self.EditClientDialog.spinBox_7.value())
				# self.comboBox_15.setCurrentIndex(self.EditClientDialog.comboBox_15.currentIndex())
				# self.textEdit.setPlainText(self.EditClientDialog.textEdit.toPlainText())
			except:
				QMessageBox.information(self.EditClientDialog,'','هنالك خطأ يرجى مراجعة العملية')
				self.EditClientDialog.close()
		else:
			warning = QMessageBox.warning(self, '', 'هل انت متأكد من حذف هذا المراجع ؟',QMessageBox.Yes | QMessageBox.No)
			if warning == QMessageBox.Yes:
				self.cur.execute(''' delete from addclient where client_name=%s''',(self.comboBox_4.currentText(),))
				self.cur.execute(''' delete from addnewitem where client_name=%s''',(self.comboBox_4.currentText(),))
				self.db.commit()
				QMessageBox.information(self.EditClientDialog,'','تم حذف هذا المراجع بنجاح')

	def Show_NormalTextDialog(self,form_edit=None):
		self.NormalTextDialog = NormalTextDialogClass.Dialog()
		self.NormalTextDialog.lineEdit_2.hide()
		if form_edit:
			self.NormalTextDialog.lineEdit_2.setText(self.comboBox_21.currentText())
		else:
			self.NormalTextDialog.lineEdit_2.setText(self.lineEdit_28.text())
		self.NormalTextDialog.comboBox.clear()
		self.NormalTextDialog.comboBox.addItems(self.RETURN_ALL_human_type())
		self.NormalTextDialog.pushButton.clicked.connect(self.Handel_NormalTextDialog)
		self.NormalTextDialog.comboBox.currentIndexChanged.connect(self.show_NormalTextDialog_data)
		self.NormalTextDialog.show()
	def Handel_NormalTextDialog(self):
		self.cur.execute(''' DELETE from analyst_normal_text where analyst_name=%s and genus=%s ''',(self.NormalTextDialog.lineEdit_2.text(),self.NormalTextDialog.comboBox.currentText(),))
		self.cur.execute(''' INSERT INTO analyst_normal_text (analyst_name,genus,normal_text) VALUES(%s,%s,%s)''',(self.NormalTextDialog.lineEdit_2.text(),self.NormalTextDialog.comboBox.currentText(),self.NormalTextDialog.lineEdit.text(),))
		self.db.commit()
		QMessageBox.information(self.NormalTextDialog,'','تم تطبيق البيانات بنجاح')
		self.NormalTextDialog.lineEdit.setText('')
	def show_NormalTextDialog_data(self):
		self.cur.execute(''' select normal_text from analyst_normal_text where analyst_name=%s and genus=%s ''',(self.NormalTextDialog.lineEdit_2.text(),self.NormalTextDialog.comboBox.currentText(),))
		data = self.cur.fetchone()
		if data:
			self.NormalTextDialog.lineEdit.setText(data[0])
		else:
			self.NormalTextDialog.lineEdit.setText('')
	def Show_all_human_type_in_combos(self):
		self.cur.execute('select name from human_type order by -date')
		data = self.cur.fetchall()
		self.comboBox_34.clear()
		complete = False
		if self.textEdit.isEnabled():
			self.comboBox_14.clear()
			complete = True
		my_list = []
		for row in data:
			if row[0] not in my_list:
				my_list.append(row[0])
		self.comboBox_34.addItem('------------')
		self.comboBox_34.addItems(my_list)
		self.comboBox_32.addItems(my_list)
		self.comboBox_34.setCurrentIndex(0)
		if complete:
			self.comboBox_14.addItems(my_list)
			self.comboBox_14.setCurrentIndex(0)
	def Add_Human_Type(self):
		if self.lineEdit_8.text() not in self.RETURN_ALL_human_type():
			self.cur.execute(''' INSERT INTO human_type (name,date) VALUES(%s,%s) ''',(self.lineEdit_8.text(),datetime.now(),))
			self.db.commit()
			QMessageBox.information(self,'','تم اضافة الجنس بنجاح')
			self.Show_all_human_type_in_combos()
		else:
			QMessageBox.information(self,'','هذا الجنس موجود بالفعل')
	def Edit_Human_Type(self):
		self.cur.execute(''' Update human_type SET name=%s,date=%s where name=%s ''',(self.lineEdit_9.text(),datetime.now(),self.comboBox_34.currentText(),))
		self.db.commit()
		QMessageBox.information(self,'','تم تعديل الجنس بنجاح')
		self.Show_all_human_type_in_combos()
	def Delete_human_type(self):
		self.cur.execute(''' Delete from human_type where name=%s ''',(self.comboBox_34.currentText(),))
		self.db.commit()
		QMessageBox.information(self,'','تم حذف الجنس بنجاح')
		self.Show_all_human_type_in_combos()
	def show_Data_in_Human_Edit(self):
		self.cur.execute('select name from human_type where name=%s',(self.comboBox_34.currentText(),))
		data = self.cur.fetchone()
		if data:
			self.lineEdit_9.setText(data[0])
	def Add_name_lqb_to_human_type(self):
		self.cur.execute(''' update human_type set report_name=%s,report_lqb=%s where name=%s ''',(self.lineEdit_16.text(),self.lineEdit_18.text(),self.comboBox_32.currentText(),))
		self.db.commit()
		QMessageBox.information(self,'','تم حفظ البيانات بنجاح')
	def VB(self):
		self.cur.execute('select report_name,report_lqb from human_type where name=%s',(self.comboBox_32.currentText(),))
		data = self.cur.fetchone()
		if data:
			self.lineEdit_16.setText(data[0])
			self.lineEdit_18.setText(data[1])
	def CC(self):
		if self.tabWidget_2.currentIndex()==1:
			self.Show_All_The_Sales()
			self.add_today_client_to_list()
			self.Auto_complete_combo7()
	def add_doctors_to_list(self):
		global all_doctors
		self.cur.execute(''' SELECT name FROM doctor ''')
		data = self.cur.fetchall()
		all_doctors.clear()
		for i in data:
			if i[0] not in all_doctors:
				all_doctors.append(i[0])

	def Add_Doctor(self):
		try:
			genus = self.comboBox_18.currentIndex()
			rgenus = ''
			if genus:
				rgenus = 'male'
			else:
				rgenus = 'female'
			self.cur.execute(''' INSERT INTO doctor (name,genus) VALUES(%s,%s)''', (self.lineEdit_5.text(), rgenus,))
			self.db.commit()
			QMessageBox.information(self, 'info', 'تم اضافة الطبيب بنجاح')
			self.Add_Doctor_Data()
			self.add_doctors_to_list()
			self.Add_Data_To_history(3, 6)
			# self.History()
		except Exception as e:
			print(e, '95kerorr')
			QMessageBox.information(self, '', 'هذا الطبيب موجود بالفعل')

	def Update_doctor(self):
		global all_doctors
		if self.comboBox_28.currentText() in all_doctors:
			try:
				genus = self.comboBox_27.currentIndex()
				rgenus = ''
				if genus:
					rgenus = 'male'
				else:
					rgenus = 'female'
				self.cur.execute(''' UPDATE doctor SET name=%s,genus=%s WHERE name=%s''',
								 ((self.lineEdit_6.text(), rgenus, self.comboBox_28.currentText(),)))
				self.db.commit()
				QMessageBox.information(self, 'info', 'تم تحديث بيانات الطبيب بنجاح')
				self.Add_Doctor_Data()
				self.add_doctors_to_list()
				self.Add_Data_To_history(4, 6)
				# self.History()
			except:
				QMessageBox.information(self, '', 'لا يمكن تكرار اسم الطبيب , يرجى تغيير اسم الطبيب')
		else:
			QMessageBox.information(self, '', 'هذا الطبيب غير موجود')

	def Delete_doctor(self):
		if self.comboBox_28.currentText() in all_doctors:
			genus = self.comboBox_27.currentIndex()
			rgenus = ''
			if genus:
				rgenus = 'male'
			else:
				rgenus = 'female'
			warning = QMessageBox.warning(self, '', f'هل انت متأكد من مسح الطبيب {self.comboBox_28.currentText()}؟',
										  QMessageBox.Yes | QMessageBox.No)
			if warning == QMessageBox.Yes:
				self.cur.execute(''' DELETE FROM doctor  WHERE name=%s''', ((self.comboBox_28.currentText(),)))
				self.db.commit()
				QMessageBox.information(self, 'info', 'تم حذف الطبيب بنجاح')
				self.Add_Doctor_Data()
				self.add_doctors_to_list()
				self.Add_Data_To_history(5, 6)
				# self.History()
		else:
			QMessageBox.information(self, '', 'هذا الطبيب غير موجود')

	def Show_Doctor_Data(self):
		global Delete_Doctor
		global Edit_Doctor
		self.cur.execute(''' select name,genus from doctor where name=%s ''', (self.comboBox_28.currentText(),))
		data = self.cur.fetchone()
		if data:
			genus = data[1]
			if genus == 'male':
				self.comboBox_27.setCurrentIndex(1)
			else:
				self.comboBox_27.setCurrentIndex(0)
			self.lineEdit_6.setText(data[0])
			if Edit_Doctor:
				self.pushButton_48.setEnabled(True)
			else:
				self.pushButton_48.setEnabled(False)
			if Delete_Doctor:
				self.pushButton_49.setEnabled(True)
			else:
				self.pushButton_49.setEnabled(False)
		else:
			self.pushButton_48.setEnabled(False)
			self.pushButton_49.setEnabled(False)
			self.comboBox_27.setCurrentIndex(0)
			self.lineEdit_6.setText('')
	def RETURN_ALL_human_type(self):
		self.cur.execute('select name from human_type order by -date')
		data = self.cur.fetchall()
		cc = []
		for row in data:
			if row[0] not in cc:
				cc.append(row[0])
		return cc
	def Show_Normal_Dialog(self,from_save=None):
		global global_from_save
		print('dmdlmld',from_save)
		self.NormalDialog = NormalDialogClass.Dialog()
		the_combo = None
		if from_save:
			print('YEEEEEEEEEEEEEEEEEEEEEEEEEEEEEES')
			the_combo = self.comboBox_25
		else:
			the_combo = self.comboBox_22
		if the_combo.currentIndex()==4 or the_combo.currentIndex()==5:
			self.NormalDialog.groupBox.hide()
			self.NormalDialog.groupBox_2.show()
			self.NormalDialog.comboBox.clear()
			self.NormalDialog.comboBox.addItems(self.RETURN_ALL_human_type())
			# if from_save:
			# 	normal = self.lineEdit_39.text()
			# else:
			# 	normal = self.lineEdit_40.text()
			# normal_list=re.findall("[-+]?[.]?[\d]+(?:,\d\d\d)*[\.]?\d*(?:[eE][-+]?\d+)?", normal)
			# try:
			# 	print(normal_list)
			# 	self.NormalDialog.doubleSpinBox_3.setValue(float(normal_list[0]))
			# 	self.NormalDialog.doubleSpinBox.setValue(float(normal_list[1]))
			# except Exception as e:
			# 	print(e)
		else:
			self.NormalDialog.groupBox_2.hide()
			self.NormalDialog.groupBox.show()
			self.NormalDialog.comboBox_2.clear()
			self.NormalDialog.comboBox_2.addItems(self.RETURN_ALL_human_type())
			comboX =None
			if from_save:
				print('#################### YES ############################3')
				comboX = self.comboBox_31
			else:
				comboX = self.comboBox_11
			for row in range(0,comboX.count()):
				self.NormalDialog.tableWidget.insertRow(row)
				comboX.setCurrentIndex(row)
				self.NormalDialog.tableWidget.setItem(row,0,QTableWidgetItem(str(comboX.currentText())))
				combo = QComboBox()
				combo.addItems(['طبيعي','غير طبيعي'])
				if from_save:
					analyst_name = self.lineEdit_29.text()
					self.cur.execute(''' select normal_value2 from analystnormal where analyst_name=%s and genus_type=%s and normal_type=%s ''',(analyst_name,self.NormalDialog.comboBox_2.currentText(),'list',))
					dataCX = self.cur.fetchone()
					if dataCX:
						x = '['+str(dataCX[0])+']'
						list_data = literal_eval(str(x))
						indexC = combo.findText(str(list_data[row]),Qt.MatchFixedString)
						combo.setCurrentIndex(indexC)
				else:
					combo.setCurrentIndex(0)
				self.NormalDialog.tableWidget.setCellWidget(row,1,combo)
			comboX.setCurrentIndex(0)
		self.NormalDialog.pushButton.clicked.connect(lambda : self.Handel_Normal_Dialog(from_save))
		self.NormalDialog.comboBox.currentIndexChanged.connect(self.NormalDialogGenusComboChanged)
		self.NormalDialog.comboBox_2.currentIndexChanged.connect(self.NormalDialogGenusComboChanged)
		# self.lineEdit_28.setText('')  # analyst_name =
		# self.comboBox_22.setCurrentIndex(0)  # analyst_result_category =
		# self.spinBox_5.setValue(0)  # analyst_price =
		# self.comboBox_23.setCurrentIndex(0)  # sub_category =
		# self.lineEdit_40.setText('')
		# self.lineEdit_48.setText('')
		# QMessageBox.information(self, 'info', 'تم اضافة التحليل بنجاح')
		self.NormalDialog.show()
		# self.NormalDialog.CloseEvent()
		if not from_save:
			self.lineEdit_28.setEnabled(False)
			self.comboBox_22.setEnabled(False)
			self.spinBox_5.setEnabled(False)
			self.comboBox_23.setEnabled(False)
			self.lineEdit_40.setEnabled(False)
			self.lineEdit_48.setEnabled(False)
		else:
			self.comboBox_21.setEnabled(False)
			self.lineEdit_29.setEnabled(False)
			self.comboBox_25.setEnabled(False)
			self.lineEdit_39.setEnabled(False)
			self.lineEdit_38.setEnabled(False)
			self.spinBox_6.setEnabled(False)
			self.comboBox_26.setEnabled(False)
		self.NormalDialog.closeEvent = self.RealCloseEvent.__get__(self.NormalDialog,NormalDialogClass.Dialog())
		global_from_save = from_save
	def RealCloseEvent(self,event,from_save=None):
		global global_from_save
		if global_from_save:
			self.comboBox_21.setEnabled(True)
			self.lineEdit_29.setEnabled(True)
			self.comboBox_25.setEnabled(True)
			self.lineEdit_39.setEnabled(True)
			self.lineEdit_38.setEnabled(True)
			self.spinBox_6.setEnabled(True)
			self.comboBox_26.setEnabled(True)
			self.comboBox_21.setCurrentIndex(0)  # analyst_current_name =
			self.lineEdit_29.setText('')  # analyst_name =
			self.comboBox_25.setCurrentIndex(0)  # analyst_result_category =
			self.lineEdit_39.setText('')  # defult =
			self.lineEdit_38.setText('')  # unit =
			self.spinBox_6.setValue(0)  # analyst_price =
			self.comboBox_26.setCurrentIndex(0)  # sub_category =
			QMessageBox.information(self, 'info', 'تم تعديل التحليل بنجاح')
			self.Show_all_analysts_in_combo()
		else:
			self.lineEdit_28.setEnabled(True)
			self.comboBox_22.setEnabled(True)
			self.spinBox_5.setEnabled(True)
			self.comboBox_23.setEnabled(True)
			self.lineEdit_40.setEnabled(True)
			self.lineEdit_48.setEnabled(True)
			self.lineEdit_28.setText('')  # analyst_name =
			self.comboBox_22.setCurrentIndex(0)  # analyst_result_category =
			self.spinBox_5.setValue(0)  # analyst_price =
			self.comboBox_23.setCurrentIndex(0)  # sub_category =
			self.lineEdit_40.setText('')
			self.lineEdit_48.setText('')
			QMessageBox.information(self, 'info', 'تم اضافة التحليل بنجاح')
			self.Show_all_analysts_in_combo()
		print('Yes Window closed')
	def Handel_Normal_Dialog(self,from_save=None):
		if from_save:
			analyst_name = self.lineEdit_29.text()
		else:
			analyst_name = self.lineEdit_28.text()
		if self.NormalDialog.groupBox.isHidden():
			genus_type = self.NormalDialog.comboBox.currentText()
			normal_value1 = self.NormalDialog.doubleSpinBox_3.value()
			normal_value2 = self.NormalDialog.doubleSpinBox.value()
			normal_type = "number"
			self.cur.execute(""" DELETE FROM analystnormal where analyst_name=%s and genus_type=%s """,(analyst_name,genus_type,))
			self.cur.execute(''' INSERT INTO analystnormal (analyst_name,genus_type,normal_type,normal_value1,normal_value2) VALUES(%s,%s,%s,%s,%s) ''',(analyst_name,genus_type,normal_type,str(normal_value1),str(normal_value2),))
		else:
			genus_type = self.NormalDialog.comboBox_2.currentText()
			normal_type = "list"
			results = []
			is_Normal = []
			for row in range(0,self.NormalDialog.tableWidget.rowCount() - 1):
				results.append(str(self.NormalDialog.tableWidget.item(row,0).text()))
				is_Normal.append(str(self.NormalDialog.tableWidget.cellWidget(row,1).currentText()))
			self.cur.execute(""" DELETE FROM analystnormal where analyst_name=%s and genus_type=%s """,(analyst_name,genus_type,))
			self.cur.execute(''' INSERT INTO analystnormal (analyst_name,genus_type,normal_type,normal_value1,normal_value2) VALUES(%s,%s,%s,%s,%s) ''',(analyst_name,genus_type,normal_type,str(results)[1:-1],str(is_Normal)[1:-1]))
		self.db.commit()
		QMessageBox.information(self.NormalDialog, 'info', 'تم تطبيق البيانات بنجاح')
	def NormalDialogGenusComboChanged(self):
		global global_from_save
		if global_from_save:
			print('$$$$$$$$$$$$$$$$$$$$$$')
			analyst_name = self.lineEdit_29.text()
		else:
			analyst_name = self.lineEdit_28.text()
		if self.NormalDialog.groupBox.isHidden():
			genus_type = self.NormalDialog.comboBox.currentText()
			self.cur.execute(''' select normal_value1,normal_value2 from analystnormal where analyst_name=%s and genus_type=%s ''',(analyst_name,genus_type,))
			data=self.cur.fetchone()
			if data:
				self.NormalDialog.doubleSpinBox_3.setValue(float(data[0]))
				self.NormalDialog.doubleSpinBox.setValue(float(data[1]))
			else:
				self.NormalDialog.doubleSpinBox_3.setValue(float(0))
				self.NormalDialog.doubleSpinBox.setValue(float(0))
		else:
			genus_type = self.NormalDialog.comboBox_2.currentText()
			self.cur.execute(''' select normal_value1,normal_value2 from analystnormal where analyst_name=%s and genus_type=%s ''',(analyst_name,genus_type,))
			data=self.cur.fetchone()
			print(analyst_name,data)
			for row in range(0,self.NormalDialog.tableWidget.rowCount() - 1):
				if data:
					try:
						x = '['+str(data[1])+']'
						list_data = literal_eval(str(x))
						index = self.NormalDialog.tableWidget.cellWidget(row,1).findText(list_data[row],Qt.MatchFixedString)
						self.NormalDialog.tableWidget.cellWidget(row,1).setCurrentIndex(index)
					except Exception as e:
						print(e)
				else:
					self.NormalDialog.tableWidget.cellWidget(row,1).setCurrentIndex(0)
	def mr(self):
		print('finaly')
	def Show_client_data_by_current_index(self):
		self.cur.execute(
			'''SELECT client_name,doctor_name,client_age,genus,notes,client_id FROM addnewitem WHERE client_name = %s''',
			(self.comboBox_4.currentText(),))
		analyst_data = self.cur.fetchone()
		if analyst_data:
			self.spinBox.setValue(int(analyst_data[5]))
			self.comboBox_15.setCurrentText(str(analyst_data[1]))
			self.comboBox_15.setEnabled(False)
			self.comboBox_14.setEnabled(False)
			self.comboBox_4.setCurrentText(analyst_data[0])
			self.spinBox_7.setValue(int(analyst_data[2]))
			self.spinBox_7.setEnabled(False)
			self.comboBox_14.setCurrentText(analyst_data[3])
			self.comboBox_14.setEnabled(False)
			self.textEdit.setPlainText(str(analyst_data[4]))
			self.textEdit.setEnabled(False)
		else:
			self.spinBox.setValue(0)
			self.comboBox_15.setCurrentText('')
			self.comboBox_15.setEnabled(True)
			self.comboBox_14.setEnabled(True)
			self.spinBox_7.setValue(20)
			self.spinBox_7.setEnabled(True)
			self.comboBox_14.setCurrentIndex(0)
			self.comboBox_14.setEnabled(True)
			self.textEdit.setPlainText('')
			self.textEdit.setEnabled(True)
		name = self.comboBox_4.currentText()
		if name != '' or name != None:
			self.pushButton_44.setEnabled(True)
			self.pushButton_17.setEnabled(True)
			self.comboBox_16.setEnabled(True)
		if name == '' or name == None:
			self.pushButton_44.setEnabled(False)
			self.pushButton_17.setEnabled(False)
			self.comboBox_16.setEnabled(False)

	def echo_mode(self):
		global echo_mode_num
		if echo_mode_num % 2 == 0:
			self.lineEdit_17.setEchoMode(QLineEdit.Normal)
		else:
			self.lineEdit_17.setEchoMode(QLineEdit.Password)
		echo_mode_num += 1

	def echo_mode2(self):
		global echo_mode_num2
		if echo_mode_num2 % 2 == 0:
			self.lineEdit_2.setEchoMode(QLineEdit.Normal)
		else:
			self.lineEdit_2.setEchoMode(QLineEdit.Password)
		echo_mode_num2 += 1

	def add_all_subCategory_toList(self):
		mylist = []
		mylist.clear()
		self.cur.execute(''' SELECT name FROM category ''')
		data = self.cur.fetchall()
		for i in data:
			if i[0] not in mylist:
				mylist.append(i[0])
		return mylist

	def Set_Chick_State2(self, item):
		nitem = self.comboBox_16.model().itemFromIndex(item)
		ritem = self.comboBox_16.model().itemFromIndex(item)
		self.comboBox_16.setCurrentText(ritem.text())
		self.Chick_analyst_category(item)

	def Set_Chick_State(self, item=None):
		try:
			if item.checkState() == Qt.Checked:
				item.setCheckState(Qt.Unchecked)
			else:
				item.setCheckState(Qt.Checked)
		except Exception as e:
			print(e, '95poerorr')

	def setCheckSateForAllItems(self, table=None, ListWidget=None):
		if table.text() == 'الكل':
			if table.checkState() == Qt.Checked:
				for row in range(0, ListWidget.count()):
					row_item = ListWidget.item(row)
					row_item.setCheckState(Qt.Checked)
			else:
				for row in range(0, ListWidget.count()):
					row_item = ListWidget.item(row)
					row_item.setCheckState(Qt.Unchecked)
		for row in range(0, ListWidget.count()):
				print('CC',str(ListWidget.item(row).text()),str(ListWidget.item(row).checkState()))

	def Add_all_analysts_items(self):
		Add_all_analysts_items_list = []
		for row in range(0, self.tableWidget_5.rowCount() - 1):
			Add_all_analysts_items_list.append(str(self.tableWidget_5.item(row, 0).text()))
		return Add_all_analysts_items_list

	def Show_multy_Dialog(self):
		self.comboBox_17.clear()
		self.doubleSpinBox_7.clear()
		self.spinBox_4.clear()
		self.lineEdit_21.setText('')
		self.Update_addNewItem_Data()
		self.Show_All_one_client_analyst(from_add_multy='Not None')
		# self.Add_all_analysts_items()
		self.Dialog = multyclass.MultyDialog()
		self.cur.execute(''' SELECT name,sub_category FROM addanalyst order by sub_category,analyst_index''')
		data = self.cur.fetchall()
		my_list = []
		for index1, i in enumerate(data):
			if data[index1][1] not in my_list:
				my_list.append(data[index1][1])
		all_listWidget = []
		for index, ii in enumerate(my_list):
			tab = QWidget()
			self.Dialog.pushButton.setText('اضافة')
			self.Dialog.tabWidget.addTab(tab, str(ii))
			self.Dialog.my_listWidget = QListWidget(tab)
			self.Dialog.my_listWidget.setObjectName("listWidget_+" + str(index + 1))
			all_listWidget.append(str("listWidget_+" + str(index + 1)))
			self.Dialog.my_listWidget.setGeometry(QRect(0, 10, 641, 491))
			font = QFont()
			font.setPointSize(11)
			self.Dialog.my_listWidget.setFont(font)
			item = QListWidgetItem()
			item.setFlags(item.flags() | Qt.ItemIsUserCheckable)
			item.setCheckState(Qt.Checked)
		for index2 in range(0, self.Dialog.tabWidget.count()):
			if index2 < self.Dialog.tabWidget.count():
				listWidget_object = self.Dialog.findChild(QListWidget, f"listWidget_+{index2 + 1}")
				listWidget_object.itemClicked.connect(self.Set_Chick_State)
				item2 = QListWidgetItem()
				item2.setFlags(item2.flags() | Qt.ItemIsUserCheckable)
				item2.setCheckState(Qt.Unchecked)
				item2.setText('الكل')
				listWidget_object.itemClicked.connect((lambda arg=listWidget_object,
															  arg2=listWidget_object: self.setCheckSateForAllItems(
					table=arg, ListWidget=arg2)))
				listWidget_object.addItem(item2)
				for iy in range(0, len(data)):
					if self.Dialog.tabWidget.tabText(index2) == str(data[iy][1]):
						item = QListWidgetItem()
						item.setFlags(item.flags() | Qt.ItemIsUserCheckable)
						item.setCheckState(Qt.Unchecked)
						item.setText(str(data[iy][0]))
						listWidget_object.addItem(item)
						if item.text() in self.Add_all_analysts_items():
							item.setFlags(Qt.NoItemFlags)
						else:
							item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable | Qt.ItemIsUserCheckable)
			else:
				break
		self.Dialog.pushButton.clicked.connect(self.Handel_multy_Dialog)
		self.Dialog.exec_()

	def Handel_multy_Dialog(self):
		start_time = time.time()
		global clients_name_glo
		global clients_name_glo2
		all_analysts = []
		first = True
		end2=False
		bb =False
		for i in range(0, self.Dialog.tabWidget.count()):
			if bb:
				break
			listWidget_object = self.Dialog.findChild(QListWidget, f"listWidget_+{i + 1}")
			for row in range(0, listWidget_object.count()):
				if listWidget_object.item(row).checkState() == Qt.Checked:
					self.comboBox_16.setCurrentText(str(listWidget_object.item(row).text()))
					if not end2:
						self.cur.execute(''' select quantity from addanalyst where name =%s ''',(str(listWidget_object.item(row).text()),))
						da=self.cur.fetchone()
						if da:
							if da[0]==0:
								warning = QMessageBox.warning(self, '','هنالك تحاليل نُفذت كميتها هل تريد المواصلة؟',
											QMessageBox.Yes | QMessageBox.No)
								if warning == QMessageBox.Yes:
									first=False
									end2=True
								else:
									bb=True
									break
					self.Sales_Page()
					first =False
					all_analysts.append(str(listWidget_object.item(row).text()))
		self.Dialog.close()
		# self.add_client_to_list()
		if not first:
			if self.comboBox_4.currentText() not in clients_name_glo:
				clients_name_glo.append(self.comboBox_4.currentText())
			self.Show_All_one_client_analyst('not')
			# # # self.Add_buttons_combo_spin_to_tableWidget()
			self.get_total_price()
			# # self.my_def2()
			if self.comboBox_4.currentText() not in clients_name_glo2:
				clients_name_glo2.append(self.comboBox_4.currentText())
			# self.add_today_client_to_list()
			# self.Auto_complete_combo7()
			# # self.Show_All_The_Sales()
			self.Add_Data_To_history(3, 1)
			# # self.History()
			# self.Add_all_analysts_items()
			for rowCX in all_analysts:
					if rowCX in self.all_analyst_in_buys():
						# self.cur.execute(''' update addbuys set item_quantity=item_quantity -1 where to_analysts=%s and item_quantity!=0''',(rowCX,))
						self.cur.execute(''' update addanalyst set quantity = quantity - 1 where name = %s and quantity!=0''',(rowCX,))
			self.db.commit()
		print("--- %s seconds ---" % (time.time() - start_time))
	def all_analyst_in_buys(self):
		self.cur.execute(''' select to_analysts from addbuys ''')
		data=self.cur.fetchall()
		cc=[]
		for row in data:
			cc.append(row[0])
		return cc
	def Show_search_Widget(self):
		self.searchWidget = searchDialog.Dialog()
		self.searchWidget.spinBox.setValue(self.spinBox.value())
		self.searchWidget.show()
		self.searchWidget.dateEdit_6.setDate(date.today())
		self.searchWidget.dateEdit_5.setDate(date.today())
		self.searchWidget.pushButton_20.clicked.connect(self.Search_by_date)

	def Show_search_Widget2(self):
		global search_info_by_date
		search_info_by_date = True
		self.searchWidget2 = searchDialog2.Dialog()
		self.searchWidget2.spinBox.setValue(self.spinBox_2.value())
		self.searchWidget2.show()
		self.searchWidget2.dateEdit_6.setDate(date.today())
		self.searchWidget2.dateEdit_5.setDate(date.today())
		self.searchWidget2.pushButton_20.clicked.connect(self.Clients_Page)

	def Search_by_date(self):
		global select_by_date
		select_by_date = True
		self.Show_All_one_client_analyst()
		self.searchWidget.close()

	def add_clients_to_combo(self):
		clients = []
		self.cur.execute(''' SELECT client_name FROM addclient''')
		data = self.cur.fetchall()
		for row in data:
			if row[0] not in clients:
				clients.append(row[0])
		clients.reverse()
		self.comboBox_4.clear()
		self.comboBox_4.addItem('')
		self.comboBox_4.addItems(clients)

	def my_def2(self):
		self.cur.execute(''' SELECT * FROM addclient group by client_name,date(date) ''')
		data = self.cur.fetchall()
		all_clients = []
		all_clients_id = []
		for index in range(0, len(data)):
			all_clients.append(data[index][1])
			all_clients_id.append(data[index][0])
		self.tableWidget_2.setSortingEnabled(False)
		self.tableWidget_2.setRowCount(0)
		self.tableWidget_2.insertRow(0)
		for row, form in enumerate(data):
			for col, item in enumerate(form):
				if col == 0:
					self.tableWidget_2.setItem(row, col,
											   QTableWidgetItem(str(all_clients_id[all_clients.index(data[row][1])])))
				elif col == 3:
					self.tableWidget_2.setItem(row, col, QTableWidgetItem(str(data[row][4])))
				elif col == 4:
					self.tableWidget_2.setItem(row, col, QTableWidgetItem(str(data[row][5])[:10]))
				else:
					self.tableWidget_2.setItem(row, col, QTableWidgetItem(str(item)))
				col += 1
			row_pos = self.tableWidget_2.rowCount()
			self.tableWidget_2.insertRow(row_pos)
		self.tableWidget_2.setSortingEnabled(True)

	def add_Analyst_to_list(self):
		global analysts_name_glo
		self.cur.execute(''' SELECT name FROM addanalyst ''')
		data = self.cur.fetchall()
		analysts_name_glo.clear()
		for i in data:
			if i[0] not in analysts_name_glo:
				analysts_name_glo.append(i[0])

	def Auto_Complete(self, model):
		model.setStringList(self.RETURN_ALL_ANALYST())

	def Auto_complete_combo(self):
		combo = self.comboBox_16, self.comboBox_21
		completer = QCompleter()
		for i in combo:
			i.setCompleter(completer)
			model = QStringListModel()
			completer.setModel(model)
			completer.setCaseSensitivity(Qt.CaseInsensitive)
			self.Auto_Complete(model)

	def add_client_to_list(self):
		global clients_name_glo
		clients_name_glo.clear()
		self.cur.execute(''' SELECT client_name FROM addclient  ''')
		data = self.cur.fetchall()
		for i in data:
			if i[0] not in clients_name_glo:
				clients_name_glo.append(i[0])

	def Auto_Complete2(self, model):
		model.setStringList(clients_name_glo)

	def Auto_complete_combo2(self):
		combo = self.lineEdit_27
		completer = QCompleter()
		combo.setCompleter(completer)
		model = QStringListModel()
		completer.setModel(model)
		completer.setCaseSensitivity(Qt.CaseInsensitive)
		self.Auto_Complete2(model)

	def add_today_client_to_list(self):
		global clients_name_glo2
		clients_name_glo2.clear()
		self.cur.execute(''' SELECT client_name FROM addclient where date(date)=%s ''', (str(date.today()),))
		data = self.cur.fetchall()
		for i in data:
			if i[0] not in clients_name_glo2:
				clients_name_glo2.append(i[0])

	def Auto_Complete7(self, model):
		model.setStringList(clients_name_glo2)

	def Auto_complete_combo7(self):
		combo = self.lineEdit_25
		completer = QCompleter()
		combo.setCompleter(completer)
		model = QStringListModel()
		completer.setModel(model)
		completer.setCaseSensitivity(Qt.CaseInsensitive)
		self.Auto_Complete7(model)

	def add_client_to_list4(self):
		self.cur.execute(''' SELECT client_name FROM addclient''')
		data = self.cur.fetchall()
		for i in data:
			if i[0] not in clients_name_glo_clients_page:
				clients_name_glo_clients_page.append(i[0])

	def Auto_Complete4(self, model):
		model.setStringList(clients_name_glo_clients_page)

	def Auto_complete_combo4(self):
		combo = self.lineEdit_27
		completer = QCompleter()
		combo.setCompleter(completer)
		model = QStringListModel()
		completer.setModel(model)
		completer.setCaseSensitivity(Qt.CaseInsensitive)
		self.Auto_Complete4(model)

	def Delete_Files(self):
		self.cur.execute(''' SELECT * FROM paths WHERE id=1 ''')
		mydata = self.cur.fetchone()
		word_files = mydata[1]
		save_word_files = mydata[2]
		try:
			if path.exists(r'%s\result.docx' % save_word_files):
				remove(r'%s\result.docx' % save_word_files)
			if path.exists(r'%s\result2.docx' % save_word_files):
				remove(r'%s\result2.docx' % save_word_files)
			if path.exists(r'%s\result3.docx' % save_word_files):
				remove(r'%s\result3.docx' % save_word_files)
		except Exception as e:
			word = client.Dispatch("Word.Application")
			word.ActiveDocument.Close()
			self.Delete_Files()
			print(e, '7erorr')

	def Show_paths(self):
		self.cur.execute(''' SELECT * FROM paths WHERE id=1 ''')
		mydata = self.cur.fetchone()
		word_files = mydata[1]
		save_word_files = mydata[2]
		self.lineEdit_14.setText(word_files)
		self.lineEdit_15.setText(save_word_files)

	def Add_Path(self):
		global word_files
		global save_word_files
		global word_data
		files_path = self.lineEdit_14.text()
		save_files_path = self.lineEdit_15.text()
		self.cur.execute(''' UPDATE paths SET file_path=%s,save_file_path=%s WHERE id=1''',
						 (files_path, save_files_path))
		self.db.commit()
		word_files = files_path
		save_word_files = save_files_path
		QMessageBox.information(self, '', 'تم تطبيق المعلومات بنجاح')
		self.Show_paths()
		self.Update_Word_Info()

	def Preview(self):
		global if_print
		self.Print_Sale_Data('T')
		self.cur.execute(''' SELECT * FROM paths WHERE id=1 ''')
		mydata = self.cur.fetchone()
		word_files = mydata[1]
		save_word_files = mydata[2]
		warning = QMessageBox.warning(self, '', 'هل انتهيت من المعاينة؟', QMessageBox.Yes | QMessageBox.No)
		try:
			if_print = True
			if path.exists(r'%s\result.docx' % save_word_files):
				remove(r'%s\result.docx' % save_word_files)
			if path.exists(r'%s\result2.docx' % save_word_files):
				remove(r'%s\result2.docx' % save_word_files)
			if path.exists(r'%s\result3.docx' % save_word_files):
				remove(r'%s\result3.docx' % save_word_files)
		except:
			word = client.Dispatch("Word.Application")
			try:
				word.ActiveDocument.Close()
			except:
				try:
					system("TASKKILL /F /IM WINWORD.exe")
					system('start WINWORD.exe')
				except:
					QMessageBox.information(self, '', 'يرجى اغلاق برنامج MS Word')
			self.Delete_Files()

	def Reset_password(self):
		try:
			user_name = self.lineEdit_7.text()
			self.cur.execute(''' SELECT * FROM adduser ''')
			data = self.cur.fetchall()
			ruser_name = ''
			a = 0
			for row in data:
				if row[1] == user_name:
					ruser_name = row[1]
				else:
					a = 5
			if a == 5:
				QMessageBox.information(self, 'info', 'اسم المستخدم الذي ادخلته غير صحيح')
			self.cur.execute(''' SELECT user_email,user_password FROM adduser WHERE user_name=%s ''', (ruser_name,))
			email_data = self.cur.fetchone()
			email = "ameersaad810@gmail.com"
			password = "aahmpredtiddvxlo"
			send_to_email = email_data[0]
			subject = "n"
			message = f'Hello.\n your password is {email_data[1]} '
			msg = MIMEMultipart()
			msg["From"] = email
			msg["To"] = send_to_email
			msg["Subject"] = subject
			msg.attach(MIMEText(message, 'plain'))
			server = smtplib.SMTP("smtp.gmail.com", 587)
			server.starttls()
			server.login(email, password)
			text = msg.as_string()
			server.sendmail(email, send_to_email, text)
			server.quit()
			print('ok')
			QMessageBox.information(self, 'تم بنجاح', f'تم ارسال كلمة المرور الى{send_to_email} ')
		except Exception as e:
			print(e, '97erorr')
			QMessageBox.information(self, 'خطأ', f'هنالك خطأ يرجى\n ارسال هذا النص {e} الى ameersaad810@gmail.com')

	def clear_data_in_sales(self):
		self.Update_addNewItem_Data()
		self.comboBox_17.clear()
		self.tableWidget_5.setRowCount(0)
		self.tableWidget_5.insertRow(0)
		self.comboBox_4.clear()
		self.lineEdit_21.setText('')
		self.spinBox_7.setValue(20)
		self.doubleSpinBox_7.setValue(0)
		self.spinBox_4.setValue(0)
		self.comboBox_14.setCurrentIndex(0)
		self.comboBox_16.setCurrentIndex(0)
		self.comboBox_17.setCurrentIndex(0)
		self.comboBox_15.setEnabled(True)
		self.comboBox_15.setCurrentIndex(0)
		self.textEdit.setPlainText('')
		self.comboBox_4.setEnabled(True)
		self.spinBox_7.setEnabled(True)
		self.comboBox_14.setEnabled(True)
		self.textEdit.setEnabled(True)
		self.spinBox.setValue(0)
		self.add_clients_to_combo()

	def Show_All_Clients(self):
		global show_clients_check
		if show_clients_check:
			client_name = self.lineEdit_27.text()
			if client_name != '0' and client_name != '':
				self.cur.execute(''' SELECT * FROM addclient where client_name=%s group by client_name,date(date) ''',
								 (client_name,))
			data = self.cur.fetchall()
			all_clients = []
			all_clients_id = []
			for index in range(0, len(data)):
				all_clients.append(data[index][1])
				all_clients_id.append(data[index][0])
			self.tableWidget_2.setSortingEnabled(False)
			self.tableWidget_2.setRowCount(0)
			self.tableWidget_2.insertRow(0)
			try:
				if client_name != '0' and client_name != '':
					for row, form in enumerate(data):
						for col, item in enumerate(form):
							if col == 0:
								self.tableWidget_2.setItem(row, col, QTableWidgetItem(
									str(all_clients_id[all_clients.index(data[row][1])])))
							elif col == 3:
								self.tableWidget_2.setItem(row, col, QTableWidgetItem(str(data[row][4])))
							elif col == 4:
								self.tableWidget_2.setItem(row, col, QTableWidgetItem(str(data[row][5])[:10]))
							else:
								self.tableWidget_2.setItem(row, col, QTableWidgetItem(str(item)))
							col += 1
						row_pos = self.tableWidget_2.rowCount()
						self.tableWidget_2.insertRow(row_pos)
				else:
					self.my_def2()
			except Exception as e:
				print(e, '9gy5erorr')
			self.tableWidget_2.setSortingEnabled(True)
		else:
			self.tableWidget_2.setRowCount(0)
			self.tableWidget_2.insertRow(0)

	def Search_In_History(self):
		self.tableWidget_8.setSortingEnabled(False)
		actionsd = self.comboBox_24.currentIndex()
		tabley = self.comboBox_20.currentIndex()
		if actionsd != 0 and tabley == 0:
			try:
				self.cur.execute(f'SELECT uid,action,tabled,dates FROM his WHERE action = {actionsd} ORDER BY -dates')
			except Exception as e:
				print(e, '8erorr')
		elif tabley != 0 and actionsd == 0:
			try:
				self.cur.execute(f'SELECT uid,action,tabled,dates FROM his WHERE tabled={tabley} ORDER BY -dates')
			except Exception as e:
				print(e, '9erorr')
		elif actionsd != 0 and tabley != 0:
			try:
				self.cur.execute(
					f' SELECT uid,action,tabled,dates FROM his WHERE action={actionsd} AND tabled={tabley} ORDER BY -dates')
			except Exception as e:
				print(e, '10erorr')
		else:
			try:
				self.cur.execute(
					f' SELECT uid,action,tabled,dates FROM his ORDER BY -dates')
			except Exception as e:
				print(e, '11erorr')
		data = self.cur.fetchall()
		self.tableWidget_8.setRowCount(0)
		self.tableWidget_8.insertRow(0)
		for row, form in enumerate(data):
			for col, item in enumerate(form):
				if col == 0:
					self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(data[row][0])))
				if col == 1:
					action = ''
					if data[row][1] == 1:
						action = 'تسجيل الدخول'
					if data[0][1] == 2:
						action = 'تسجيل الخروج'
					if data[row][1] == 3:
						action = 'اضافة'
					if data[row][1] == 4:
						action = 'تعديل'
					if data[row][1] == 5:
						action = 'حذف'
					if data[row][1] == 6:
						action = 'بحث'
					if data[row][1] == 7:
						action = 'طباعة'
					if data[row][1] == 8:
						action = 'معاينة'
					self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(action)))
				if col == 2:
					tables = ''
					if data[row][2] == 1:
						tables = 'مبيع يومي'
					if data[row][2] == 2:
						tables = 'تحليل'
					if data[row][2] == 3:
						tables = 'مشتريات'
					if data[row][2] == 4:
						tables = 'مراجعين'
					if data[row][2] == 5:
						tables = 'مستخدم'
					if data[row][2] == 6:
						tables = 'طبيب'
					if data[row][2] == 7:
						tables = 'تقرير'
					if data[row][2] == 8:
						tables = 'احصائيات'
					self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(tables)))
				if col == 3:
					self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(data[row][3])))
				col += 1
			row_pos = self.tableWidget_8.rowCount()
			self.tableWidget_8.insertRow(row_pos)
		self.tableWidget_8.setSortingEnabled(True)

	def Search_In_All_Sales(self):
		client_name = self.lineEdit_25.text()
		if client_name != '0':
			self.cur.execute(''' SELECT * FROM addclient WHERE client_name=%s AND DATE(date)=%s ORDER BY id''',
							 (client_name, date.today(),))
		else:
			self.cur.execute(''' SELECT * FROM addclient WHERE DATE(date)=%s ORDER BY -date''',
							 (date.today(),))
		analyst_data = self.cur.fetchall()
		self.tableWidget_6.setSortingEnabled(True)
		self.tableWidget_6.setRowCount(0)
		self.tableWidget_6.insertRow(0)
		try:
			if client_name != '0' and client_name != '':
				for i, k in enumerate(analyst_data[0]):
					if i == 3:
						self.tableWidget_6.setItem(0, i, QTableWidgetItem(str(analyst_data[0][4])))
					else:
						self.tableWidget_6.setItem(0, i, QTableWidgetItem(str(k)))
			else:
				self.Show_All_The_Sales()
		except Exception as e:
			print(e, '95erlplorr')
			self.tableWidget_6.setRowCount(0)
			self.tableWidget_6.insertRow(0)
		self.Add_Data_To_history(6, 1)
		# self.History()
		self.tableWidget_6.setSortingEnabled(True)

	def Print_Sale_Data(self, prev):
		genuses = self.comboBox_14.currentText()
		all_analyst = []
		all_result = []
		date = datetime.now()
		day = date.year
		month = date.month
		year = date.day
		real_name = self.comboBox_4.currentText()
		real_doctor = self.comboBox_15.currentText()
		categorys = []
		prev_category = self.tableWidget_5.item(0, 4).text()
		all_analyst.append(prev_category)
		all_result.append('')
		for row in range(0, self.tableWidget_5.rowCount() - 1):
			result = ''
			categorys.append(self.tableWidget_5.item(row, 4).text())
			if row != 0:
				if self.tableWidget_5.item(row, 4).text() != prev_category:
					all_analyst.append(self.tableWidget_5.item(row, 4).text())
					all_result.append('')
					prev_category = self.tableWidget_5.item(row, 4).text()
			try:
				analyst = self.tableWidget_5.item(row, 0).text()
				all_analyst.append(analyst)
			except Exception as e:
				print(e, '95ehiororr')
			try:
				result = self.tableWidget_5.cellWidget(row, 1).value()
			except:
				try:
					result = self.tableWidget_5.cellWidget(row, 1).currentText()
					result = result.replace("$", "\n")

				except:
					try:
						result = self.tableWidget_5.cellWidget(row, 1).text()
						result = result.replace("$", "\n")
					except Exception as e:
						print(e, '95ek;okpororr')
			all_result.append(result)
		# try:
		self.Bio_Word(real_name, real_doctor, all_analyst, all_result, year, month, day, prev, genuses, categorys)
		# except Exception as e:
		#     print(e, '12erorr')
		#     QMessageBox.information(self, 'خطأ',
		#                             f'هنالك خطأ يرجى مراجعة العملية او\n ارسال هذا النص {e} الى ameersaad810@gmail.com')
		# print(all_analyst)
		# print(all_result)
		if prev != 'T':
			self.Delete_Files()
		self.Update_addNewItem_Data()

	def Delete_Row(self, item):
		try:
			if self.tableWidget_5.rowCount() > 2:
				analyst_name = self.tableWidget_5.item(item, 0).text()
				client_name = self.comboBox_4.currentText()
				warning = QMessageBox.warning(self, 'احذر', f"سوف يتم مسح العنصر{analyst_name} هل انت متأكد؟",
											  QMessageBox.Yes | QMessageBox.No)
				if warning == QMessageBox.Yes:
					self.tableWidget_5.removeRow(item)
					self.cur.execute(''' select MIN(id) from addclient WHERE client_name=%s''',
									 (client_name,))
					min_id = self.cur.fetchone()
					self.cur.execute(
						''' DELETE FROM addnewitem WHERE client_name=%s AND analyst_name=%s AND DATE(date)=%s''',
						(client_name, analyst_name, str(date.today()),))
					self.db.commit()
					self.cur.execute(''' DELETE FROM addclient WHERE client_name=%s AND id!=%s limit 1''',
									 (client_name, min_id[0],))
					self.db.commit()
			else:
				analyst_name = self.tableWidget_5.item(item, 0).text()
				QMessageBox.information(self, 'تحذير',
										f'لا يمكن حذف كل العناصر اضف عنصر جديد لكي تستطيع حذف{analyst_name}')
		except Exception as e:
			print(e, '13erorr')

	def Show_Category_Dialog(self):
		self.Category_Dailog = add_delete_category_dialogpy.Dialog()
		self.Category_Dailog.comboBox_11.currentTextChanged.connect(self.C_OR_U_Category)
		self.Category_Dailog.comboBox_11.setEditable(True)
		self.Category_Dailog.comboBox_11.addItems(self.add_all_subCategory_toList())
		self.Category_Dailog.pushButton_18.clicked.connect(self.CUD_Category)
		self.Category_Dailog.pushButton_16.clicked.connect(self.CUD_Category)
		self.Category_Dailog.pushButton_17.clicked.connect(self.Delete_Category)
		self.Category_Dailog.exec_()

	def C_OR_U_Category(self):
		comboText = self.Category_Dailog.comboBox_11.currentText()
		if comboText in self.add_all_subCategory_toList():
			self.Category_Dailog.pushButton_18.show()
			self.Category_Dailog.pushButton_16.hide()
			self.lineEdit_13.setText(self.comboBox_11.currentText())
		else:
			self.Category_Dailog.pushButton_18.hide()
			self.Category_Dailog.pushButton_16.show()

	def CUD_Category(self):
		name = self.Category_Dailog.lineEdit_13.text()
		b_name = self.Category_Dailog.comboBox_11.currentText()
		if self.sender().text() == 'اضافة':
			try:
				self.cur.execute(''' INSERT INTO category (name) VALUES(%s) ''', (name,))
				self.db.commit()
				QMessageBox.information(self, '', 'تم اضافة صنف جديد بنجاح')
			except Exception as e:
				print('jjjjj', e)
				QMessageBox.information(self, '', 'هذا الصنف موجود بالفعل')
		else:
			self.cur.execute(''' UPDATE category set name=%s WHERE name=%s ''', (name, b_name))
			self.cur.execute(''' UPDATE addanalyst set sub_category=%s where sub_category=%s ''', (name, b_name,))
			self.cur.execute(''' UPDATE addnewitem set sub_category=%s where sub_category=%s ''', (name, b_name,))
			self.db.commit()
			QMessageBox.information(self, '', 'تم تعديل الصنف بنجاح')
		self.Category_Dailog.close()
		self.comboBox_23.clear()
		self.comboBox_23.addItems(self.add_all_subCategory_toList())

	def Delete_Category(self):
		b_name = self.Category_Dailog.comboBox_11.currentText()
		warning = QMessageBox.warning(self, 'احذر',
									  f"سوف يتم مسح الصنف {b_name} وجميع التحاليل الخاصة به, هل انت متأكد؟",
									  QMessageBox.Yes | QMessageBox.No)
		if warning == QMessageBox.Yes:
			self.cur.execute(''' DELETE FROM category where name=%s ''', (b_name,))
			self.cur.execute(''' DELETE FROM addanalyst where sub_category=%s ''', (b_name,))
			self.db.commit()
			QMessageBox.information(self, '', 'تم حذف الصنف بنجاح')
		self.comboBox_23.clear()
		self.comboBox_23.addItems(self.add_all_subCategory_toList())
		self.Category_Dailog.close()

	def ShowIsInCategory(self):
		if self.tabWidget_3.currentIndex() == 2:
			self.tabWidget_3.setCurrentIndex(0)
			self.Show_Category_Dialog()

	def Delete_Row2(self, item):
		self.tableWidget_3.setSortingEnabled(False)
		try:
			if self.tableWidget_3.rowCount() > 1:
				item_name = self.tableWidget_3.item(item, 1).text()
				item_quantity = self.tableWidget_3.item(item, 1).text()
				item_price = self.tableWidget_3.item(item, 2).text()
				warning = QMessageBox.warning(self, 'احذر', f"سوف يتم مسح العنصر{item_name} هل انت متأكد؟",
											  QMessageBox.Yes | QMessageBox.No)
				if warning == QMessageBox.Yes:
					self.cur.execute(
						''' DELETE FROM addbuys WHERE item_name=%s AND signal_item_price=%s AND quantity=%s''',
						(item_name, int(item_price), int(item_quantity),))
					self.db.commit()
					self.Show_all_buys()
		except Exception as e:
			print(e, '95erorrlddeee')
		self.tableWidget_3.setSortingEnabled(True)

	def Add_Doctor_Data(self):
		self.comboBox_28.clear()
		self.comboBox_15.clear()
		self.comboBox_28.addItem('---------------')
		self.cur.execute(''' select * from doctor ''')
		data = self.cur.fetchall()
		for i in range(0, len(data)):
			self.comboBox_28.addItem(str(data[i][1]))
			self.comboBox_15.addItem(str(data[i][1]))

	def buttonClicked(self):
		button = self.sender()
		index = self.tableWidget_5.indexAt(button.pos())
		self.Delete_Row(index.row())

	def get_total_price(self):# MCV
		# self.cur.execute(''' select full_price from price_task where client_name=%s and date(date)=%s ''',(self.comboBox_4.currentText(),date.today()))
		# data = self.cur.fetchone()
		# if data:
		# 	self.lineEdit_24.setText(str(data[0]))
		# else:
		total_price = 0
		for row in range(0, self.tableWidget_5.rowCount() - 1):
			try:
				analyst_name = self.tableWidget_5.item(row, 0).text()
			except:
				analyst_name = ''
			try:
				a = self.tableWidget_5.item(row, 2).text()
			except:
				a = 0
			try:
				total_price += int(a)
			except ValueError:
				a = 0
		self.lineEdit_24.setText(str(total_price))

	def Update_addNewItem_Data(self):
		r2_client_name = self.comboBox_4.currentText()
		r2_doctor = self.comboBox_15.currentText()
		for rowj in range(0, self.tableWidget_5.rowCount() - 1):
			# price = self.tableWidget_5.item(rowj, 2).text()
			try:
				r2_analyst_name = self.tableWidget_5.item(rowj, 0).text()
			except:
				r2_analyst_name = ''
			r2_result = None
			try:
				r2_result = self.tableWidget_5.cellWidget(rowj, 1).value()
			except:
				try:
					r2_result = self.tableWidget_5.cellWidget(rowj, 1).currentText()
				except:
					try:
						r2_result = self.tableWidget_5.cellWidget(rowj, 1).text()
					except:
						try:
							r2_result = self.tableWidget_5.cellWidget(rowj, 1).text()
						except Exception as e:
							print(e, '99erorr')
							r2_result = ''
			try:
				self.cur.execute(
					''' UPDATE addnewitem SET  analyst_name=%s ,analyst_result=%s WHERE client_name=%s AND analyst_name=%s''',
					(r2_analyst_name, r2_result, r2_client_name, r2_analyst_name))
				self.db.commit()
			except:
				print('except')

	def Analyst_sub_category(self, analyst_name):
		self.cur.execute(''' select sub_category,price from addanalyst where name=%s''', (analyst_name,))
		data = self.cur.fetchone()
		return data

	def Sales_Page(self, for_loop2=None):
		global clients_name_glo
		not_in_all_analysts_items = False
		# if for_loop2==False:
		for mjrow in range(0, self.tableWidget_5.rowCount() - 1):
			if self.comboBox_16.currentText() == str(self.tableWidget_5.item(mjrow, 0).text()):
				not_in_all_analysts_items = True
		if not not_in_all_analysts_items:
			if self.comboBox_4.isEnabled()== True:
				self.comboBox_4.setEnabled(False)
				self.spinBox_7.setEnabled(False)
				self.comboBox_14.setEnabled(False)
				self.comboBox_15.setEnabled(False)
				self.textEdit.setEnabled(False)
			global client_id_glob
			global chick_if_add_new
			global clients_name_glo
			analyst_name = self.comboBox_16.currentText()
			analyst_lineEdit_result = self.lineEdit_21.text()
			analyst_combo_result = self.comboBox_17.currentText()
			analyst_float_result = self.doubleSpinBox_7.value()
			analyst_number_result = self.spinBox_4.value()
			
			self.cur.execute('''SELECT price,category,sub_category,analyst_index,quantity FROM addanalyst WHERE name = %s''',
							 (analyst_name,))
			analyst_price = self.cur.fetchone()
			
			CUI2=False
			if analyst_price:
				if for_loop2 == False:
					if analyst_price[4]!=None:
						if analyst_price[4] == 0:
							warning = QMessageBox.warning(self,"احذر","نفذت كمية هذا التحليل هل تريد المواصلة؟",QMessageBox.Yes | QMessageBox.No)
							if warning == QMessageBox.Yes:
								CUI2=True
						else:
							CUI2 =True
					else:
						CUI2 =True
				else:
					CUI2=True
			if CUI2:
				analyst_index = 1
				if analyst_price:
					if analyst_price[3]:
						analyst_index = int(analyst_price[3])
				latest_result = 1
				total_price = 0
				if analyst_price != None:
					if analyst_price[1] == 'عدد':
						latest_result = analyst_number_result
					if analyst_price[1] == 'خيارات':
						latest_result = analyst_combo_result
					if analyst_price[1] == 'حقل كتابة':
						latest_result = analyst_lineEdit_result
					if analyst_price[1] == 'خيارات مع تعديل':
						latest_result = analyst_combo_result
					if analyst_price[1] == 'عدد عشري':
						latest_result = analyst_float_result

					total_price = int(analyst_price[0])
					client_name = self.comboBox_4.currentText()
					client_age = self.spinBox_7.value()
					client_doctor = self.comboBox_15.currentText()
					client_genus = self.comboBox_14.currentText()
					analyst_or_clients_notes = self.textEdit.toPlainText()
					if client_name not in clients_name_glo:					
						self.cur.execute('''
								INSERT INTO addclient (client_name,client_age,client_genus,client_doctor,date)
								VALUES (%s,%s,%s,%s,%s)
								''', (
							str(client_name), str(client_age), str(client_genus), str(client_doctor),
							str(datetime.now())))
						self.db.commit()
					self.cur.execute('''SELECT id FROM addclient WHERE client_name = %s''', (client_name,))
					real_client_id = self.cur.fetchone()
					self.cur.execute('''
					INSERT INTO addnewitem (client_name,client_id,client_age,genus,doctor_name,notes,analyst_name,analyst_result,price,total_price,date,sub_category,analyst_index)
					VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
					''', (
						client_name, real_client_id, client_age, client_genus, client_doctor, analyst_or_clients_notes,
						analyst_name,
						latest_result,
						total_price, total_price, datetime.now(), analyst_price[2], analyst_index,))
					self.db.commit()
					clients_name_glo.append(str(client_name))
					if for_loop2 == False:
						self.Update_addNewItem_Data()
						self.Show_All_one_client_analyst('lets see')	
				else:
					if for_loop2 == False:
						QMessageBox.information(self, '', 'هذا التحليل غير موجود')
				if for_loop2 == False:
					chick_if_add_new = True
					# # self.Add_buttons_combo_spin_to_tableWidget()
					self.tableWidget_5.scrollToBottom()
					self.get_total_price()
					# self.my_def2()
					self.Add_Data_To_history(3, 1)
					# self.History()
					# self.cur.execute(''' update addbuys set item_quantity=item_quantity -1 where to_analysts=%s and item_quantity!=0''',(analyst_name,))
					self.cur.execute(''' update addanalyst set quantity = quantity - 1 where name = %s and quantity!=0''',(analyst_name,))
					self.db.commit()
				else:
					pass
		else:
			if for_loop2 == False:
				QMessageBox.information(self, 'تحذير', 'هذا العنصر موجود بالفعل')
				# self.Add_all_analysts_items()
			else:
				pass
		not_in_all_analysts_items = False
		# if for_loop2 == False:
			# self.add_today_client_to_list()
			# self.Auto_complete_combo7()
			# self.Show_All_The_Sales()

	def Add_buttons_combo_spin_to_tableWidget(self):
		name = self.comboBox_4.currentText()
		for rowd in range(0, self.tableWidget_5.rowCount() - 1):
			mypush_button = QPushButton(self)
			mypush_button.setText('حذف')
			mypush_button.setStyleSheet(
				'''QPushButton:pressed{background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:0.687, y2:0.704545, stop:0.0646766 rgba(255, 155, 155, 255), stop:0.751244 rgba(235, 54, 30, 255));}QPushButton{border-radius:12px;background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:1, y2:1, stop:0.512438 rgba(235, 54, 30, 255), stop:0.875622 rgba(255, 155, 155, 255));}''')
			mypush_button.clicked.connect(self.buttonClicked)
			try:
				row_analyst_name = self.tableWidget_5.item(rowd, 0).text()
				if row_analyst_name and row_analyst_name != '':
					self.tableWidget_5.setCellWidget(rowd, 3, mypush_button)
			except Exception as e:
				print(e, '95erorr')
			all_name_items = []
			rs_name = ''
			try:
				rs_name = self.tableWidget_5.item(rowd, 0).text()
			except Exception as e:
				print(e, '93erorr')
				rs_name = ''
			self.cur.execute('''SELECT results,category FROM addanalyst WHERE name=%s''', (rs_name,))
			results_data = self.cur.fetchall()
			mycobmbo = None
			object_type = ''
			if results_data:
				if results_data[0][1] == 'خيارات':
					mycobmbo = QComboBox(self)
					x = '[' + str(results_data[0][0]) + ']'
					list_data = literal_eval(str(x))
					# mycobmbo.addItems(list_data)
					model = mycobmbo.model()
					for rowx in list_data:
						item = QStandardItem(str(rowx)+'gg')
						item.setForeground(QColor('red'))
						model.appendRow(item)
					object_type = 'خيارات'
				if results_data[0][1] == 'عدد':
					mycobmbo = QSpinBox(self)
					mycobmbo.setMaximum(2147483647)
					mycobmbo.setMinimum(-2147483647)
					object_type = 'عدد'
				if results_data[0][1] == 'عدد عشري':
					mycobmbo = QDoubleSpinBox(self)
					mycobmbo.setDecimals(5)
					mycobmbo.setMaximum(2147483647)
					mycobmbo.setMinimum(-2147483647)
					object_type = 'عدد عشري'
				if results_data[0][1] == 'خيارات مع تعديل':
					mycobmbo = QComboBox(self)
					mycobmbo.setEditable(True)
					x = '[' + str(results_data[0][0]) + ']'
					list_data = literal_eval(str(x))
					print(list_data, type(list))
					for rr in list_data:
						mycobmbo.addItem(str(rr))
					object_type = 'خيارات مع تعديل'
				if results_data[0][1] == 'حقل كتابة':
					mycobmbo = QLineEdit(self)
					object_type = 'حقل كتابة'
				try:
					r2_analyst_name = self.tableWidget_5.item(rowd, 0).text()
					self.cur.execute(
						''' SELECT  analyst_result  FROM addnewitem WHERE client_name=%s AND analyst_name=%s ''',
						(name, r2_analyst_name))
					myrs = self.cur.fetchall()
					if myrs != None:
						if object_type == 'خيارات' or object_type == 'خيارات مع تعديل':
							index = mycobmbo.findText(myrs[0][0], Qt.MatchFixedString)
							print('here index', index)
							if index != -1:
								mycobmbo.setCurrentIndex(index)
							else:
								print('countinu')
								if myrs[0][0] != mycobmbo.currentText():
									print('couy3')
									mycobmbo.addItem(str(myrs[0][0]))
									index = mycobmbo.findText(myrs[0][0], Qt.MatchFixedString)
									mycobmbo.setCurrentIndex(index)
						if object_type == 'عدد':
							if myrs[0][0] and myrs[0][0] != '':
								mycobmbo.setValue(int(myrs[0][0]))
						if object_type == 'عدد عشري':
							if myrs[0][0] and myrs[0][0] != '':
								mycobmbo.setValue(float(myrs[0][0]))
						if object_type == 'حقل كتابة':
							mycobmbo.setText(str(myrs[0][0]))
					self.tableWidget_5.setItem(rowd, 1, QTableWidgetItem(str('')))
					self.tableWidget_5.setCellWidget(rowd, 1, mycobmbo)
					if object_type == 'خيارات' or object_type == 'خيارات مع تعديل':
						if mycobmbo.currentText() == '' or mycobmbo.currentText() == ' ':
							mycobmbo.setCurrentIndex(0)
				except Exception as e:
					print(e, '91erorr')

	def Show_All_one_client_analyst(self, from_add_multy=None):
		global client_id_glob
		global chick_if_add_new
		global select_by_date
		global CJ
		global MAX_ID
		client_name = self.comboBox_4.currentText()
		CV = []
		if from_add_multy:
			self.cur.execute('''SELECT id FROM addclient WHERE client_name = %s''', (client_name,))
			real_client_id = self.cur.fetchone()
			if real_client_id:
				if client_name != '' and client_name != ' ':
					self.spinBox.setValue(real_client_id[0])
		if self.spinBox.value() == 0 and chick_if_add_new == False:
			self.tableWidget_5.setRowCount(0)
			self.tableWidget_5.insertRow(0)
			self.lineEdit_21.setText('')
			self.spinBox_7.setValue(0)
			self.doubleSpinBox_7.setValue(0)
			self.spinBox_4.setValue(0)
			if not from_add_multy:
				self.comboBox_14.setCurrentIndex(0)
				self.comboBox_15.setCurrentIndex(0)
			self.comboBox_17.setCurrentIndex(0)
			self.textEdit.setPlainText('')
			self.comboBox_4.setEnabled(True)
			self.spinBox_7.setEnabled(True)
			self.comboBox_14.setEnabled(True)
			self.textEdit.setEnabled(True)
		if True:
			if select_by_date == True:
				idn = self.searchWidget.spinBox.value()
				from_date = self.searchWidget.dateEdit_6.date()
				to_date = self.searchWidget.dateEdit_5.date()
				self.cur.execute(
					'''SELECT client_name,analyst_name,analyst_result,doctor_name,total_price,sub_category,client_age,genus,notes,analyst_index FROM addnewitem WHERE client_id = %s AND DATE(date)>=%s AND DATE(date)<=%s ORDER BY sub_category ASC, analyst_index ASC''',
					(idn, str(from_date.toPyDate()), str(to_date.toPyDate()),))
			elif select_by_date == 'gg':
				idn = self.spinBox.value()
				self.cur.execute(
					'''SELECT client_name,analyst_name,analyst_result,doctor_name,total_price,sub_category,client_age,genus,notes,analyst_index FROM addnewitem WHERE client_id = %s ORDER BY sub_category ASC, analyst_index ASC''',
					(idn,))
			else:
				self.cur.execute(
					'''SELECT client_name,analyst_name,analyst_result,doctor_name,total_price,sub_category,client_age,genus,notes,analyst_index FROM addnewitem  WHERE client_id = %s AND DATE(date)=%s ORDER BY sub_category ASC, analyst_index ASC''',
					(self.spinBox.value(), date.today(),))
			analyst_data = self.cur.fetchall()
			if analyst_data:
				# if self.comboBox_14.isEnabled() == True:
				self.comboBox_15.setCurrentText(str(analyst_data[0][3]))
				self.comboBox_15.setEnabled(False)
				self.comboBox_14.setEnabled(False)
				self.comboBox_4.setCurrentText(analyst_data[0][0])
				self.comboBox_4.setEnabled(False)
				self.spinBox_7.setValue(int(analyst_data[0][6]))
				self.spinBox_7.setEnabled(False)
				self.comboBox_14.setCurrentText(analyst_data[0][7])
				self.comboBox_14.setEnabled(False)
				self.textEdit.setPlainText(str(analyst_data[0][8]))
				self.textEdit.setEnabled(False)
				self.tableWidget_5.setRowCount(0)
				self.tableWidget_5.insertRow(0)
				for row, form in enumerate(analyst_data):
					for col, item in enumerate(form):
							if col == 0:
								self.tableWidget_5.setItem(row, col, QTableWidgetItem(str(analyst_data[row][1])))
								my_item = self.tableWidget_5.item(row, col)
								my_item.setFlags(Qt.ItemIsEditable)
								my_item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable)
							elif col == 1:
								self.cur.execute('''SELECT results,category FROM addanalyst WHERE name=%s''', (analyst_data[row][1],))
								results_data = self.cur.fetchall()
								mycobmbo = None
								object_type = ''
								if results_data:
									if results_data[0][1] == 'خيارات':
										genus_type = self.comboBox_14.currentText()
										self.cur.execute(''' select normal_value1,normal_value2 from analystnormal where analyst_name=%s and genus_type=%s''',(analyst_data[row][1],genus_type,))
										NormalData = self.cur.fetchone()
										mycobmbo = QComboBox()
										x = '[' + str(results_data[0][0]) + ']'
										list_data = literal_eval(str(x))
										mycobmbo.addItems(list_data)
										if NormalData:
											x3 = '[' + str(NormalData[1]) + ']'
											list_data3 = literal_eval(str(x3))
											for cindex,tt in enumerate(list_data3):
												if tt =='غير طبيعي':
													mycobmbo.setItemData(cindex,QColor(Qt.red),Qt.ForegroundRole)
										object_type = 'خيارات'

									if results_data[0][1] == 'عدد':
										mycobmbo = QSpinBox()
										mycobmbo.setMaximum(2147483647)
										mycobmbo.setMinimum(-2147483647)
										object_type = 'عدد'
										CJ = None
										mycobmbo.valueChanged.connect(self.IsNormalForAddNewItem)
									if results_data[0][1] == 'عدد عشري':
										mycobmbo = QDoubleSpinBox()
										mycobmbo.setDecimals(5)
										mycobmbo.setMaximum(2147483647)
										mycobmbo.setMinimum(-2147483647)
										object_type = 'عدد عشري'
										CJ = None
										mycobmbo.valueChanged.connect(self.IsNormalForAddNewItem)
									if results_data[0][1] == 'خيارات مع تعديل':
										genus_type = self.comboBox_14.currentText()
										
										self.cur.execute(''' select normal_value1,normal_value2 from analystnormal where analyst_name=%s and genus_type=%s''',(analyst_data[row][1],genus_type,))
										NormalData = self.cur.fetchone()
										mycobmbo = QComboBox()
										mycobmbo.setEditable(True)
										x = '[' + str(results_data[0][0]) + ']'
										list_data = literal_eval(str(x))
										mycobmbo.addItems(list_data)
										object_type = 'خيارات مع تعديل'
										if NormalData:
											x3 = '[' + str(NormalData[1]) + ']'
											list_data3 = literal_eval(str(x3))
											for cindex,tt in enumerate(list_data3):
												if tt =='غير طبيعي':
													mycobmbo.setItemData(cindex,QColor(Qt.red),Qt.ForegroundRole)
									if results_data[0][1] == 'حقل كتابة':
										mycobmbo = QLineEdit()
										object_type = 'حقل كتابة'
									if analyst_data[row][2]!=None:
										if object_type == 'خيارات' or object_type == 'خيارات مع تعديل':
											index = mycobmbo.findText(str(analyst_data[row][2]), Qt.MatchFixedString)
											if index != -1:
												mycobmbo.setCurrentIndex(index)
											else:
												if analyst_data[row][2] != mycobmbo.currentText():
													mycobmbo.addItem(str(analyst_data[row][2]))
													index = mycobmbo.findText(analyst_data[row][2], Qt.MatchFixedString)
													mycobmbo.setCurrentIndex(index)
										if object_type == 'عدد':
											if analyst_data[row][2] and analyst_data[row][2] != '':
												mycobmbo.setValue(int(analyst_data[row][2]))
												CV.append(row)
										if object_type == 'عدد عشري':
											if analyst_data[row][2] and analyst_data[row][2] != '':
												mycobmbo.setValue(float(analyst_data[row][2]))
												CV.append(row)
										if object_type == 'حقل كتابة':
											mycobmbo.setText(str(analyst_data[row][2]))
									self.tableWidget_5.setItem(row, 1, QTableWidgetItem(str('')))
									self.tableWidget_5.setCellWidget(row, 1, mycobmbo)
									if object_type == 'خيارات' or object_type == 'خيارات مع تعديل':
										if mycobmbo.currentText() == '' or mycobmbo.currentText() == ' ':
											mycobmbo.setCurrentIndex(0)
							elif col == 2:
								self.tableWidget_5.setItem(row, col, QTableWidgetItem(str(analyst_data[row][4])))
								my_item = self.tableWidget_5.item(row, col)
								my_item.setFlags(Qt.ItemIsEditable)
								my_item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable)
							elif col == 3:
								mypush_button = QPushButton(self)
								mypush_button.setText('حذف')
								mypush_button.setStyleSheet(
									'''QPushButton:pressed{background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:0.687, y2:0.704545, stop:0.0646766 rgba(255, 155, 155, 255), stop:0.751244 rgba(235, 54, 30, 255));}QPushButton{border-radius:12px;background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:1, y2:1, stop:0.512438 rgba(235, 54, 30, 255), stop:0.875622 rgba(255, 155, 155, 255));}''')
								mypush_button.clicked.connect(self.buttonClicked)
								self.tableWidget_5.setItem(row, 3, QTableWidgetItem(str('')))
								self.tableWidget_5.setCellWidget(row, 3, mypush_button)
							elif col == 4:
								self.tableWidget_5.setItem(row, col, QTableWidgetItem(str(analyst_data[row][5])))
					col += 1
					row_pos = self.tableWidget_5.rowCount()
					self.tableWidget_5.insertRow(row_pos)
				chick_if_add_new = False
				# # self.Show_All_The_Sales()
				# # self.Add_buttons_combo_spin_to_tableWidget()
			else:
				if not from_add_multy:
					if select_by_date != 'gg':
						QMessageBox.information(self, 'Error','الرقم الذي ادخلته غير موجود في مبيعات اليوم يرجى مراجعة صفحة "مبيعات اليوم" للتأكد من الرقم')
					else:
						QMessageBox.information(self, 'Error','لا يوجد مراجع  في هذا التسلسل')
						self.spinBox.setValue(MAX_ID)
			self.get_total_price()
			for iB in CV:	
				CJ = iB
				self.IsNormalForAddNewItem()
		CJ = None
		self.tableWidget_5.resizeColumnsToContents()
		# self.Add_all_analysts_items()
		self.Add_Data_To_history(6, 1)
		# # self.History()
		select_by_date = False
	def IsNormalForAddNewItem(self):
		# print('ss4'+random.choice(['1','2','3','4','5','r','7','8']))
		global CJ
		try:
			if CJ == None:	
				index = self.sender().pos()	
				index = self.tableWidget_5.indexAt(index)	
				index = index.row()	
			else:	
				index = CJ
			spin = self.tableWidget_5.cellWidget(index,1)
			name = self.tableWidget_5.item(index,0).text()
			genus_type = self.comboBox_14.currentText()
			self.cur.execute(''' select normal_value1,normal_value2 from analystnormal where analyst_name=%s and genus_type=%s ''',(name,genus_type,))
			data =self.cur.fetchone()
			if spin.value() < float(data[0]) or spin.value() > float(data[1]):
				spin.setStyleSheet(';font: 10pt "Segoe UI"; color:#cf2525')
			else:
				spin.setStyleSheet('font: 10pt "Segoe UI";')
		except:
			pass
	def Show_Type_of_result_category(self):
		global addTrue
		analyst_name = self.comboBox_16.currentText()
		self.cur.execute('''SELECT category,results FROM addanalyst WHERE name = %s''', (analyst_name,))
		analyst_category = self.cur.fetchone()
		try:
			if analyst_category:
				if analyst_category[0] == 'عدد':
					self.comboBox_17.hide()
					self.lineEdit_21.hide()
					self.doubleSpinBox_7.hide()
					self.spinBox_4.show()
				if analyst_category[0] == 'عدد عشري':
					self.comboBox_17.hide()
					self.lineEdit_21.hide()
					self.spinBox_4.hide()
					self.doubleSpinBox_7.show()
				if analyst_category[0] == 'خيارات':
					self.lineEdit_21.hide()
					self.doubleSpinBox_7.hide()
					self.spinBox_4.hide()
					self.comboBox_17.show()
					self.comboBox_17.setEditable(False)
				if analyst_category[0] == 'حقل كتابة':
					self.lineEdit_21.show()
					self.doubleSpinBox_7.hide()
					self.spinBox_4.hide()
					self.comboBox_17.hide()
				if analyst_category[0] == 'خيارات مع تعديل':
					self.lineEdit_21.hide()
					self.doubleSpinBox_7.hide()
					self.spinBox_4.hide()
					self.comboBox_17.show()
					self.comboBox_17.setEditable(True)
				try:
					self.comboBox_17.clear()
					x = '[' + str(analyst_category[1]) + ']'
					list_data = literal_eval(str(x))
					for rr in list_data:
						self.comboBox_17.addItem(str(rr))
				except Exception as e:
					print(e, '959erorr')
		except Exception as e:
			print(e, '2erorr')
			if analyst_name not in self.add_all_subCategory_toList():
				if not addTrue:
					QMessageBox.information(self, 'تحذير', "يرجى اختيار تحليل صحيح")

	def analyst_category_function(self, name):
		if name in self.add_all_subCategory_toList():
			self.cur.execute('''SELECT name FROM addanalyst WHERE sub_category=%s ''', (str(name),))
			my_DATA = self.cur.fetchall()
			for new in my_DATA:
				self.comboBox_16.setCurrentText(str(new[0]))
				self.Sales_Page()
			self.add_client_to_list()
			self.Show_All_one_client_analyst()
			self.Update_addNewItem_Data()
			# # self.Add_buttons_combo_spin_to_tableWidget()
			self.tableWidget_5.scrollToBottom()
			self.get_total_price()
			# self.my_def2()
			# self.Show_All_The_Sales()
			# self.add_today_client_to_list()
			# self.Auto_complete_combo7()
			self.Add_Data_To_history(3, 1)
			# self.History()
			# self.Add_all_analysts_items()

	def Chick_analyst_category(self, sd=None):
		if self.comboBox_16.currentText() not in self.add_all_subCategory_toList() and self.comboBox_16.currentIndex() != 0:
			self.analyst_category_function(self.comboBox_16.currentText())
		if self.comboBox_16.currentText() in self.add_all_subCategory_toList():
			try:
				ritem = self.comboBox_16.model().itemFromIndex(sd)
				if ritem.checkState() == Qt.Unchecked:
					ritem.setCheckState(Qt.Checked)
					self.analyst_category_function(ritem.text())
					ritem.setCheckState(Qt.Unchecked)
				else:
					ritem.setCheckState(Qt.Unchecked)
				ritem.setCheckState(Qt.Unchecked)
			except Exception as e:
				print(e, '3erorr')

	def get_client_id(self):
		global client_id_glob
		self.Update_addNewItem_Data()
		client_id_glob = self.spinBox.value()
		self.Show_All_one_client_analyst()

	def Show_All_The_Sales(self):
		global show_all_sales_in_clients_page
		self.cur.execute(''' SELECT client_name FROM addclient WHERE DATE(date)=%s ORDER BY -date ''',
						 (date.today(),))
		analyst_data = self.cur.fetchall()
		all_names = {}
		all_data = []
		id = 0
		for item in analyst_data:
			if item[0] not in all_names:
				all_names[item[0]] = 0
		for s in all_names.keys():
			self.cur.execute(''' SELECT id FROM addclient WHERE client_name=%s ORDER BY id ''', (s,))
			analyst_data2 = self.cur.fetchall()
			all_names[s] = analyst_data2[0][0]
			id = analyst_data2[0][0]
			self.cur.execute(''' SELECT * FROM addclient WHERE id=%s ORDER BY -date ''', (id,))
			latest_data = self.cur.fetchall()
			data = [{'name': s, 'value': latest_data}]
			all_data.append(data)
		self.tableWidget_6.setRowCount(0)
		self.tableWidget_6.insertRow(0)
		l_data = ()
		for ih in all_data:
			hs_data = ih[0]['value']
			l_data += tuple(hs_data)
		for row, form in enumerate(l_data):
			for col, item in enumerate(form):
				if col == 3:
					self.tableWidget_6.setItem(row, col, QTableWidgetItem(str(l_data[row][4])))
				else:
					self.tableWidget_6.setItem(row, col, QTableWidgetItem(str(item)))
				col += 1
			row_pos = self.tableWidget_6.rowCount()
			self.tableWidget_6.insertRow(row_pos)

	def Show_All_The_Analysts(self):
		global from_start
		self.tableWidget_7.setSortingEnabled(False)
		search_type = self.comboBox_19.currentText()
		search_words = self.lineEdit_26.text()
		current_index = self.comboBox_19.currentIndex()
		if current_index == 1:
			self.cur.execute(
				''' SELECT id,name,category,defult,unit,price,sub_category,analyst_index FROM addanalyst WHERE name=%s ''',
				(search_words,))
		if current_index == 2:
			self.cur.execute(
				''' SELECT id,name,category,defult,unit,price,sub_category,analyst_index FROM addanalyst WHERE price=%s ''',
				(search_words,))
		if current_index == 3:
			self.cur.execute(
				''' SELECT id,name,category,defult,unit,price,sub_category,analyst_index FROM addanalyst WHERE sub_category=%s ''',
				(search_words,))
		if current_index == 0:
			self.cur.execute(
				''' SELECT id,name,category,defult,unit,price,sub_category,analyst_index FROM addanalyst ORDER BY sub_category ASC ,analyst_index ASC''')
		analyst_data = self.cur.fetchall()
		self.tableWidget_7.setRowCount(0)
		self.tableWidget_7.insertRow(0)
		for row in range(0, len(analyst_data)):
			for col in range(0, 10):
				if col == 2:
					my_combo2 = QComboBox()
					my_combo2.addItems(['عدد عشري','خيارات', 'عدد', 'خيارات مع تعديل', 'حقل كتابة'])
					index = my_combo2.findText(analyst_data[row][2], Qt.MatchFixedString)
					my_combo2.setCurrentIndex(index)
					self.tableWidget_7.setCellWidget(row, col, my_combo2)
				elif col == 6:
					my_combo = QComboBox()
					my_combo.addItems(self.add_all_subCategory_toList())
					index = my_combo.findText(analyst_data[row][6], Qt.MatchFixedString)
					my_combo.setCurrentIndex(index)
					self.tableWidget_7.setCellWidget(row, col, my_combo)
				elif col == 5:
					my_spin = QSpinBox()
					my_spin.setMaximum(2147483647)
					my_spin.setMinimum(0)
					my_spin.setValue(analyst_data[row][5])
					self.tableWidget_7.setCellWidget(row, col, my_spin)
				elif col == 7:
					mycobmbo = QComboBox(self)
					self.cur.execute(''' select name,analyst_index from addanalyst where sub_category=%s''',
									 (analyst_data[row][6],))
					frData = self.cur.fetchall()
					liost = []
					for ikp in range(1, len(frData) + 1):
						liost.append(str(ikp))
					mycobmbo.addItems(liost)
					try:
						mycobmbo.setCurrentIndex(int(analyst_data[row][7]) - 1)
					except Exception as e:
						mycobmbo.setCurrentIndex(0)
					self.tableWidget_7.setCellWidget(row, col, mycobmbo)
				elif col == 8:
					my_push = QPushButton()
					my_push.setText('حفظ')
					my_push.setStyleSheet('border-radius:12px;')
					my_push.clicked.connect(self.Handel_Save_Delete_analyst)
					self.tableWidget_7.setCellWidget(row, col, my_push)
				elif col == 9:
					my_push2 = QPushButton()
					my_push2.setText('حذف')
					my_push2.setStyleSheet(
						'''QPushButton:pressed{background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:0.687, y2:0.704545, stop:0.0646766 rgba(255, 155, 155, 255), stop:0.751244 rgba(235, 54, 30, 255));}QPushButton{border-radius:12px;background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:1, y2:1, stop:0.512438 rgba(235, 54, 30, 255), stop:0.875622 rgba(255, 155, 155, 255));}''')
					my_push2.clicked.connect(self.Handel_Save_Delete_analyst)
					self.tableWidget_7.setCellWidget(row, col, my_push2)
				else:
					if analyst_data[row][col]:
						self.tableWidget_7.setItem(row, col, QTableWidgetItem(str(analyst_data[row][col])))
				col += 1
			row_pos = self.tableWidget_7.rowCount()
			self.tableWidget_7.insertRow(row_pos)
		self.tableWidget_7.setSortingEnabled(False)
		if from_start:
			from_start = False
		else:
			self.Add_Data_To_history(6, 2)
			# self.History()
		self.tableWidget_7.resizeColumnsToContents()

	def Handel_Save_Delete_analyst(self):
		index = self.sender().pos()
		index = self.tableWidget_7.indexAt(index)
		index = index.row()
		r_id = self.tableWidget_7.item(index, 0).text()
		if self.sender().text() == 'حفظ':
			try:
				name = self.tableWidget_7.item(index, 1).text()
			except:
				name = ''
			category = self.tableWidget_7.cellWidget(index, 2).currentText()
			try:
				default = self.tableWidget_7.item(index, 3).text()
			except:
				default = ''
			try:
				unit = self.tableWidget_7.item(index, 4).text()
			except:
				unit = ''
			price = self.tableWidget_7.cellWidget(index, 5).value()
			sub_category = self.tableWidget_7.cellWidget(index, 6).currentText()
			analyst_index = int(self.tableWidget_7.cellWidget(index, 7).currentText())
			if name != '' and name:
				self.cur.execute(
					''' UPDATE addanalyst set name=%s,price=%s,category=%s,sub_category=%s,unit=%s,defult=%s,analyst_index=%s where id=%s''',
					(name, price, category, sub_category, unit, default, analyst_index, r_id,))
				self.db.commit()
				self.cur.execute(''' select to_analysts from addbuys where id=%s ''',(r_id,))
				data_CBO = self.cur.fetchone()
				if data_CBO:
					self.cur.execute(''' update addbuys set to_analysts=%s where to_analysts=%s ''',(name,data_CBO[0],))
					self.db.commit()
				QMessageBox.information(self, '', 'تم حفظ بيانات التحليل بنجاح')
				self.Add_Data_To_history(4, 2)
				# self.History()
			else:
				QMessageBox.information(self, '', 'يرجى اختيار اسم تحليل صحيح')
		else:
			warning = QMessageBox.warning(self, 'احذر', "هل انت متأكد من انك تريد مسح التحليل {name}؟",
										  QMessageBox.Yes | QMessageBox.No)
			if warning == QMessageBox.Yes:
				self.cur.execute(''' select name from addanalyst where id =%s ''', (r_id,))
				data = self.cur.fetchone()
				self.cur.execute(''' delete from addanalyst  where id =%s''', (r_id,))
				self.cur.execute(''' delete from addbuys where to_analysts = %s''', (data[0],))
				self.db.commit()
				self.tableWidget_7.removeRow(index)
				QMessageBox.information(self, '', 'تم حذف التحليل بنجاح')
				self.Add_Data_To_history(5, 2)
				# self.History()

	def tslsol_wout_b(self, uio=None):
		print(uio, 'jkej')
		if uio == 'in edit':
			self.comboBox_30.clear()
			self.cur.execute(''' select name from addanalyst where sub_category=%s ''',
							 (self.comboBox_26.currentText(),))
			data = self.cur.fetchall()
			for i in range(1, len(data) + 2):
				self.comboBox_30.addItem(str(i))
		else:
			self.comboBox_29.clear()
			self.cur.execute(''' select name from addanalyst where sub_category=%s ''',
							 (self.comboBox_23.currentText(),))
			data = self.cur.fetchall()
			for i2 in range(1, len(data) + 2):
				self.comboBox_29.addItem(str(i2))

	def Close_tslsol(self, typet=None):
		for i in range(0, self.Analyst_Dialog2.tableWidget.rowCount() - 1):
			if typet == 'no':
				self.cur.execute(''' update addanalyst set analyst_index=%s where name=%s and sub_category=%s ''', (
				self.Analyst_Dialog2.tableWidget.cellWidget(i, 1).currentText(),
				self.Analyst_Dialog2.tableWidget.item(i, 0).text(), self.comboBox_23.currentText(),))
			else:
				self.cur.execute(''' update addanalyst set analyst_index=%s where name=%s and sub_category=%s ''', (
				self.Analyst_Dialog2.tableWidget.cellWidget(i, 1).currentText(),
				self.Analyst_Dialog2.tableWidget.item(i, 0).text(), self.comboBox_26.currentText(),))
		self.db.commit()
		self.Leo()
		self.Analyst_Dialog2.close()

	def Show_all_taslsol(self):
		self.Analyst_Dialog2 = add_delete_analyst_choicesrupy2.Dialog()
		self.Analyst_Dialog2.setWindowFlags(Qt.Window | Qt.FramelessWindowHint)
		self.Analyst_Dialog2.tableWidget.setRowCount(0)
		self.Analyst_Dialog2.tableWidget.insertRow(0)
		if self.sender().text() == '.':
			self.Analyst_Dialog2.pushButton_27.clicked.connect(lambda: self.Close_tslsol('from edit'))
		else:
			self.Analyst_Dialog2.pushButton_27.clicked.connect(lambda: self.Close_tslsol('no'))
		self.Analyst_Dialog2.show()
		if self.sender().text() == '.':
			self.cur.execute(
				''' select name,analyst_index from addanalyst where sub_category=%s order by analyst_index ASC ''',
				(self.comboBox_26.currentText(),))
		else:
			self.cur.execute(
				''' select name,analyst_index from addanalyst where sub_category=%s order by analyst_index ASC ''',
				(self.comboBox_23.currentText(),))
		data = self.cur.fetchall()
		mulist = []
		for iq in range(0, len(data)):
			mulist.append(str(iq + 1))
		if data:
			count = 0
			for index, i in enumerate(data):
				self.Analyst_Dialog2.tableWidget.setItem(count, 0, QTableWidgetItem(str(data[index][0])))
				combo = QComboBox()
				combo.addItems(mulist)
				if data[index][1]:
					combo.setCurrentIndex(int(data[index][1]) - 1)
				self.Analyst_Dialog2.tableWidget.setCellWidget(count, 1, combo)
				self.Analyst_Dialog2.tableWidget.insertRow(count + 1)
				count += 1

	def Add_Analyst(self):
		analyst_name = self.lineEdit_28.text()
		if analyst_name not in self.RETURN_ALL_ANALYST():
			analyst_result_category = self.comboBox_22.currentText()
			analyst_price = self.spinBox_5.value()
			categoryes_number = self.comboBox_23.count()
			categoryes = []
			all_results = []
			results_counts = self.comboBox_11.count()
			for i in range(0, results_counts):
				self.comboBox_11.setCurrentIndex(i)
				all_results.append(self.comboBox_11.currentText())
			self.comboBox_11.setCurrentIndex(0)
			sub_category = self.comboBox_23.currentText()
			date = datetime.now()
			defult = self.lineEdit_40.text()
			unit = self.lineEdit_48.text()
			analyst_index = self.comboBox_29.currentText()
			if analyst_index == '' or analyst_index == None:
				analyst_index = 1
			self.cur.execute(
				''' INSERT INTO addanalyst (name,price,category,sub_category,date,defult,unit,results,analyst_index) VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s) ''',
				(str(analyst_name), analyst_price, analyst_result_category, sub_category, date, defult, unit,
				 str(all_results)[1:-1], int(analyst_index)))
			self.db.commit()
			# self.lineEdit_28.setText('')  # analyst_name =
			# self.comboBox_22.setCurrentIndex(0)  # analyst_result_category =
			# self.spinBox_5.setValue(0)  # analyst_price =
			# self.comboBox_23.setCurrentIndex(0)  # sub_category =
			# self.lineEdit_40.setText('')
			# self.lineEdit_48.setText('')
			# QMessageBox.information(self, 'info', 'تم اضافة التحليل بنجاح')
			# self.add_Analyst_to_list()
			# self.add_Analyst_to_list()
			# self.add_Analyst_to_list()
			self.Show_Normal_Dialog()
			self.Auto_complete_combo()
			# self.Show_all_analysts_in_combo()
			self.Show_All_The_Analysts()
			self.Add_Data_To_history(3, 2)
			# self.History()
		else:
			QMessageBox.information(self, '', 'هذا التحليل موجود بالفعل')

	def CLOSE_Dialog(self, in_edit):
		if in_edit == 'Yes':
			self.comboBox_31.clear()
		else:
			self.comboBox_11.clear()
		for i in range(0, self.Analyst_Dialog.tableWidget.rowCount() - 1):
			item = self.Analyst_Dialog.tableWidget.item(i, 0).text()
			if in_edit == 'Yes':
				self.comboBox_31.addItem(item)
			else:
				self.comboBox_11.addItem(item)
		self.Analyst_Dialog.close()

	def Show_analyst_chioces_dialog(self):
		self.Analyst_Dialog = add_delete_analyst_choices.Dialog()
		self.Analyst_Dialog.setWindowFlags(Qt.Window | Qt.FramelessWindowHint)
		self.Analyst_Dialog.tableWidget.insertRow(0)
		if self.sender().text() != '.':
			self.Analyst_Dialog.pushButton_27.clicked.connect(lambda: self.CLOSE_Dialog('No'))
			self.Analyst_Dialog.pushButton_26.clicked.connect(self.ADD_CHOICE)
		else:
			self.Analyst_Dialog.pushButton_27.clicked.connect(lambda: self.CLOSE_Dialog('Yes'))
			self.Analyst_Dialog.pushButton_26.clicked.connect(self.ADD_CHOICE)
			self.cur.execute(''' select results from addanalyst where name=%s''', (self.comboBox_21.currentText(),))
			data = self.cur.fetchone()
			if data:
				x = '[' + str(data[0]) + ']'
				results_list = literal_eval(str(x))
				count = 0
				for i in results_list:
					self.Analyst_Dialog.tableWidget.setItem(count, 0, QTableWidgetItem(str(i)))
					my_push = QPushButton()
					my_push.setText('حفظ')
					my_push.setStyleSheet('border-radius:12px;')
					my_push.clicked.connect(self.Handel_Save_Delete_analyst_choice)
					self.Analyst_Dialog.tableWidget.setCellWidget(count, 1, my_push)
					my_push2 = QPushButton()
					my_push2.setText('حذف')
					my_push2.setStyleSheet(
						'''QPushButton:pressed{background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:0.687, y2:0.704545, stop:0.0646766 rgba(255, 155, 155, 255), stop:0.751244 rgba(235, 54, 30, 255));}QPushButton{border-radius:12px;background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:1, y2:1, stop:0.512438 rgba(235, 54, 30, 255), stop:0.875622 rgba(255, 155, 155, 255));}''')
					my_push2.clicked.connect(self.Handel_Save_Delete_analyst_choice)
					self.Analyst_Dialog.tableWidget.setCellWidget(count, 2, my_push2)
					self.Analyst_Dialog.tableWidget.insertRow(count + 1)
					count += 1
		self.Analyst_Dialog.show()

	def ADD_CHOICE(self):
		text = self.Analyst_Dialog.lineEdit.text()
		row = self.Analyst_Dialog.tableWidget.rowCount()
		self.Analyst_Dialog.tableWidget.setItem(row - 1, 0, QTableWidgetItem(str(text)))
		my_push = QPushButton()
		my_push.setText('حفظ')
		my_push.setStyleSheet('border-radius:12px;')
		my_push.clicked.connect(self.Handel_Save_Delete_analyst_choice)
		self.Analyst_Dialog.tableWidget.setCellWidget(row - 1, 1, my_push)
		my_push2 = QPushButton()
		my_push2.setText('حذف')
		my_push2.setStyleSheet(
			'''QPushButton:pressed{background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:0.687, y2:0.704545, stop:0.0646766 rgba(255, 155, 155, 255), stop:0.751244 rgba(235, 54, 30, 255));}QPushButton{border-radius:12px;background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:1, y2:1, stop:0.512438 rgba(235, 54, 30, 255), stop:0.875622 rgba(255, 155, 155, 255));}''')
		my_push2.clicked.connect(self.Handel_Save_Delete_analyst_choice)
		self.Analyst_Dialog.tableWidget.setCellWidget(row - 1, 2, my_push2)
		self.Analyst_Dialog.tableWidget.insertRow(row)

	def Handel_Save_Delete_analyst_choice(self):
		index = self.sender().pos()
		index = self.Analyst_Dialog.tableWidget.indexAt(index)
		index = index.row()
		r_id = self.Analyst_Dialog.tableWidget.item(index, 0).text()
		if self.sender().text() != 'حفظ':
			self.Analyst_Dialog.tableWidget.removeRow(index)
	def Leo(self):	
		self.cur.execute(''' select analyst_index FROM addanalyst WHERE name=%s ''',(self.comboBox_21.currentText(),))	
		data = self.cur.fetchone()	
		self.comboBox_30.setCurrentIndex(int(data[0])-1)
	def Show_analyst_in_Edit_Or_Delete(self):
		self.comboBox_31.clear()
		analyst_current_name = self.comboBox_21.currentText()
		analyst_current_index = self.comboBox_21.currentIndex()
		if analyst_current_index == 0:
			# QMessageBox.information(self, 'تحذير', "يرجى اختيار تحليل صحيح")
			self.comboBox_26.setCurrentText('')  # sub_category =
			self.comboBox_26.setCurrentIndex(0)  # sub_category =
			self.comboBox_25.setCurrentIndex(0)  # sub_category =
			self.lineEdit_29.setText('')  # analyst_name =
			self.comboBox_25.setCurrentText('')  # analyst_result_category =
			self.lineEdit_39.setText('')  # defult =
			self.lineEdit_38.setText('')  # unit =
			self.spinBox_6.setValue(0)  # analyst_price =
		else:
			self.cur.execute(
				''' SELECT name,defult,unit,price,category,sub_category,results,analyst_index FROM addanalyst WHERE name=%s ''',
				(analyst_current_name,))
			data = self.cur.fetchall()
			if data:
				if data[0][6]:
					x = '[' + str(data[0][6]) + ']'
					list_data = literal_eval(str(x))
					for rr in list_data:
						self.comboBox_31.addItem(rr)  # results =
				self.comboBox_26.setCurrentText(str(data[0][5]))  # sub_category =
				self.lineEdit_29.setText(str(data[0][0]))  # analyst_name =
				self.comboBox_25.setCurrentText(str(data[0][4]))  # analyst_result_category =
				self.lineEdit_39.setText(data[0][1])  # defult =
				self.lineEdit_38.setText(data[0][2])  # unit =
				self.spinBox_6.setValue(data[0][3])  # analyst_price =
				self.comboBox_30.setCurrentIndex(int(data[0][7])-1)
			# self.Add_Data_To_history(6, 2)
			# # self.History()
			else:
				# QMessageBox.information(self, 'تحذير', "يرجى اختيار تحليل صحيح")
				self.comboBox_26.setCurrentText('')  # sub_category =
				self.comboBox_26.setCurrentIndex(0)  # sub_category =
				self.comboBox_25.setCurrentIndex(0)  # sub_category =
				self.lineEdit_29.setText('')  # analyst_name =
				self.comboBox_25.setCurrentText('')  # analyst_result_category =
				self.lineEdit_39.setText('')  # defult =
				self.lineEdit_38.setText('')  # unit =
				self.spinBox_6.setValue(0)  # analyst_price =

	def Update_addNewItemAnalysts(self):
		global analyst_name_for_update
		global analyst_name_for_update_before
		self.cur.execute('''UPDATE addnewitem set analyst_name=%s where analyst_name=%s''',
						 (analyst_name_for_update, analyst_name_for_update_before))
		self.db.commit()

	def Edit_Analyst(self):
		global analyst_name_for_update
		global analyst_name_for_update_before
		analyst_name = self.lineEdit_29.text()
		analyst_name_before = self.comboBox_21.currentText()
		if self.comboBox_21.currentText() in self.RETURN_ALL_ANALYST():
			analyst_result_category = self.comboBox_25.currentText()
			defult = self.lineEdit_39.text()
			unit = self.lineEdit_38.text()
			analyst_price = self.spinBox_6.value()
			sub_category = self.comboBox_26.currentText()
			date = datetime.now()
			analyst_index = int(self.comboBox_30.currentText())
			results_number = self.comboBox_31.count()
			results = []
			for i in range(0, results_number):
				self.comboBox_31.setCurrentIndex(i)
				results.append(str(self.comboBox_31.currentText()))
			mysql = '''UPDATE addanalyst SET name=%s,defult=%s,unit=%s,price=%s,category=%s,sub_category=%s,date=%s,results=%s,analyst_index=%s where name=%s'''
			values = (
				str(analyst_name), defult, unit, analyst_price, analyst_result_category, sub_category, date,
				str(results)[1:-1], analyst_index,
				analyst_name_before)
			self.cur.execute(mysql, values)
			self.cur.execute(''' update addbuys set to_analysts=%s where to_analysts=%s ''',(analyst_name,analyst_name_before))
			self.cur.execute(''' update analystnormal set analyst_name=%s where analyst_name=%s ''',(analyst_name,analyst_name_before))
			self.db.commit()
			analyst_name_for_update = analyst_name
			analyst_name_for_update_before = analyst_name_before
			# self.comboBox_21.setCurrentIndex(0)  # analyst_current_name =
			# self.lineEdit_29.setText('')  # analyst_name =
			# self.comboBox_25.setCurrentIndex(0)  # analyst_result_category =
			# self.lineEdit_39.setText('')  # defult =
			# self.lineEdit_38.setText('')  # unit =
			# self.spinBox_6.setValue(0)  # analyst_price =
			# self.comboBox_26.setCurrentIndex(0)  # sub_category =
			# QMessageBox.information(self, 'info', 'تم تعديل التحليل بنجاح')
			# self.add_Analyst_to_list()
			self.Show_Normal_Dialog('True')
			self.Auto_complete_combo()
			# self.Show_all_analysts_in_combo()
			self.Show_All_The_Analysts()
			self.Add_Data_To_history(4, 2)
			# self.History()
		else:
			QMessageBox.information(self, '', 'هذا التحليل غير موجود')
		self.Update_addNewItemAnalysts()

	def Delete_Analyst(self):
		warning = QMessageBox.warning(self, 'احذر', "هل انت متأكد من انك تريد حذف التحليل",
									  QMessageBox.Yes | QMessageBox.No)
		if warning == QMessageBox.Yes:
			analyst_current_name = self.comboBox_21.currentText()
			sql = ''' DELETE FROM addanalyst WHERE name=%s '''
			self.cur.execute(sql, [(analyst_current_name)])
			self.cur.execute(''' delete from analystnormal where analyst_name=%s ''',(analyst_current_name,))
			self.cur.execute(''' delete from analyst_normal_text where analyst_name=%s ''',(analyst_current_name,))
			self.db.commit()
			QMessageBox.information(self, 'info', 'تم حذف التحليل بنجاح')
			self.Show_all_analysts_in_combo()
			self.Add_Data_To_history(5, 2)
			# self.History()

	def Show_all_analysts_in_combo(self):
		global addTrue
		addTrue = True
		self.cur.execute(
			''' SELECT name FROM addanalyst ORDER BY sub_category''')
		data = self.cur.fetchall()
		self.comboBox_21.clear()
		self.comboBox_16.clear()
		self.comboBox_26.clear()
		self.comboBox_23.clear()
		self.comboBox_26.addItem('----------------')
		self.comboBox_23.addItem('----------------')
		self.comboBox_21.addItem('----------------')
		self.comboBox_16.addItem('----------------')
		self.comboBox_5.addItem('الكل')
		self.comboBox_16.addItems(self.add_all_subCategory_toList())
		self.comboBox_5.addItems(self.add_all_subCategory_toList())
		self.comboBox_26.addItems(self.add_all_subCategory_toList())
		self.comboBox_23.addItems(self.add_all_subCategory_toList())
		for item in data:
			self.comboBox_21.addItem(str(item[0]))
			self.comboBox_16.addItem(str(item[0]))
			self.comboBox_5.addItem(str(item[0]))
		for ii in range(0, self.comboBox_16.count()):
			self.comboBox_16.setCurrentIndex(ii)
			if self.comboBox_16.currentText() in self.add_all_subCategory_toList():
				myitem = self.comboBox_16.model().item(self.comboBox_16.currentIndex(), self.comboBox_16.modelColumn())
				myitem.setCheckState(Qt.Unchecked)
		self.comboBox_16.setCurrentIndex(0)
		addTrue = False

	def Clients_Page(self):
		global search_info_by_date
		self.tableWidget_4.setSortingEnabled(False)
		self.tableWidget_9.setSortingEnabled(False)
		self.tableWidget_4.setRowCount(0)
		self.tableWidget_4.insertRow(0)
		self.tableWidget_9.setRowCount(0)
		self.tableWidget_9.insertRow(0)
		idn = self.spinBox_2.value()
		if search_info_by_date:
			idn = self.searchWidget2.spinBox.value()
			from_date = str(self.searchWidget2.dateEdit_6.date().toPyDate())
			to_date = str(self.searchWidget2.dateEdit_5.date().toPyDate())
			self.cur.execute(
				''' SELECT price,analyst_name,analyst_result,client_name,date FROM addnewitem WHERE client_id=%s AND DATE(date)>=%s AND DATE(date)<=%s''',
				(str(idn), from_date, to_date,))
			client_analyst_data = self.cur.fetchall()
			self.cur.execute(
				''' SELECT client_name,client_age,client_genus,client_doctor FROM addclient WHERE id=%s''',
				(str(idn),))
			client_data = self.cur.fetchall()
			self.cur.execute(''' select latest_price,pushed_price from price_task where client_name=%s AND DATE(date)>=%s AND DATE(date)<=%s ''',(client_data[0][0],from_date, to_date,))
			FCBARCELONA = self.cur.fetchone()
		else:
			self.cur.execute(
				''' SELECT price,analyst_name,analyst_result,client_name,date FROM addnewitem WHERE client_id=%s''',
				(str(idn),))
			client_analyst_data = self.cur.fetchall()
			self.cur.execute(''' SELECT client_name,client_age,client_genus,client_doctor FROM addclient WHERE id=%s''',
							 (str(idn),))
			client_data = self.cur.fetchall()
			self.cur.execute(''' select SUM(latest_price),SUM(pushed_price) from price_task where client_name=%s ''',(client_data[0][0],))
			FCBARCELONA = self.cur.fetchone()
		num = 0
		all_client_analyst = []
		total = 0
		for i in client_analyst_data:
			num += 1
		for price in range(0, num):
			total += client_analyst_data[price][0]
		for j in range(0, num):
			all_client_analyst.append(str(client_analyst_data[j][1]))
		self.tableWidget_4.setRowCount(0)
		self.tableWidget_4.insertRow(0)
		for row, form in enumerate(client_data):
			for col, item in enumerate(form):
				self.tableWidget_4.setItem(row, 4, QTableWidgetItem(str(num)))
				self.tableWidget_4.setItem(row, 5, QTableWidgetItem(str(','.join(all_client_analyst))))
				self.tableWidget_4.setItem(row, 6, QTableWidgetItem(str(FCBARCELONA[0])))
				self.tableWidget_4.setItem(row, 7, QTableWidgetItem(str(FCBARCELONA[1])))
				self.tableWidget_4.setItem(row, 8, QTableWidgetItem(str(int(FCBARCELONA[0])-int(FCBARCELONA[1]))))
				self.tableWidget_4.setItem(row, col, QTableWidgetItem(str(item)))
				col += 1
			row_pos = self.tableWidget_4.rowCount()
			self.tableWidget_4.insertRow(row_pos)
		self.tableWidget_9.setRowCount(0)
		self.tableWidget_9.insertRow(0)
		for row, form in enumerate(client_analyst_data):
			for col, item in enumerate(form):
				if col == 0:
					self.tableWidget_9.setItem(row, col, QTableWidgetItem(str(client_analyst_data[0][3])))
				if col == 1:
					self.tableWidget_9.setItem(row, col, QTableWidgetItem(str(client_analyst_data[row][1])))
				if col == 2:
					self.tableWidget_9.setItem(row, col, QTableWidgetItem(str(client_analyst_data[row][2])))
				if col == 3:
					self.tableWidget_9.setItem(row, col, QTableWidgetItem(str(client_analyst_data[row][4])))
				col += 1
			row_pos = self.tableWidget_9.rowCount()
			self.tableWidget_9.insertRow(row_pos)
		if search_info_by_date:
			self.searchWidget2.close()
		self.Add_Data_To_history(6, 6)
		# self.History()
		search_info_by_date = False
		self.tableWidget_4.setSortingEnabled(True)
		self.tableWidget_9.setSortingEnabled(True)

	def RETURN_ALL_ANALYST(self):
		self.cur.execute(''' select name from addanalyst ''')
		data = self.cur.fetchall()
		analysts = []
		if data:
			for i in data:
				if i[0]:
					analysts.append(i[0])
		return analysts

	def Show_Add_Buys_Complete_Dialog(self):
		global current_group_box
		self.Add_Buys_Dialog = add_buys_completepy.Dialog()
		if self.comboBox_12.currentIndex() == 0:
			current_group_box = 'groupBox_2'
			self.Add_Buys_Dialog.groupBox_2.show()
			self.Add_Buys_Dialog.groupBox.hide()
			self.Add_Buys_Dialog.comboBox_2.addItems(self.RETURN_ALL_ANALYST())
			# for row2 in range(1, self.Add_Buys_Dialog.comboBox_2.count()):
			# 	item = self.Add_Buys_Dialog.comboBox_2.model().item(row2)
			# 	item.setCheckState(Qt.Unchecked)
		else:
			current_group_box = 'groupBox'
			self.Add_Buys_Dialog.groupBox_2.hide()
			self.Add_Buys_Dialog.groupBox.show()
			self.Add_Buys_Dialog.comboBox.addItems(self.RETURN_ALL_ANALYST())
		self.Add_Buys_Dialog.pushButton_16.clicked.connect(self.Add_Buys)
		self.Add_Buys_Dialog.dateEdit.setDate(date.today())
		self.Add_Buys_Dialog.dateEdit_2.setDate(date.today())
		self.Add_Buys_Dialog.dateEdit_4.setDate(date.today())
		self.Add_Buys_Dialog.show()

	def RETURN_ALL_MANDOBS(self):
		self.cur.execute(''' select name from mandob ''')
		data =self.cur.fetchall()
		cc =[]
		if data:
			for i in data:
				cc.append(i[0])
		return cc
	def ADD_MANDOBS_TO_COMBO(self):
		self.comboBox_13.clear()
		self.comboBox_13.addItem('')
		self.comboBox_13.addItems(self.RETURN_ALL_MANDOBS())
	def ADD_OR_UPDATE_Mandob(self):
		if self.MandobDialog.comboBox_11.currentText() in self.RETURN_ALL_MANDOBS():
			self.cur.execute(''' select * from mandob ''')
			data=self.cur.fetchone()
			self.MandobDialog.pushButton_16.setText("حفظ")
			self.MandobDialog.lineEdit_16.setText(str(data[1]))
			self.MandobDialog.lineEdit_13.setText(str(data[2]))
			self.MandobDialog.lineEdit_14.setText(str(data[3]))
			self.MandobDialog.lineEdit_15.setText(str(data[4]))
		else:
			self.MandobDialog.pushButton_16.setText("اضافة")
			self.MandobDialog.lineEdit_16.setText('')
			self.MandobDialog.lineEdit_13.setText('')
			self.MandobDialog.lineEdit_14.setText('')
			self.MandobDialog.lineEdit_15.setText('')
		data =self.cur.fetchall()
	def DELETE_Mandob(self):
		if self.MandobDialog.comboBox_11.currentText() in self.RETURN_ALL_MANDOBS():
			warning = QMessageBox.warning(self, '',
												"هل انت متأكد انك تريد حذف هذا المجهز؟",
												QMessageBox.Yes | QMessageBox.No)
			if warning == QMessageBox.Yes:
				self.cur.execute(''' delete from mandob where name=%s ''',(self.MandobDialog.comboBox_11.currentText()))
				self.db.commit()
				QMessageBox.information(self, 'info', 'تم حذف المجهز بنجاح')
				self.ADD_MANDOBS_TO_COMBO()
		else:
			QMessageBox.information(self, 'info', 'عذرا هذا المجهز غير موجود')
	def INSERT_OR_UPDATE_Mandob(self):
		if self.MandobDialog.pushButton_16.text()=="اضافة":
			self.cur.execute(''' insert into mandob (name,phone,gps,notes,date) values(%s,%s,%s,%s,%s)''',(self.MandobDialog.lineEdit_16.text(),self.MandobDialog.lineEdit_13.text(),self.MandobDialog.lineEdit_14.text(),self.MandobDialog.lineEdit_15.text(),str(datetime.now())))
			QMessageBox.information(self, 'info', 'تم اضافة المجهز بنجاح')
		else:
			self.cur.execute(''' update mandob set name=%s,phone=%s,gps=%s,notes=%s,date=%s where name=%s''',(self.MandobDialog.lineEdit_16.text(),self.MandobDialog.lineEdit_13.text(),self.MandobDialog.lineEdit_14.text(),self.MandobDialog.lineEdit_15.text(),str(datetime.now()),self.MandobDialog.comboBox_11.currentText()))
			QMessageBox.information(self, 'info', 'تم تحديث معلومات المجهز بنجاح')
		self.db.commit()
		self.ADD_MANDOBS_TO_COMBO()
		self.MandobDialog.close()
	def show_mandob_dialog(self):
		self.MandobDialog = mandobuipy.Dialog()
		self.MandobDialog.comboBox_11.addItem(" ")
		self.MandobDialog.comboBox_11.addItems(self.RETURN_ALL_MANDOBS())
		self.MandobDialog.comboBox_11.currentTextChanged.connect(self.ADD_OR_UPDATE_Mandob)
		self.MandobDialog.pushButton_16.clicked.connect(self.INSERT_OR_UPDATE_Mandob)
		self.MandobDialog.pushButton_17.clicked.connect(self.DELETE_Mandob)
		self.MandobDialog.show()

	def Add_Buys(self):
		global current_group_box
		item_name = self.lineEdit_13.text()
		quantity = self.spinBox_3.value()
		signal_item_price = self.spinBox_8.value()
		to_analyst = None
		still_price = None
		notification_date = None
		pushed_price = None
		if current_group_box == 'groupBox_2':
			signal_item_quantity = self.Add_Buys_Dialog.spinBox_9.value() * quantity
			date_create_before = str(self.Add_Buys_Dialog.dateEdit_4.date().toPyDate())
			to_analyst = self.Add_Buys_Dialog.comboBox_2.currentText()

		else:
			signal_item_quantity = self.Add_Buys_Dialog.spinBox_8.value() * quantity
			date_create_before = str(self.Add_Buys_Dialog.dateEdit.date().toPyDate())
			notification_date = str(self.Add_Buys_Dialog.dateEdit_2.date().toPyDate())
			pushed_price = self.Add_Buys_Dialog.spinBox_3.value()
			to_analyst = self.Add_Buys_Dialog.comboBox.currentText()
		item_type = self.comboBox_12.currentText()
		mandob = self.comboBox_13.currentText()
		category = ''
		self.cur.execute(""" select sub_category from addanalyst where name=%s limit 1""",(to_analyst,))
		datax=self.cur.fetchone()
		if datax:
			category = str(datax[0])
		total_item_price = signal_item_price * quantity
		if current_group_box != 'groupBox_2':
			still_price = int(total_item_price) - int(pushed_price)
		self.cur.execute(
			''' INSERT INTO addbuys (item_name,signal_item_price,total_price,buys_type,quantity,item_quantity,to_analysts,still_price,pushed_price,notification_date,mandob,category,date) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)'''
			, (
				item_name, signal_item_price, total_item_price, item_type, quantity, signal_item_quantity,
				to_analyst,
				still_price, pushed_price, notification_date, mandob, category, date_create_before))
		# self.cur.execute(''' update addanalyst set quantity=%s where name=%s''',(signal_item_price,to_analyst,))
		self.db.commit()
		self.cur.execute(''' select sum(item_quantity) from addbuys where to_analysts=%s ''',(to_analyst,))
		ccxdata = self.cur.fetchone()
		self.cur.execute(''' update addanalyst set quantity=%s where name=%s ''',(ccxdata[0],to_analyst,))
		self.db.commit()
		self.Add_Buys_Dialog.close()
		QMessageBox.information(self, '', 'تمت الاضافة الى المشتريات')
		self.Show_all_buys()
		self.lineEdit_13.setText('')
		self.spinBox_3.setValue(0)
		self.spinBox_8.setValue(0)
		self.Add_Data_To_history(3, 4)
		# self.History()
	def Buy_Type_Changed(self):
		senderz =self.sender()
		index = self.tableWidget_3.indexAt(senderz.pos()).row()
		try:
			combo = self.tableWidget_3.cellWidget(index,3)
			if combo.currentText() == 'تم الشراء':
				self.tableWidget_3.cellWidget(index,9).setEnabled(False)
			else:
				self.tableWidget_3.cellWidget(index,9).setEnabled(True)
		except:
			pass
	def Change_Buy_Total_Price(self):
		senderz =self.sender()
		index = self.tableWidget_3.indexAt(senderz.pos()).row()
		try:
			combo = self.tableWidget_3.item(index,4).text()
			combo2 = self.tableWidget_3.cellWidget(index,5).value()
			self.tableWidget_3.item(index,7).setText(str(int(combo)*int(combo2)))
		except Exception as e:
			print(e)
		try:
			combo = self.tableWidget_3.item(index,7).text()
			combo2 = self.tableWidget_3.cellWidget(index,9).value()
			self.tableWidget_3.cellWidget(index,10).setValue(int(combo)-combo2)
			if self.tableWidget_3.cellWidget(index,10).value() == 0:
				self.tableWidget_3.cellWidget(index, 3).setCurrentIndex(0)
			else:
				self.tableWidget_3.cellWidget(index, 3).setCurrentIndex(1)
		except Exception as e:
			print(e)
	def Show_all_buys(self):
		self.cur.execute(
			''' SELECT id,item_name,to_analysts,buys_type,quantity,signal_item_price,item_quantity,total_price,mandob,pushed_price,still_price,date,still_price,still_price FROM addbuys ''')
		data = self.cur.fetchall()
		self.tableWidget_3.setSortingEnabled(False)
		self.tableWidget_3.setRowCount(0)
		self.tableWidget_3.insertRow(0)
		for row, form in enumerate(data):
			for col, item in enumerate(form):
				if col == 2:
					my_combo2 = QComboBox()
					my_combo2.addItems(self.RETURN_ALL_ANALYST())
					my_combo2.addItem("لا يوجد")
					index = my_combo2.findText(str(item), Qt.MatchFixedString)
					my_combo2.setCurrentIndex(index)
					self.tableWidget_3.setCellWidget(row, col, my_combo2)
				elif col == 3:
					my_combo2 = QComboBox()
					my_combo2.addItems(["تم الشراء", "تم الشراء , دُفع مبلغ"])
					index = my_combo2.findText(str(item), Qt.MatchFixedString)
					my_combo2.currentIndexChanged.connect(self.Buy_Type_Changed)
					my_combo2.setCurrentIndex(index)
					self.tableWidget_3.setCellWidget(row, col, my_combo2)
				elif col == 4:
					# my_spin = QSpinBox()
					# my_spin.setMaximum(2147483647)
					# my_spin.setMinimum(0)
					# if item:
					# 	my_spin.setValue(int(item))
					# my_spin.valueChanged.connect(self.Change_Buy_Total_Price)
					# my_spin.setValue(378)
					# my_spin.setValue(int(item))
					# self.tableWidget_3.setCellWidget(row, col, my_spin)
					self.tableWidget_3.setItem(row, col, QTableWidgetItem(str(item)))
					my_item =self.tableWidget_3.item(row, col)
					my_item.setFlags(Qt.ItemIsEditable)
					my_item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable)

				elif col == 5:
					my_spin = QSpinBox()
					my_spin.setMaximum(2147483647)
					my_spin.setMinimum(0)
					if item:
						my_spin.setValue(int(item))
					my_spin.valueChanged.connect(self.Change_Buy_Total_Price)
					my_spin.setValue(378)
					my_spin.setValue(int(item))
					self.tableWidget_3.setCellWidget(row, col, my_spin)
				elif col == 6:
					# my_spin = QSpinBox()
					# my_spin.setMaximum(2147483647)
					# my_spin.setMinimum(0)
					# if item:
					# 	my_spin.setValue(int(item))
					self.tableWidget_3.setItem(row, col, QTableWidgetItem(str(item)))
					my_item =self.tableWidget_3.item(row, col)
					my_item.setFlags(Qt.ItemIsEditable)
					my_item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable)
				elif col == 9:
					my_spin = QSpinBox()
					my_spin.setMaximum(int(data[row][7]))
					my_spin.setMinimum(0)
					if item:
						my_spin.setValue(int(item))
					my_spin.valueChanged.connect(self.Change_Buy_Total_Price)
					
					if item:
						my_spin.setValue(378)
						my_spin.setValue(int(item))
					else:
						my_spin.setValue(0)
					self.tableWidget_3.setCellWidget(row, col, my_spin)
				elif col == 10:
					my_spin = QSpinBox()
					my_spin.setMaximum(2147483647)
					my_spin.setMinimum(0)
					if item:
						my_spin.setValue(int(item))
					my_spin.setEnabled(False)
					self.tableWidget_3.setCellWidget(row, col, my_spin)
				elif col == 7:
					self.tableWidget_3.setItem(row, col, QTableWidgetItem(str(item)))
					my_item = self.tableWidget_3.item(row, col)
					my_item.setFlags(Qt.ItemIsEditable)
					my_item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable)
				elif col == 8:
					my_combo2 = QComboBox()
					self.cur.execute("select name from mandob")
					data_zx = self.cur.fetchall()
					for ibb in data_zx:
						my_combo2.addItem(ibb[0])
					index = my_combo2.findText(str(item), Qt.MatchFixedString)
					my_combo2.setCurrentIndex(index)
					self.tableWidget_3.setCellWidget(row, col, my_combo2)
				elif col == 11:
				    my_date = QDateEdit()
				    my_date.setDate(item)
					# my_date.setCalendarPopup(True)
				    self.tableWidget_3.setCellWidget(row, col, my_date)
				elif col == 12:
					self.cur.execute(''' select sum(item_quantity) from addbuys where to_analysts=%s ''',(data[row][2],))
					dataXCd = self.cur.fetchone()
					cb = ''
					if dataXCd:
						cb = dataXCd[0]
					self.tableWidget_3.setItem(row, col, QTableWidgetItem(str(cb)))
					my_item = self.tableWidget_3.item(row, col)
					my_item.setFlags(Qt.ItemIsEditable)
					my_item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable)
				elif col == 13:
					self.cur.execute(''' select quantity from addanalyst where name=%s ''',(data[row][2],))
					dataXCd = self.cur.fetchone()
					cb = ''
					if dataXCd:
						cb = dataXCd[0]
					self.tableWidget_3.setItem(row, col, QTableWidgetItem(str(cb)))
					my_item = self.tableWidget_3.item(row, col)
					my_item.setFlags(Qt.ItemIsEditable)
					my_item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable)
				else:
					self.tableWidget_3.setItem(row, col, QTableWidgetItem(str(item)))
				col += 1
			row_pos = self.tableWidget_3.rowCount()
			self.tableWidget_3.insertRow(row_pos)
		for rowd in range(0, self.tableWidget_3.rowCount() - 1):
			mypush_button = QPushButton(self)
			mypush_button.setText('حفظ')
			mypush_button.setStyleSheet('border-radius:12px;')
			mypush_button.clicked.connect(self.buttonClicked2Save)
			self.tableWidget_3.setCellWidget(rowd, 14, mypush_button)
			mypush_button = QPushButton(self)
			mypush_button.setText('حذف')
			mypush_button.setStyleSheet(
				'''QPushButton:pressed{background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:0.687, y2:0.704545, stop:0.0646766 rgba(255, 155, 155, 255), stop:0.751244 rgba(235, 54, 30, 255));}QPushButton{border-radius:12px;background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:1, y2:1, stop:0.512438 rgba(235, 54, 30, 255), stop:0.875622 rgba(255, 155, 155, 255));}''')
			mypush_button.clicked.connect(self.buttonClicked2)
			self.tableWidget_3.setCellWidget(rowd, 15, mypush_button)
		self.tableWidget_3.setSortingEnabled(True)
		try:
			for cud in range(0, self.tableWidget_3.rowCount() - 1):
				combo = self.tableWidget_3.cellWidget(cud,3)
				if combo.currentText() == 'تم الشراء':
					self.tableWidget_3.cellWidget(cud,9).setEnabled(False)
				else:
					self.tableWidget_3.cellWidget(cud,9).setEnabled(True)
		except:
			pass
		self.tableWidget_3.resizeColumnsToContents()
		# self.Buy_Type_Changed()

	def buttonClicked2(self):
		button = self.sender()
		self.tableWidget_3.setSortingEnabled(False)
		index = self.tableWidget_3.indexAt(button.pos()).row()
		the_id = self.tableWidget_3.item(index, 0).text()
		to_analysts = self.tableWidget_3.cellWidget(index, 2).currentText()
		self.cur.execute(''' select to_analysts from addbuys where id=%s ''',(the_id,))
		datan = self.cur.fetchone()
		if datan:
			self.cur.execute(''' update addanalyst set quantity=%s where name=%s ''',(None,datan[0],))
		self.cur.execute(''' select sum(item_quantity) from addbuys where to_analysts=%s ''',(to_analysts,))
		ccxdata = self.cur.fetchone()
		if ccxdata:
			self.cur.execute(''' update addanalyst set quantity=%s where name=%s ''',(ccxdata[0],to_analysts,))
		self.cur.execute(""" delete from addbuys where id=%s """, (the_id,))
		self.db.commit()
		self.tableWidget_3.setSortingEnabled(True)
		self.Show_all_buys()
		QMessageBox.information(self, 'info', 'تم حذف العنصر بنجاح')

	def buttonClicked2Save(self):
		button = self.sender()
		self.tableWidget_3.setSortingEnabled(False)
		index = self.tableWidget_3.indexAt(button.pos()).row()
		the_id = self.tableWidget_3.item(index, 0).text()
		item_name = self.tableWidget_3.item(index, 1).text()
		to_analysts = self.tableWidget_3.cellWidget(index, 2).currentText()
		buys_type = self.tableWidget_3.cellWidget(index, 3).currentText()
		quantity = self.tableWidget_3.item(index, 4).text()
		signal_item_price = self.tableWidget_3.cellWidget(index, 5).value()
		# item_quantity = self.tableWidget_3.cellWidget(index, 6).value()# * quantity
		total_price = int(quantity) * int(signal_item_price)
		mandob = self.tableWidget_3.cellWidget(index, 8).currentText()
		pushed_price = self.tableWidget_3.cellWidget(index, 9).value()
		still_price = self.tableWidget_3.cellWidget(index, 10).value()
		date_buy = str(self.tableWidget_3.cellWidget(index, 11).date().toPyDate())
		category = ''
		self.cur.execute(''' select to_analysts from addbuys where id=%s ''',(the_id,))
		datan = self.cur.fetchone()
		if datan:
			self.cur.execute(''' update addanalyst set quantity=%s where name=%s ''',(None,datan[0],))
		self.cur.execute(''' select sum(item_quantity) from addbuys where to_analysts=%s ''',(to_analysts,))
		ccxdata = self.cur.fetchone()
		if ccxdata:
			self.cur.execute(''' update addanalyst set quantity=%s where name=%s ''',(ccxdata[0],to_analysts,))
		self.cur.execute(""" select sub_category from addanalyst where name=%s limit 1""",(to_analysts,))
		datax=self.cur.fetchone()
		if datax:
			category = str(datax[0])
		self.cur.execute(
			''' UPDATE addbuys SET item_name=%s,to_analysts=%s,buys_type=%s,signal_item_price=%s,total_price=%s,mandob=%s,category=%s,pushed_price=%s,still_price=%s,date=%s WHERE id=%s ''',
			(item_name, to_analysts, buys_type, signal_item_price, total_price, mandob,category,
			 pushed_price, still_price, date_buy, the_id,))
		self.db.commit()
		self.cur.execute(''' select sum(item_quantity) from addbuys where to_analysts=%s ''',(to_analysts,))
		ccxdata = self.cur.fetchone()
		if ccxdata:
			self.cur.execute(''' update addanalyst set quantity=%s where name=%s ''',(ccxdata[0],to_analysts,))
		else:
			self.cur.execute(''' update addanalyst set quantity=%s where name=%s ''',(None,to_analysts,))
		self.db.commit()
		self.tableWidget_3.setSortingEnabled(True)
		QMessageBox.information(self, 'info', 'تم حفظ هذا العنصر بنجاح')
	def Show_statics(self):
		self.lineEdit_30.setText('')
		self.lineEdit_33.setText('')
		self.lineEdit_36.setText('')
		self.lineEdit_31.setText('')
		self.lineEdit_32.setText('')
		all_categorys = []
		self.cur.execute(''' SELECT sub_category FROM addanalyst ''')
		datagh = self.cur.fetchall()
		for ijko in datagh:
			if ijko[0] not in all_categorys:
				all_categorys.append(ijko[0])
		from_date = self.dateEdit_6.date().toPyDate()
		to_date = self.dateEdit_5.date().toPyDate()
		item = self.comboBox_5.currentText()
		if self.comboBox_5.currentIndex() == 0:
			self.cur.execute(''' SELECT sum(total_price) FROM addbuys WHERE DATE(date)>=%s AND DATE(date)<=%s''',
							 (str(from_date), str(to_date),))
			data = self.cur.fetchone()
			total_buys_price = 0
			if data:
				if data[0]:
					total_buys_price = data[0]
			self.cur.execute(
				''' select sum(price) AS Price from addnewitem where DATE(date)>=%s and DATE(date)<=%s Group by DATE(date) Order by Price DESC  ''',
				(str(from_date), str(to_date)))
			data2 = self.cur.fetchall()
			total_sales_price = 0
			if data2:
				for i in range(0, len(data2)):
					if data2[i][0]:
						total_sales_price += data2[i][0]
				self.lineEdit_33.setText(str(total_sales_price))
			else:
				self.lineEdit_33.setText('0')
			if data:
				if data[0] != None:
					self.lineEdit_30.setText(str(total_buys_price))
				else:
					self.lineEdit_30.setText('0')
			else:
				self.lineEdit_30.setText('0')
			try:
				self.lineEdit_31.setText(str(total_sales_price - total_buys_price))
				if total_buys_price:
					self.lineEdit_32.setText(
						str(((total_sales_price - total_buys_price) / total_buys_price) * 100) + '%')
				else:
					self.lineEdit_32.setText("0%")
			except Exception as e:
				print(e)
				self.lineEdit_31.setText("0")
				self.lineEdit_32.setText("0%")
			self.cur.execute(''' select id from addnewitem where DATE(date)>=%s and DATE(date)<=%s ''',
							 (str(from_date), str(to_date)))
			data = self.cur.fetchall()
			sales_num = 0
			for iueh in data:
				sales_num += 1
			self.lineEdit_36.setText(str(sales_num))
		else:
			if item not in all_categorys:
				self.cur.execute(''' SELECT sum(total_price) FROM addbuys WHERE DATE(date)>=%s AND DATE(date)<=%s AND to_analysts=%s''',
								 (str(from_date), str(to_date),item,))
				data = self.cur.fetchone()
				total_buys_price = 0
				if data:
					if data[0]:
						total_buys_price = data[0]
				self.cur.execute(
					''' select sum(price) AS Price from addnewitem where DATE(date)>=%s and DATE(date)<=%s AND analyst_name=%s Group by DATE(date) Order by Price DESC  ''',
					(str(from_date), str(to_date), item))
				data2 = self.cur.fetchall()
				total_sales_price = 0
				if data2:
					for i in range(0, len(data2)):
						total_sales_price += data2[i][0]
					self.lineEdit_33.setText(str(total_sales_price))
				else:
					self.lineEdit_33.setText('0')
				if data:
					self.lineEdit_30.setText(str(total_buys_price))
				else:
					self.lineEdit_30.setText('0')
				try:
					self.lineEdit_31.setText(str(total_sales_price - total_buys_price))
					if total_buys_price:
						self.lineEdit_32.setText(
							str(((total_sales_price - total_buys_price) / total_buys_price) * 100) + '%')
					else:
						self.lineEdit_32.setText("0%")
				except Exception as e:
					print(e)
					self.lineEdit_31.setText("0")
					self.lineEdit_32.setText("0%")
				self.cur.execute(
					''' select count(analyst_name) from addnewitem where DATE(date)>=%s and DATE(date)<=%s AND analyst_name=%s ''',
					(str(from_date), str(to_date), item))
				data = self.cur.fetchall()
				if data:
					self.lineEdit_36.setText(str(data[0][0]))
				else:
					self.lineEdit_36.setText('0')
			else:
				self.cur.execute(''' SELECT sum(total_price) FROM addbuys WHERE DATE(date)>=%s AND DATE(date)<=%s AND category=%s''',
								 (str(from_date), str(to_date),item,))
				data = self.cur.fetchone()
				total_buys_price = 0
				if data:
					if data[0]:
						total_buys_price = data[0]
				self.cur.execute(
					''' select sum(price) AS Price from addnewitem where DATE(date)>=%s and DATE(date)<=%s AND sub_category=%s Group by DATE(date) Order by Price DESC  ''',
					(str(from_date), str(to_date), item))
				data2 = self.cur.fetchall()
				total_sales_price = 0
				if data2:
					for i in range(0, len(data2)):
						if data2[i][0]:
							total_sales_price += data2[i][0]
						if total_buys_price:
							self.lineEdit_30.setText(str(total_buys_price))
						else:
							self.lineEdit_30.setText('0')
				if data2:
					self.lineEdit_33.setText(str(total_sales_price))
				else:
					self.lineEdit_33.setText('0')
				try:
					self.lineEdit_31.setText(str(total_sales_price - total_buys_price))
					if total_buys_price:
						self.lineEdit_32.setText(
							str(((total_sales_price - total_buys_price) / total_buys_price) * 100) + '%')
					else:
						self.lineEdit_32.setText("0%")
				except Exception as e:
					print(e)
					self.lineEdit_31.setText("0")
					self.lineEdit_32.setText("0%")
				self.cur.execute(
					''' select count(sub_category) from addnewitem where DATE(date)>=%s and DATE(date)<=%s AND sub_category=%s ''',
					(str(from_date), str(to_date), item))
				data = self.cur.fetchall()
				if data:
					self.lineEdit_36.setText(str(data[0][0]))
				else:
					self.lineEdit_36.setText('0')
		self.Add_Data_To_history(6, 8)
		# self.History()

	def Show_default_statics(self):
		if self.tabWidget_4.currentIndex() == 4:
			try:
				self.cur.execute(
					'''select analyst_name,count(price) AS Price from addnewitem Group by analyst_name Order by Price DESC limit 5''')
				data = self.cur.fetchall()
				self.comboBox_7.clear()
				for i in range(0, len(data)):
					self.comboBox_7.addItem(data[i][0] + '     ' + str(data[i][1]))
				self.cur.execute(
					'''select sub_category,count(sub_category) AS Price from addnewitem Group by sub_category Order by Price DESC''')
				data = self.cur.fetchall()
				self.comboBox_8.clear()
				if data:
					for i7 in range(0, len(data)):
						self.comboBox_8.addItem(data[i7][0] + '       ' + str(data[i7][1]))
				self.cur.execute(
					''' select analyst_name,sum(price) AS Price from addnewitem Group by analyst_name Order by Price DESC limit 5''')
				data = self.cur.fetchall()
				self.comboBox.clear()
				for iqw in range(0, len(data)):
					self.comboBox.addItem(data[iqw][0] + '     ' + str(data[iqw][1]))
				self.cur.execute(
					''' select sub_category,sum(price) AS Price from addnewitem Group by sub_category Order by Price DESC''')
				data = self.cur.fetchall()
				self.comboBox_6.clear()
				for iqw2 in range(0, len(data)):
					self.comboBox_6.addItem(data[iqw2][0] + '     ' + str(data[iqw2][1]))
				self.cur.execute(
					''' select date(date),count(date(date)) AS Price from addnewitem Group by date(date) Order by Price DESC limit 10 ''')
				data = self.cur.fetchall()
				self.comboBox_9.clear()
				for iqw3 in range(0, len(data)):
					self.comboBox_9.addItem(str(data[iqw3][0]) + '     ' + str(data[iqw3][1]))
				self.cur.execute(
					''' select date(date),sum(price) AS Price from addnewitem Group by date(date) Order by Price DESC limit 10 ''')
				data = self.cur.fetchall()
				self.comboBox_10.clear()
				for iqw4 in range(0, len(data)):
					self.comboBox_10.addItem(str(data[iqw4][0]) + '     ' + str(data[iqw4][1]))
			except Exception as e:
				print(e, '98erorr')
	# here def
	def Add_Data_To_history(self, action, table):
		global user_id
		self.cur.execute(
			''' INSERT INTO his VALUES (DEFAULT,%s,%s,%s,%s,%s)''',
			(user_id, action, table, datetime.now(), 1))
		self.db.commit()

	def History(self):
		self.tableWidget_8.setSortingEnabled(False)
		self.cur.execute('''SELECT uid,action,tabled,dates FROM his ORDER BY -dates''')
		analyst_data = self.cur.fetchall()
		self.tableWidget_8.setRowCount(0)
		self.tableWidget_8.insertRow(0)
		for row, form in enumerate(analyst_data):
			for col, item in enumerate(form):
				if col == 0:
					self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(analyst_data[row][0])))
				if col == 1:
					action = ''
					if analyst_data[row][1] == 1:
						action = 'تسجيل الدخول'
					if analyst_data[0][1] == 2:
						action = 'تسجيل الخروج'
					if analyst_data[row][1] == 3:
						action = 'اضافة'
					if analyst_data[row][1] == 4:
						action = 'تعديل'
					if analyst_data[row][1] == 5:
						action = 'حذف'
					if analyst_data[row][1] == 6:
						action = 'بحث'
					if analyst_data[row][1] == 7:
						action = 'طباعة'
					if analyst_data[row][1] == 8:
						action = 'معاينة'
					self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(action)))
				if col == 2:
					tables = ''
					if analyst_data[row][2] == 1:
						tables = 'مبيع يومي'
					if analyst_data[row][2] == 2:
						tables = 'تحليل'
					if analyst_data[row][2] == 3:
						tables = 'مشتريات'
					if analyst_data[row][2] == 4:
						tables = 'مراجعين'
					if analyst_data[row][2] == 5:
						tables = 'مستخدم'
					if analyst_data[row][2] == 6:
						tables = 'طبيب'
					if analyst_data[row][2] == 7:
						tables = 'تقرير'
					if analyst_data[row][2] == 8:
						tables = 'احصائيات'
					self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(tables)))
				if col == 3:
					self.tableWidget_8.setItem(row, col, QTableWidgetItem(str(analyst_data[row][3])))
				col += 1
			row_pos = self.tableWidget_8.rowCount()
			self.tableWidget_8.insertRow(row_pos)
		self.tableWidget_8.setSortingEnabled(True)

	def Log_In_Chieck(self):
		global user_id
		global Edit_Doctor
		global Delete_Doctor
		global edit_employee_check
		global show_clients_check
		user_name = self.lineEdit.text()
		user_password = self.lineEdit_2.text()
		self.cur.execute(''' SELECT id,user_password,user_name FROM adduser WHERE user_name=%s ''', (user_name,))
		data = self.cur.fetchone()
		if data != None:
			if data[2] == user_name and data[1] == user_password:
				self.groupBox.setEnabled(True)
				user_id = data[2]
				try:
					self.cur.execute(f' SELECT * FROM userper WHERE employee_name=%s ', (str(user_name),))
					PerData = self.cur.fetchone()
					if PerData:
						if not PerData[27]:
							self.pushButton_47.setEnabled(False)
						else:
							self.pushButton_47.setEnabled(True)
						if not PerData[28]:
							Edit_Doctor = False
							self.pushButton_48.setEnabled(False)
						else:
							Edit_Doctor = True
							self.pushButton_48.setEnabled(True)
						if not PerData[29]:
							Delete_Doctor = False
							self.pushButton_49.setEnabled(False)
						else:
							Delete_Doctor = True
							self.pushButton_49.setEnabled(True)
						if not PerData[28] and not PerData[29]:
							self.groupBox_13.setEnabled(False)
						else:
							self.groupBox_13.setEnabled(True)
						if not PerData[2]:
							self.pushButton_17.setEnabled(False)
							self.pushButton_44.setEnabled(False)
						else:
							self.pushButton_17.setEnabled(True)
							self.pushButton_44.setEnabled(True)
						if not PerData[3]:
							self.pushButton.setEnabled(False)
						else:
							self.pushButton.setEnabled(True)
						if not PerData[4]:
							self.pushButton_2.setEnabled(False)
						else:
							self.pushButton_2.setEnabled(True)
						if not PerData[5]:
							self.pushButton_23.setEnabled(False)
							self.comboBox.clear()
							self.comboBox_6.clear()
							self.comboBox_7.clear()
							self.comboBox_8.clear()
							self.comboBox_9.clear()
							self.comboBox_10.clear()
						else:
							self.Show_default_statics()
							self.pushButton_23.setEnabled(True)
						if not PerData[6]:
							self.pushButton_5.setEnabled(False)
						else:
							self.pushButton_5.setEnabled(True)
						if not PerData[7]:
							self.pushButton_3.setEnabled(False)
						else:
							self.pushButton_3.setEnabled(True)
						if not PerData[8]:
							self.pushButton_4.setEnabled(False)
						else:
							self.pushButton_4.setEnabled(True)
						if not PerData[9]:
							self.groupBox_4.setEnabled(False)
						else:
							self.groupBox_4.setEnabled(True)
						if not PerData[10]:
							self.pushButton_27.setEnabled(False)
						else:
							self.pushButton_27.setEnabled(True)
						if not PerData[11]:
							self.pushButton_28.setEnabled(False)
						else:
							self.pushButton_28.setEnabled(True)
						if not PerData[10] and not PerData[11]:
							self.groupBox_9.setEnabled(False)
						else:
							self.groupBox_9.setEnabled(True)
						if not PerData[12]:
							self.pushButton_15.setEnabled(False)
							self.pushButton_13.setEnabled(False)
							self.pushButton_9.setEnabled(False)
							self.pushButton_14.setEnabled(False)
							self.pushButton_11.setEnabled(False)
						else:
							self.pushButton_15.setEnabled(True)
							self.pushButton_13.setEnabled(True)
							self.pushButton_9.setEnabled(True)
							self.pushButton_14.setEnabled(True)
							self.pushButton_11.setEnabled(True)
						if not PerData[13]:
							self.pushButton_16.setEnabled(False)
						else:
							self.pushButton_16.setEnabled(True)
						if not PerData[14]:
							self.pushButton_31.setEnabled(False)
							self.groupBox_14.setEnabled(False)
						else:
							self.pushButton_31.setEnabled(True)
							self.groupBox_14.setEnabled(True)
						if not PerData[15]:
							self.pushButton_24.setEnabled(False)
						else:
							self.pushButton_24.setEnabled(True)
						if not PerData[16]:
							self.pushButton_22.setEnabled(False)
						else:
							self.pushButton_22.setEnabled(True)
						if not PerData[17]:
							self.pushButton_53.setEnabled(False)
						else:
							self.pushButton_53.setEnabled(True)
						if not PerData[18]:
							self.pushButton_25.setEnabled(False)
						else:
							self.pushButton_25.setEnabled(True)
						if not PerData[17] and not PerData[18] and not PerData[19]:
							self.pushButton_53.setEnabled(False)
							self.pushButton_25.setEnabled(False)
							self.pushButton_37.setEnabled(False)
							self.pushButton_33.setEnabled(False)
							self.comboBox_3.setEnabled(False)
							self.lineEdit_17.setEnabled(False)
							self.lineEdit_3.setEnabled(False)
							self.checkBox.setEnabled(False)
							self.comboBox_2.setEnabled(False)
						else:
							self.pushButton_53.setEnabled(True)
							self.pushButton_25.setEnabled(True)
							self.pushButton_37.setEnabled(True)
							self.pushButton_33.setEnabled(True)
							self.comboBox_3.setEnabled(True)
							self.lineEdit_17.setEnabled(True)
							self.lineEdit_3.setEnabled(True)
							self.checkBox.setEnabled(True)
							self.comboBox_2.setEnabled(True)
						if not PerData[19]:
							self.pushButton_37.setEnabled(False)
						else:
							self.pushButton_37.setEnabled(True)
						if not PerData[20]:
							self.pushButton_43.setEnabled(False)
						else:
							self.pushButton_43.setEnabled(True)
						if not PerData[21]:
							show_clients_check = False
							self.tableWidget_2.setRowCount(0)
							self.tableWidget_2.insertRow(0)
						else:
							show_clients_check = True
						if not PerData[22]:
							self.pushButton_36.setEnabled(False)
						else:
							self.pushButton_36.setEnabled(True)
						if not PerData[23]:
							self.pushButton_29.setEnabled(False)
							self.pushButton_42.setEnabled(False)
						else:
							self.pushButton_29.setEnabled(True)
							self.pushButton_42.setEnabled(True)
						if not PerData[24]:
							self.pushButton_34.setEnabled(False)
						else:
							self.pushButton_34.setEnabled(True)
						if not PerData[25]:
							self.pushButton_18.setEnabled(False)
						else:
							self.pushButton_18.setEnabled(True)
						if not PerData[26]:
							self.pushButton_36.setEnabled(False)
						else:
							self.pushButton_36.setEnabled(True)
				except Exception as e:
					print(e, '5erorr')
				self.lineEdit.setText('')
				self.lineEdit_2.setText('')
				self.tabWidget.setCurrentIndex(3)
				self.Add_Data_To_history(1, 5)
				# self.History()
				self.cur.execute(""" select to_analysts,total_price,still_price,item_name from addbuys where notification_date=%s """,(str(date.today()),))
				data_Cx = self.cur.fetchall()
				for CN in range(0,len(data_Cx)):
					QMessageBox.information(self,'تنبيه',f'عليك ديون لدفعها {data_Cx[CN][2]} من {data_Cx[CN][1]} للمنتج {data_Cx[CN][3]} الخاص بالتحليل {data_Cx[CN][0]}')

			else:
				warning = QMessageBox.warning(self, '',
											  "كلمة المرور او اسم المستخدم غير صحيحة هل تريد استعادة كلمة المرور؟",
											  QMessageBox.Yes | QMessageBox.No)
				if warning == QMessageBox.Yes:
					self.Open_ResetPassword_Page()
			self.pushButton_44.setEnabled(False)
			self.pushButton_17.setEnabled(False)
			self.comboBox_16.setEnabled(False)
		else:
			warning = QMessageBox.warning(self, '',
										  "كلمة المرور او اسم المستخدم غير صحيحة هل تريد استعادة كلمة المرور؟",
										  QMessageBox.Yes | QMessageBox.No)
			if warning == QMessageBox.Yes:
				self.Open_ResetPassword_Page()

	def Add_all_employee_to_comboBox(self):
		self.comboBox_3.clear()
		self.cur.execute(''' SELECT user_name FROM adduser ''')
		data = self.cur.fetchall()
		self.comboBox_3.addItem('-----------------')
		for i in data:
			self.comboBox_3.addItem(i[0])

	def False_checkState(self):
		self.checkBox_58.setCheckState(False)
		self.checkBox_57.setCheckState(False)
		self.checkBox_56.setCheckState(False)
		self.checkBox_8.setCheckState(False)
		self.checkBox_15.setCheckState(False)
		self.checkBox_9.setCheckState(False)
		self.checkBox_55.setCheckState(False)
		self.checkBox_44.setCheckState(False)
		self.checkBox_40.setCheckState(False)
		self.checkBox_13.setCheckState(False)
		self.checkBox_16.setCheckState(False)
		self.checkBox_14.setCheckState(False)
		self.checkBox_7.setCheckState(False)
		self.checkBox_11.setCheckState(False)
		self.checkBox_12.setCheckState(False)
		self.checkBox_46.setCheckState(False)
		self.checkBox_49.setCheckState(False)
		self.checkBox_50.setCheckState(False)
		self.checkBox_47.setCheckState(False)
		self.checkBox_48.setCheckState(False)
		self.checkBox_51.setCheckState(False)
		self.checkBox_52.setCheckState(False)
		self.checkBox_53.setCheckState(False)
		self.checkBox_54.setCheckState(False)
		self.checkBox_44.setCheckState(False)
		self.checkBox_45.setCheckState(False)
		self.checkBox_42.setCheckState(False)
		self.checkBox_43.setCheckState(False)
		self.checkBox_41.setCheckState(False)

	def Add_employee(self):
		global Edit_employee
		employee_name = self.comboBox_3.currentText()
		employee_password = self.lineEdit_17.text()
		employee_email = self.lineEdit_3.text()
		if self.sender().text() != 'حفظ':
			self.cur.execute(''' INSERT INTO adduser (user_name,user_password,user_email,date) VALUES (%s,%s,%s,%s) ''',
							 (employee_name, employee_password, employee_email, str(datetime.now())))
			self.db.commit()
			QMessageBox.information(self, '', "تم اضافة الموظف بنجاح")
			self.Add_Data_To_history(3, 5)
			# self.History()
		else:
			self.cur.execute(
				''' UPDATE adduser SET user_name=%s,user_password=%s,user_email=%s,date=%s WHERE user_name=%s''',
				(employee_name, employee_password, employee_email, str(datetime.now()), employee_name))
			self.db.commit()
			QMessageBox.information(self, '', "تم تعديل الموظف بنجاح")
			self.Add_Data_To_history(4, 5)
			# self.History()
		self.lineEdit_17.setText('')
		self.lineEdit_3.setText('')
		self.comboBox_3.setCurrentIndex(0)
		self.comboBox_2.setCurrentIndex(0)
		self.frame.show()
		self.Add_all_employee_to_comboBox()
		self.False_checkState()

	def Show_employee_data(self):
		global Edit_employee
		if self.comboBox_3.currentIndex() != 0:
			self.False_checkState()
			# try:
			self.cur.execute(''' SELECT user_name,user_password,user_email FROM adduser WHERE user_name=%s ''',
							 (self.comboBox_3.currentText(),))
			data = self.cur.fetchone()
			if data:
				Edit_employee = True
				self.pushButton_53.hide()
				self.pushButton_25.show()
				self.lineEdit_17.setText(data[1])
				self.lineEdit_3.setText(data[2])
				self.cur.execute(''' SELECT * FROM userper WHERE employee_name=%s ''', (self.comboBox_3.currentText(),))
				perData = self.cur.fetchone()
				if perData:
					if perData[27]:
						self.checkBox_58.setCheckState(True)
					if perData[28]:
						self.checkBox_57.setCheckState(True)
					if perData[29]:
						self.checkBox_56.setCheckState(True)
					if perData[3]:
						self.checkBox_9.setCheckState(True)
					if perData[6]:
						self.checkBox_55.setCheckState(True)
					if perData[4]:
						self.checkBox_40.setCheckState(True)
					if perData[26]:
						self.checkBox_13.setCheckState(True)
					if perData[2]:
						self.checkBox_8.setCheckState(True)
					if perData[24]:
						self.checkBox_15.setCheckState(True)
					if perData[25]:
						self.checkBox_16.setCheckState(True)
					if perData[20]:
						self.checkBox_14.setCheckState(True)
					if perData[21]:
						self.checkBox_7.setCheckState(True)
					if perData[22]:
						self.checkBox_11.setCheckState(True)
					if perData[23]:
						self.checkBox_12.setCheckState(True)
					if perData[8]:
						self.checkBox_46.setCheckState(True)
					if perData[12]:
						self.checkBox_49.setCheckState(True)
					if perData[15]:
						self.checkBox_50.setCheckState(True)
					if perData[14]:
						self.checkBox_47.setCheckState(True)
					if perData[13]:
						self.checkBox_48.setCheckState(True)
					if perData[17]:
						self.checkBox_51.setCheckState(True)
					if perData[18]:
						self.checkBox_52.setCheckState(True)
					if perData[19]:
						self.checkBox_53.setCheckState(True)
					if perData[5]:
						self.checkBox_54.setCheckState(True)
					if perData[7]:
						self.checkBox_44.setCheckState(True)
					if perData[16]:
						self.checkBox_45.setCheckState(True)
					if perData[9]:
						self.checkBox_42.setCheckState(True)
					if perData[10]:
						self.checkBox_43.setCheckState(True)
					if perData[11]:
						self.checkBox_41.setCheckState(True)
			else:
				self.pushButton_25.hide()
				self.pushButton_53.show()
				self.comboBox_2.setCurrentIndex(0)
				self.False_checkState()
				self.lineEdit_17.setText('')
				self.lineEdit_3.setText('')
		else:
			self.comboBox_2.setCurrentIndex(0)
			self.False_checkState()
			self.lineEdit_17.setText('')
			self.lineEdit_3.setText('')

	def Show_permissions(self):
		if self.comboBox_2.currentIndex() != 0:
			if self.comboBox_2.currentIndex() == 1:
				self.frame.hide()
				self.groupBox_21.show()
				self.groupBox_16.hide()
				self.groupBox_20.hide()
				self.groupBox_19.hide()
				self.groupBox_15.hide()
			if self.comboBox_2.currentIndex() == 2:
				self.frame.hide()
				self.groupBox_15.show()
				self.groupBox_21.hide()
				self.groupBox_16.hide()
				self.groupBox_20.hide()
				self.groupBox_19.hide()
			if self.comboBox_2.currentIndex() == 3:
				self.frame.hide()
				self.groupBox_16.show()
				self.groupBox_15.hide()
				self.groupBox_21.hide()
				self.groupBox_20.hide()
				self.groupBox_19.hide()
			if self.comboBox_2.currentIndex() == 4:
				self.frame.hide()
				self.groupBox_19.show()
				self.groupBox_15.hide()
				self.groupBox_21.hide()
				self.groupBox_16.hide()
				self.groupBox_20.hide()
			if self.comboBox_2.currentIndex() == 5:
				self.frame.hide()
				self.groupBox_20.show()
				self.groupBox_15.hide()
				self.groupBox_21.hide()
				self.groupBox_16.hide()
				self.groupBox_19.hide()
		else:
			self.frame.show()

	def Add_permissions(self):
		employee_name = self.comboBox_3.currentText()
		sales_page = 0
		analysts_page = 0
		settings_page = 0
		clients_page = 0
		history_page = 0
		add_sales = 0
		print_report = 0
		prev_report = 0
		search_by_date_in_sales = 0
		add_employee = 0
		edit_employee = 0
		delete_employee = 0
		add_doctor = 0
		edit_doctor = 0
		delete_doctor = 0
		add_analyst = 0
		edit_analyst = 0
		delete_analyst = 0
		show_client_info = 0
		change_path = 0
		change_theme = 0
		edit_report = 0
		add_buys = 0
		statics = 0
		search_in_all_sales = 0
		show_all_clients = 0
		search_client = 0
		delete_all_history = 0
		if self.checkBox_58.isChecked():
			add_doctor = 1
		if self.checkBox_57.isChecked():
			edit_doctor = 1
		if self.checkBox_56.isChecked():
			delete_doctor = 1
		if self.checkBox_9.isChecked():
			sales_page = 1
		if self.checkBox_55.isChecked():
			clients_page = 1
		if self.checkBox_44.isChecked():
			history_page = 1
		if self.checkBox_40.isChecked():
			analysts_page = 1
		if self.checkBox_13.isChecked():
			search_in_all_sales = 1
		if self.checkBox_8.isChecked():
			add_sales = 1
		if self.checkBox_15.isChecked():
			prev_report = 1
		if self.checkBox_16.isChecked():
			print_report = 1
		if self.checkBox_13.isChecked():
			search_in_all_sales = 1
		if self.checkBox_14.isChecked():
			search_by_date_in_sales = 1
		if self.checkBox_7.isChecked():
			show_all_clients = 1
		if self.checkBox_11.isChecked():
			search_client = 1
		if self.checkBox_12.isChecked():
			show_client_info = 1
		if self.checkBox_46.isChecked():
			settings_page = 1
		if self.checkBox_49.isChecked():
			change_theme = 1
		if self.checkBox_50.isChecked():
			change_path = 1
		if self.checkBox_47.isChecked():
			edit_report = 1
		if self.checkBox_48.isChecked():
			add_buys = 1  #
		if self.checkBox_51.isChecked():
			add_employee = 1
		if self.checkBox_52.isChecked():
			edit_employee = 1
		if self.checkBox_53.isChecked():
			delete_employee = 1
		if self.checkBox_54.isChecked():
			statics = 1
		if self.checkBox_44.isChecked():
			history_page = 1
		if self.checkBox_45.isChecked():
			delete_all_history = 1
		if self.checkBox_40.isChecked():
			analysts_page = 1
		if self.checkBox_42.isChecked():
			add_analyst = 1
		if self.checkBox_43.isChecked():
			edit_analyst = 1
		if self.checkBox_41.isChecked():
			delete_analyst = 1
		if self.sender().text() == 'اضافة':
			self.cur.execute(
				''' INSERT INTO userper (employee_name,add_sale_item_page,sales_page,analyst_page,statics,clients_page,history_page,settings_page,add_analyst,edit_analyst,delete_analyst,change_theme,add_buys,edit_report,change_path,delet_history,add_user,edit_user,delete_user,search_by_date_in_sales_page,show_clients,search_client,search_client_info,prev_report,print_report,search_in_all_sales,add_doctor,edit_doctor,delete_doctor) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s) ''',
				(
					employee_name, add_sales, sales_page, analysts_page, statics, clients_page, history_page,
					settings_page,
					add_analyst, edit_analyst, delete_analyst, change_theme, add_buys, edit_report, change_path,
					delete_all_history, add_employee, edit_employee, delete_employee, search_by_date_in_sales,
					show_all_clients, search_client, show_client_info, prev_report, print_report, search_in_all_sales,
					add_doctor, edit_doctor, delete_doctor,))
			self.db.commit()
			QMessageBox.information(self, 'info', 'تم اضافة الموظف بنجاح')
		else:
			self.cur.execute(
				''' UPDATE userper SET employee_name=%s,add_sale_item_page=%s,sales_page=%s,analyst_page=%s,statics=%s,clients_page=%s,history_page=%s,settings_page=%s,add_analyst=%s,edit_analyst=%s,delete_analyst=%s,change_theme=%s,add_buys=%s,edit_report=%s,change_path=%s,delet_history=%s,add_user=%s,edit_user=%s,delete_user=%s,search_by_date_in_sales_page=%s,show_clients=%s,search_client=%s,search_client_info=%s,prev_report=%s,print_report=%s,search_in_all_sales=%s,add_doctor=%s,edit_doctor=%s,delete_doctor=%s WHERE employee_name=%s ''',
				(
					employee_name, add_sales, sales_page, analysts_page, statics, clients_page, history_page,
					settings_page,
					add_analyst, edit_analyst, delete_analyst, change_theme, add_buys, edit_report, change_path,
					delete_all_history, add_employee, edit_employee, delete_employee, search_by_date_in_sales,
					show_all_clients, search_client, show_client_info, prev_report, print_report, search_in_all_sales,
					add_doctor, edit_doctor, delete_doctor,
					self.comboBox_3.currentText()))
			self.db.commit()
			QMessageBox.information(self, 'info', 'تم تعديل بيانات الموظف بنجاح')

	def Delete_employee(self):
		warning = QMessageBox.warning(self, 'احذر', f"هل انت متأكد من انك تريد مسح الموظف {self.comboBox_4.text()}",
									  QMessageBox.Yes | QMessageBox.No)
		if warning == QMessageBox.Yes:
			self.cur.execute(''' DELETE FROM adduser WHERE user_name=%s ''', (self.comboBox_4.text(),))
			self.cur.execute(''' DELETE FROM userper WHERE employee_name=%s ''', (self.comboBox_4.text(),))
			self.db.commit()
			self.False_checkState()
			self.lineEdit_17.setText('')
			self.lineEdit_3.setText('')
			self.comboBox_3.setCurrentIndex(0)
			self.comboBox_2.setCurrentIndex(0)
			self.Add_Data_To_history(5, 5)
			# self.History()
		QMessageBox.information(self, 'info', 'تم حذف الموظف بنجاح')

	def Delete_All_History_Data(self):
		sql = '''DELETE FROM his'''
		self.cur.execute(sql)
		self.db.commit()
		QMessageBox.information(self, 'info', 'تم حذف محتويات السجل بنجاح')
		self.tableWidget_8.setRowCount(0)
		self.tableWidget_8.insertRow(0)

	def Open_Sales_Page(self):
		self.tabWidget.setCurrentIndex(3)

	def Open_Login_Page(self):
		self.tabWidget.setCurrentIndex(0)

	def Open_Settings_Page(self):
		self.tabWidget.setCurrentIndex(6)

	def Open_History_Page(self):
		self.tabWidget.setCurrentIndex(5)
		self.History()

	def Open_clients_Page(self):
		self.tabWidget.setCurrentIndex(2)
	def CC2(self):
		if self.tabWidget_5.currentIndex()==0:
			self.my_def2()
	def Open_Analyse_Page(self):
		self.tabWidget.setCurrentIndex(4)

	def Open_ResetPassword_Page(self):
		self.tabWidget.setCurrentIndex(1)

	def Light_Blue_Theme(self):
		style = open('thems/light_blue.css', 'r')
		style = style.read()
		self.setStyleSheet(style)

	def Dark_Blue_Theme(self):
		style = open('thems/darkblue.css', 'r')
		style = style.read()
		self.setStyleSheet(style)

	def Dark_Gray_Theme(self):
		style = open('thems/darkgray.css', 'r')
		style = style.read()
		self.setStyleSheet(style)

	def Dark_Orange_Theme(self):
		style = open('thems/darkorange.css', 'r')
		style = style.read()
		self.setStyleSheet(style)

	def Dark_Theme(self):
		style = open('thems/qdark.css', 'r')
		style = style.read()
		self.setStyleSheet(style)

	def Bio_Word(self, name, doctor, analysts, results, year, month, day, prev, genus, categorys3):
		global if_print
		global word_files
		global save_word_files
		global word_data
		complete = True
		if prev != "T":
			warning = QMessageBox.warning(self, '', f"سوف تتم الطباعة هل انت متأكد؟",
										  QMessageBox.Yes | QMessageBox.No)
			if warning == QMessageBox.Yes:
				complete = True
			else:
				complete = False
		else:
			system("TASKKILL /F /IM WINWORD.exe")
			system('start WINWORD.exe')
			warning = QMessageBox.warning(self, '', f"هل انت متأكد من المعاينة ؟",
										  QMessageBox.Yes | QMessageBox.No)
			if warning == QMessageBox.Yes:
				complete = True
		if complete:
			self.cur.execute(''' SELECT * FROM doctor WHERE name=%s ''', (doctor,))
			doctor_data = self.cur.fetchone()
			doctor_genus = doctor_data[2]
			files = 0
			units = []
			defults = []
			for iplz in analysts:
				if iplz in categorys3:
					units.append('')
					defults.append('')
				else:
					self.cur.execute(''' SELECT unit FROM addanalyst WHERE name=%s ''', (iplz,))
					myplzdata = self.cur.fetchone()
					self.cur.execute(''' select normal_text from analyst_normal_text where analyst_name=%s and genus=%s ''',(iplz,genus,))
					normal_text_data = self.cur.fetchone()
					if myplzdata:
						if myplzdata[0]:
							units.append(myplzdata[0])
						else:
							units.append('')
					else:
						units.append('')
					if normal_text_data:
						if normal_text_data[0]:
							defults.append(normal_text_data[0])
						else:
							defults.append('')
					else:
						self.cur.execute(''' select defult from addanalyst where name=%s ''',(iplz,))
						normal_text_data = self.cur.fetchone()
						if normal_text_data:
							if normal_text_data[0]:
								defults.append(normal_text_data[0])
							else:
								defults.append('')
						else:
							defults.append('')
			all_files = []
			f = open(r'%s\test-mydocx.docx' % word_files, 'rb')
			f.read()
			document = Document(f)
			all_files.append(document)
			document2 = None
			if len(analysts) > 23:
				files = 2
				f2 = open(r'%s\test-mydocx2.docx' % word_files, 'rb')
				f2.read()
				document2 = Document(f2)
				all_files.append(document2)
			if len(analysts) > 46:
				files = 3
				all_files.clear()
				f2 = open(r'%s\test-mydocx2.docx' % word_files, 'rb')
				f2.read()
				document2 = Document(f2)
				f3 = open(r'%s\test-mydocx3.docx' % word_files, 'rb')
				f3.read()
				document3 = Document(f3)
				all_files.append(document)
				all_files.append(document2)
				all_files.append(document3)
			from_count = 0
			from_count2 = 0
			is_normal_list = []
			Low_or_High = []	
			for It in range(0,len(results)):	
				if analysts[It] in categorys3:	
					is_normal_list.append(True)	
					Low_or_High.append('')	
				else:	
					self.cur.execute(''' select normal_value1,normal_value2,normal_type from analystnormal where analyst_name=%s and genus_type=%s ''',(analysts[It],genus,))	
					data_f = self.cur.fetchone()	
					if data_f:
						if data_f[2]=='number':	
							if results[It] < float(data_f[0]):	
								is_normal_list.append(False)	
								Low_or_High.append(' L')	
							if results[It] > float(data_f[1]):	
								is_normal_list.append(False)	
								Low_or_High.append(' H')	
							# else:	
							# 	is_normal_list.append(True)	
							# 	Low_or_High.append('')	
							if results[It] >= float(data_f[0]) and results[It] <= float(data_f[1]):	
								is_normal_list.append(True)	
								Low_or_High.append('')	
						else:	
							x = '['+str(data_f[0])+']'	
							list_data = literal_eval(str(x))	
							x2 = '['+str(data_f[1])+']'	
							list_data2 = literal_eval(str(x2))	
							try:	
								Xindex = list_data.index(str(results[It]))	
								if list_data2[Xindex] =='غير طبيعي':	
									is_normal_list.append(False)	
								else:	
									is_normal_list.append(True)	
								Low_or_High.append('')	
							except Exception as e:	
								print(e)	
								is_normal_list.append(True)	
								Low_or_High.append('')	
					else:	
						is_normal_list.append(True)	
						Low_or_High.append('')
			for big_item in all_files:
				for table_index, i in enumerate(big_item.tables):
					for row_index, k in enumerate(i.rows):
						for cell_index, j in enumerate(k.cells):
							for n in j.paragraphs:
								if n.text == 'Date     /    /20':
									n.text = f'Date  {year}/{month}/{day}'
									run = n.runs
									font = run[0].font
									font.name = 'Tahoma'
									font.bold = True
									font.size = Pt(12)
								if n.text == 'employee1':
									if word_data[4]:
										n.text = str(word_data[4])
									else:
										n.text = ""
									run = n.runs
									font = run[0].font
									font.name = 'Times New Roman'
									font.color.rgb =RGBColor.from_string('e99d39')
									font.bold = True
									font.size = Pt(11)
								if n.text == 'employee2':
									if word_data[5]:
										n.text = str(word_data[5])
									else:
										n.text = ""
									run = n.runs
									font = run[0].font
									font.name = 'Times New Roman'
									font.color.rgb =RGBColor.from_string('e99d39')
									font.bold = True
									font.size = Pt(11)
								if n.text == 'employeeshahada1':
									if word_data[6]:
										n.text = str(word_data[6])
									else:
										n.text = ""
									run = n.runs
									font = run[0].font
									font.name = 'Times New Roman'
									font.bold = True
									font.color.rgb =RGBColor.from_string('e99d39')
									font.size = Pt(10)
								if n.text == 'employeeshahada2':
									if word_data[7]:
										n.text = str(word_data[7])
									else:
										n.text = ""
									run = n.runs
									font = run[0].font
									font.name = 'Times New Roman'
									font.bold = True
									font.color.rgb =RGBColor.from_string('e99d39')
									font.size = Pt(10)
								if n.text == 'shopname':
									if word_data[1]:
										n.text = str(word_data[1])
									else:
										n.text = ""
									run = n.runs
									font = run[0].font
									font.name = 'Times New Roman'
									font.bold = True
									font.color.rgb =RGBColor.from_string('e99d39')
									font.size = Pt(20)
								if n.text == 'phone1':
									if word_data[2]:
										n.text = str(word_data[2])
									else:
										n.text = ""
									run = n.runs
									font = run[0].font
									font.name = 'Times New Roman'
									font.color.rgb =RGBColor.from_string('e99d39')
									font.bold = True
									font.size = Pt(13)
								if n.text == 'phone2':
									if word_data[3]:
										n.text = str(word_data[3])
									else:
										n.text = ""
									run = n.runs
									font = run[0].font
									font.name = 'Times New Roman'
									font.color.rgb =RGBColor.from_string('e99d39')
									font.bold = True
									font.size = Pt(13)
								if n.text == 'gps':
									if word_data[8]:
										n.text = str(word_data[8])
									else:
										n.text = ""
									run = n.runs
									font = run[0].font
									font.name = 'Times New Roman'
									font.color.rgb =RGBColor.from_string('e99d39')
									font.bold = True
									font.size = Pt(12)
								if n.text == 'client_name':
									self.cur.execute('select report_name from human_type where name=%s',(genus,))
									dataVB = self.cur.fetchone()
									n.text = dataVB[0]
									run = n.runs
									font = run[0].font
									font.name = 'Times New Roman'
									font.bold = True
									font.color.rgb =RGBColor.from_string('1f497d')
									font.size = Pt(12)
								if n.text == 'rname':
									n.text = str(name)
									run = n.runs
									font = run[0].font
									font.name = 'Times New Roman'
									font.bold = True
									font.color.rgb =RGBColor.from_string('1f497d')
									font.size = Pt(12)
								if n.text == 'doctor':
									if doctor_genus == 'female':
										if word_data[15]:
											n.text = f'{word_data[15]}'
										else:
											n.text = ""
									else:
										if word_data[11]:
											n.text = f'{word_data[11]}'
										else:
											n.text = ""
									run = n.runs
									font = run[0].font
									font.name = 'Times New Roman'
									font.color.rgb =RGBColor.from_string('1f497d')
									font.bold = True
									font.size = Pt(12)
								if n.text == 'rdoctor':
									n.text = str(doctor)
									run = n.runs
									font = run[0].font
									font.name = 'Times New Roman'
									font.bold = True
									font.color.rgb =RGBColor.from_string('1f497d')
									font.size = Pt(12)
								if n.text == 'lqb1':
									self.cur.execute('select report_lqb from human_type where name=%s',(genus,))
									dataVB = self.cur.fetchone()
									n.text = dataVB[0]
									run = n.runs
									font = run[0].font
									font.name = 'Times New Roman'
									font.color.rgb =RGBColor.from_string('1f497d')
									font.bold = True
									font.size = Pt(12)
								if n.text == 'lqb2':
									if doctor_genus == 'female':
										if word_data[16]:
											n.text = str(word_data[16])
										else:
											n.text = ""
									else:
										if word_data[12]:
											n.text = str(word_data[12])
										else:
											n.text = ""
									run = n.runs
									font = run[0].font
									font.name = 'Times New Roman'
									font.color.rgb =RGBColor.from_string('1f497d')
									font.bold = True
									font.size = Pt(12)
								# if n.text=='Test' and n.runs[0].font.underline:
								# 	n.text = 'Test'
								# 	font.name = 'Times New Roman'
								# 	font.color.rgb =RGBColor.from_string('76923c')
								# 	font.bold = True
								# 	font.size = Pt(22)
								# if n.text=='Result' and n.runs[0].font.underline:
								# 	n.text = 'Result'
								# 	font.name = 'Times New Roman'
								# 	font.color.rgb =RGBColor.from_string('76923c')
								# 	font.bold = True
								# 	font.size = Pt(22)
								# 	# for igq in range(0, 5):  # here is highlighting
								# 	# 			i.rows[row_index].cells[igq]._tc.get_or_add_tcPr().append(
								# 	# 				parse_xml(r'<w:shd {} w:fill="ffffcc"/>'.format(nsdecls('w'))))
								# if n.text=='Unit' and n.runs[0].font.underline:
								# 	n.text = 'Unit'
								# 	font.name = 'Times New Roman'
								# 	font.color.rgb =RGBColor.from_string('76923c')
								# 	font.bold = True
								# 	font.size = Pt(22)
								# if n.text=='Normal' and n.runs[0].font.underline:
								# 	n.text = 'Normal'
								# 	font.name = 'Times New Roman'
								# 	font.color.rgb =RGBColor.from_string('76923c')
								# 	font.bold = True
								# 	font.size = Pt(22)
								for row in range(from_count, len(analysts)):
									if n.text == str((row + 1) - from_count2):
										n.text = str(analysts[row])
										run = n.runs
										font = run[0].font
										font.bold = True
										if str(n.text) in categorys3:
											# font.highlight_color = WD_COLOR_INDEX.YELLOW
											for igq in range(0, 5):  # here is highlighting
												i.rows[row_index].cells[igq]._tc.get_or_add_tcPr().append(
													parse_xml(r'<w:shd {} w:fill="30b7d1"/>'.format(nsdecls('w'))))
										if str(n.text) in categorys3:
											font.size = Pt(11)
											font.color.rgb =RGBColor.from_string('ffffff')
										else:
											font.size = Pt(10)
										font.name = 'Tahoma'
									if n.text == str((row + 1) - from_count2) + 'r':
										latest_result = ''
										try:
											cc = float(results[row])
											if cc.is_integer():
												latest_result = str(int(cc))
											else:
												latest_result = results[row]
										except Exception as e:
											print(';;;;;', e)
											latest_result = results[row]
										n.text = ' ' + str(latest_result)
										run = n.runs
										font = run[0].font
										font.bold = False
										font.size = Pt(11)
										if is_normal_list[row]:
											font.color.rgb =RGBColor.from_string('548dd4')
										else:
											font.color.rgb =RGBColor.from_string('cf0000')
										font.name = 'Tahoma'
										n2 = n.add_run(Low_or_High[row])	
										n2.bold = False	
										n2.font.name = 'Tahoma'	
										n2.font.size = Pt(11)
										
									if n.text == str((row + 1) - from_count2) + 'unit':
										n.text = str(units[row])
										run1 = n
										font1 = run1.runs[0].font
										font1.bold = True
										font1.size = Pt(9)
										font1.name = 'Times New Roman'
									if n.text == str((row + 1) - from_count2) + 'defult':
										n.text = str(defults[row])
										run1 = n
										font1 = run1.runs[0].font
										font1.bold = True
										font1.size = Pt(9)
										font1.name = 'Times New Roman'
				is_break = False
				for iq1 in big_item.tables:
					if is_break:
						break
					for kq1 in iq1.rows:
						if is_break:
							print('you will dont show')
						for jq1 in kq1.cells:
							for nq1 in jq1.paragraphs:
								if str(nq1.text) == '23':
									nq1.text = ''
								if str(nq1.text) == '23r':
									nq1.text = ''
								if str(nq1.text) == '23unit':
									nq1.text = ''
									font1.size = Pt(10)
								if str(nq1.text) == '23defult' and nq1.runs[0].font.underline:
									nq1.text = ''
									# font1 = nq1.runs[0].font
									# font1.bold = True
									font1.size = Pt(10)
									# font1.name = 'Times New Roman'
									# table=iq1._tbl
									# tr = kq1._tr
									# table.remove(tr)
									break
									is_break = True
								for myq in range(0, 23):
									if nq1.text == str(myq):
										if nq1.runs[0].font.underline:
											nq1.text = ''
											font1.size = Pt(10)
									if nq1.text == str(myq) + 'unit':
										if nq1.runs[0].font.underline:
											nq1.text = ''
											font1.size = Pt(10)
									if nq1.text == str(myq) + 'defult':
										if nq1.runs[0].font.underline:
											nq1.text = ''
											font1.size = Pt(10)
									if nq1.text == str(myq) + 'r':
										if nq1.runs[0].font.underline:
											nq1.text = ''
											font1.size = Pt(10)
				from_count += 22
				from_count2 += 23
			document.save(r'%s\result.docx' % save_word_files)
			f.close()
			if files == 2:
				print('2 files')
				if document2:
					document2.save(r'%s\result2.docx' % save_word_files)
				f2.close()
			if files == 3:
				print('2 files')
				document2.save(r'%s\result2.docx' % save_word_files)
				f2.close()
				print('files3')
				document3.save(r'%s\result3.docx' % save_word_files)
				f3.close()
			try:
				if prev == 'T':
					word = client.Dispatch("Word.Application")
					if path.exists(r'%s\result.docx' % save_word_files):
						word.Documents.Open(r'%s\result.docx' % save_word_files)
					if path.exists(r'%s\result2.docx' % save_word_files):
						word.Documents.Open(r'%s\result2.docx' % save_word_files)
					if path.exists(r'%s\result3.docx' % save_word_files):
						word.Documents.Open(r'%s\result3.docx' % save_word_files)
					self.Add_Data_To_history(8, 7)
					# self.History()
				else:
					system("TASKKILL /F /IM WINWORD.exe")
					word = client.Dispatch("Word.Application")
					if path.exists(r'%s\result.docx' % save_word_files):
						word.Documents.Open(r'%s\result.docx' % save_word_files)
						word.ActiveDocument.PrintOut(Background=False)
					if path.exists(r'%s\result2.docx' % save_word_files):
						word.Documents.Open(r'%s\result2.docx' % save_word_files)
						word.ActiveDocument.PrintOut(Background=False)
					if path.exists(r'%s\result3.docx' % save_word_files):
						word.Documents.Open(r'%s\result3.docx' % save_word_files)
						word.ActiveDocument.PrintOut(Background=False)
					self.Delete_Files()
					self.Add_Data_To_history(7, 7)
					# self.History()
			except Exception as e:
				print(e, '6erorr')


def main():
	app = QApplication(argv)
	app.processEvents()
	window = mainapp()
	window.show()
	app.exec_()


if __name__ == '__main__':
	main()
